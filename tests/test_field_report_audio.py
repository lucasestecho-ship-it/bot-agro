import ast
import html
import json
import re
import unittest
import unicodedata
from pathlib import Path


MAIN_PATH = Path(__file__).parents[1] / "main.py"


def load_main_helpers(extra_names=()):
    """Carga funciones sueltas de main.py sin ejecutar el modulo entero (FastAPI, Telegram, etc.)."""
    source = MAIN_PATH.read_text(encoding="utf-8")
    tree = ast.parse(source)
    helper_names = {
        "clean_inline_markdown",
        "normalize_heading_key",
        "normalize_transcript_text",
        "is_noise_transcript",
        "valid_transcript_blocks",
        "plain_observation_lines",
        "fallback_observations_markdown",
        "general_data_markdown",
        "sanitize_observations",
        "assemble_report_markdown",
        "build_basic_report_markdown",
        "extract_markdown_section",
        "markdown_summary",
        "format_item_line",
        "summarized_item_lines",
        "build_observations_prompt",
    } | set(extra_names)
    helper_nodes = [
        node
        for node in tree.body
        if isinstance(node, ast.FunctionDef) and node.name in helper_names
    ]
    namespace = {
        "re": re,
        "json": json,
        "unicodedata": unicodedata,
        "LIGHT_REPORT_ITEM_LIMIT": 40,
        "NOISE_TRANSCRIPTS": {
            "bye",
            "bye.",
            "thank you",
            "thank you.",
            "thanks",
            "gracias",
            "sin transcripcion disponible",
            "sin transcripción disponible",
        },
    }
    exec(compile(ast.Module(body=helper_nodes, type_ignores=[]), "main.py", "exec"), namespace)
    return source, namespace


def report_section_titles():
    source = MAIN_PATH.read_text(encoding="utf-8")
    tree = ast.parse(source)
    for node in tree.body:
        if isinstance(node, ast.Assign):
            targets = [t.id for t in node.targets if isinstance(t, ast.Name)]
            if "REPORT_SECTION_TITLES" in targets:
                return ast.literal_eval(node.value)
    raise AssertionError("REPORT_SECTION_TITLES no encontrado en main.py")


SESSION = {
    "campo": "Prueba",
    "sector": "Lote 1",
    "started_at": "2026-08-05T21:18:53Z",
    "closed_at": "2026-08-05T22:00:00Z",
}


class FieldReportAudioTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.source, cls.helpers = load_main_helpers()

    def test_short_useful_audio_is_not_discarded(self):
        self.assertFalse(self.helpers["is_noise_transcript"]("prueba"))
        blocks = self.helpers["valid_transcript_blocks"]([
            {"fecha_hora": "2026-08-05T21:19:34Z", "transcript_text": "prueba"},
        ])
        self.assertEqual(["Fecha 2026-08-05T21:19:34Z:\nprueba"], blocks)

    def test_noise_transcripts_are_discarded(self):
        for noise in ("", "   ", "gracias", "Thank you.", "bye"):
            self.assertTrue(self.helpers["is_noise_transcript"](noise), noise)

    def test_fallback_report_keeps_what_was_dictated(self):
        markdown = self.helpers["build_basic_report_markdown"](
            SESSION,
            [{"fecha_hora": "2026-08-05T21:19:34Z", "transcript_text": "el alambrado del lote 4 esta caido"}],
            [],
            [],
        )
        self.assertIn("## Observaciones de la recorrida", markdown)
        self.assertIn("el alambrado del lote 4 esta caido", markdown)

    def test_report_builders_are_not_duplicated(self):
        self.assertEqual(1, self.source.count("def build_basic_report_markdown("))
        self.assertEqual(1, self.source.count("def build_report_markdown("))


class ObservationsSectionTests(unittest.TestCase):
    """El bug original: la seccion existia en el markdown pero se perdia al armar DOCX y PDF."""

    @classmethod
    def setUpClass(cls):
        cls.source, cls.helpers = load_main_helpers()

    def test_every_rendered_section_is_extractable(self):
        markdown = self.helpers["assemble_report_markdown"](
            SESSION, [], [], "Se observo que el dia estaba nublado."
        )
        for title in report_section_titles():
            extracted = self.helpers["extract_markdown_section"](markdown, title)
            self.assertNotEqual(
                "No registrado en la recorrida",
                extracted,
                f"La seccion '{title}' se renderiza pero no se puede extraer del markdown",
            )

    def test_observations_survive_extraction(self):
        markdown = self.helpers["assemble_report_markdown"](
            SESSION, [], [], "Se observo que el dia estaba nublado."
        )
        extracted = self.helpers["extract_markdown_section"](markdown, "Observaciones de la recorrida")
        self.assertIn("Se observo que el dia estaba nublado.", extracted)

    def test_headings_inside_observations_do_not_truncate_the_section(self):
        markdown = self.helpers["assemble_report_markdown"](
            SESSION, [], [], "### Observacion 1\nSe observo alambrado caido.\n### Observacion 2\nSe observo agua acumulada."
        )
        extracted = self.helpers["extract_markdown_section"](markdown, "Observaciones de la recorrida")
        self.assertIn("Se observo alambrado caido.", extracted)
        self.assertIn("Se observo agua acumulada.", extracted)

    def test_summary_uses_the_observations_not_the_metadata(self):
        markdown = self.helpers["assemble_report_markdown"](
            SESSION, [], [], "- Se observo que el dia estaba nublado."
        )
        summary = self.helpers["markdown_summary"](markdown)
        self.assertEqual("Se observo que el dia estaba nublado.", summary)

    def test_report_has_no_removed_sections(self):
        markdown = self.helpers["build_basic_report_markdown"](
            SESSION, [{"fecha_hora": "x", "transcript_text": "esta nublado"}], [], []
        )
        for removed in (
            "Notas de voz registradas",
            "Diagnostico de situacion",
            "Analisis economico",
            "Recomendaciones",
            "Resumen ejecutivo",
        ):
            self.assertNotIn(removed, markdown)

    def test_observations_prompt_carries_the_transcripts_and_the_no_invention_rule(self):
        prompt = self.helpers["build_observations_prompt"](
            SESSION,
            [{"fecha_hora": "2026-08-05T21:19:34Z", "transcript_text": "aca estoy caminando y se ve nublado"}],
            [],
            [],
        )
        self.assertIn("aca estoy caminando y se ve nublado", prompt)
        self.assertIn("No agregues ninguna observacion", prompt)
        self.assertIn("No inventes datos", prompt)


class ReportRenderingTests(unittest.TestCase):
    """Verifica el archivo entregable, no solo el markdown intermedio."""

    @classmethod
    def setUpClass(cls):
        cls.source, cls.helpers = load_main_helpers()
        cls.markdown = cls.helpers["assemble_report_markdown"](
            SESSION, [], [], "Se observo que el dia estaba nublado y que el alambrado del lote 4 estaba caido."
        )

    def test_docx_body_contains_the_observation(self):
        try:
            from docx import Document as DocxDocument
        except ImportError:
            self.skipTest("python-docx no disponible")

        add_markdown_to_doc = self._load_renderer()
        document = DocxDocument()
        for title in report_section_titles():
            document.add_heading(title, level=1)
            add_markdown_to_doc(document, self.helpers["extract_markdown_section"](self.markdown, title))
        text = "\n".join(p.text for p in document.paragraphs)
        self.assertIn("Se observo que el dia estaba nublado", text)
        self.assertIn("alambrado del lote 4 estaba caido", text)

    def test_pdf_story_contains_the_observation(self):
        rendered = []
        for title in report_section_titles():
            section_text = self.helpers["extract_markdown_section"](self.markdown, title)
            for line in section_text.splitlines():
                if line.strip():
                    rendered.append(html.escape(self.helpers["clean_inline_markdown"](line.strip())))
        joined = "\n".join(rendered)
        self.assertIn("Se observo que el dia estaba nublado", joined)
        self.assertIn("alambrado del lote 4 estaba caido", joined)

    def _load_renderer(self):
        tree = ast.parse(self.source)
        nodes = [
            node for node in tree.body
            if isinstance(node, ast.FunctionDef) and node.name in {"add_markdown_to_doc", "clean_inline_markdown", "normalize_heading_key"}
        ]
        namespace = {"re": re, "unicodedata": unicodedata}
        exec(compile(ast.Module(body=nodes, type_ignores=[]), "main.py", "exec"), namespace)
        return namespace["add_markdown_to_doc"]


if __name__ == "__main__":
    unittest.main()
