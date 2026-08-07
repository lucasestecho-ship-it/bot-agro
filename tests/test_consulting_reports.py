import tempfile
import unittest
import json
from pathlib import Path

from agent_crew import AgentCrew
from capataz import CapatazStore, heuristic_analysis
from consulting_reports import generate_consulting_report
from report_playbooks import detect_report_playbook
from document_intake import (
    LUCAS_HIGHLIGHT_END,
    LUCAS_HIGHLIGHT_START,
    extract_office_document,
)


class _FakeCompletions:
    def __init__(self, content):
        self.content = content
        self.requests = []

    def create(self, **kwargs):
        self.requests.append(kwargs)
        message = type("Message", (), {"content": self.content})()
        choice = type("Choice", (), {"message": message})()
        return type("Response", (), {"choices": [choice]})()


class _FakeOpenAI:
    def __init__(self, content):
        self.chat = type("Chat", (), {})()
        self.chat.completions = _FakeCompletions(content)


class ReportPlaybookTests(unittest.TestCase):
    def test_detects_each_professional_deliverable(self):
        examples = {
            "proyecto_agua": "Haceme un informe de agua y dimensiona la red de bebederos",
            "propuesta_comercial": "Arma una propuesta tecnica y comercial para el cliente",
            "presupuesto_profesional": "Necesito un presupuesto profesional de mi trabajo",
            "comparativo_presupuestos": "Comparar presupuestos de tres proveedores",
            "evaluacion_compra": "Informe de riesgo de compra de este campo",
            "comparativo_campos": "Hacer comparacion de islas contra La Tigra",
            "dossier_venta": "Preparar un dossier de venta del campo",
            "informe_recorrida": "Cerrar con informe de recorrida",
            "informe_tecnico": "Hacer un informe tecnico agronomico",
        }
        for expected, text in examples.items():
            with self.subTest(text=text):
                self.assertEqual(detect_report_playbook(text).key, expected)
        self.assertEqual(
            detect_report_playbook("Comparame estas ofertas y decime cual proveedor conviene").key,
            "comparativo_presupuestos",
        )
        self.assertEqual(
            detect_report_playbook("Haceme un informe mejorando la redaccion de esta recorrida").key,
            "informe_recorrida",
        )

    def test_deliverable_routes_the_complete_specialist_crew(self):
        from unittest.mock import patch
        env = patch.dict("os.environ", {"CAPATAZ_SUSPENDED_AGENTS": ""})
        env.start()
        self.addCleanup(env.stop)
        with tempfile.TemporaryDirectory() as temp_dir:
            crew = AgentCrew(CapatazStore(data_dir=Path(temp_dir)))
            draft = heuristic_analysis("Comparar presupuestos de bombas para La Susana")
            route = crew.route(draft, draft["summary"])
            self.assertEqual(route, ["tero", "margen", "comercial", "informes"])


class ConsultingReportGenerationTests(unittest.TestCase):
    def test_extracts_docx_and_xlsx_for_budget_comparisons(self):
        try:
            from openpyxl import Workbook
            from docx import Document
        except ImportError:
            self.skipTest("Dependencias de oficina no instaladas")
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            docx_path = root / "oferta.docx"
            document = Document()
            document.add_paragraph("Proveedor Uno - oferta vigente")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Bomba"
            table.cell(0, 1).text = "USD 1200"
            document.save(docx_path)
            self.assertIn("USD 1200", extract_office_document(docx_path, docx_path.name))

            xlsx_path = root / "comparativo.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Ofertas"
            sheet.append(["Proveedor", "Subtotal", "IVA"])
            sheet.append(["Uno", 1000, "=B2*0.21"])
            workbook.save(xlsx_path)
            extracted = extract_office_document(xlsx_path, xlsx_path.name)
            self.assertIn("--- Hoja: Ofertas ---", extracted)
            self.assertIn("=B2*0.21", extracted)

    def test_preserves_lucas_yellow_highlights_in_docx(self):
        try:
            from docx import Document
            from docx.enum.text import WD_COLOR_INDEX
        except ImportError:
            self.skipTest("Dependencias de documentos no instaladas")
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "recorrida.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("Dani registro la observacion. ")
            lucas = paragraph.add_run("Lucas agrego el criterio tecnico.")
            lucas.font.highlight_color = WD_COLOR_INDEX.YELLOW
            document.save(path)
            extracted = extract_office_document(path, path.name)
            self.assertIn("Dani registro la observacion", extracted)
            self.assertIn(LUCAS_HIGHLIGHT_START, extracted)
            self.assertIn("Lucas agrego el criterio tecnico", extracted)
            self.assertIn(LUCAS_HIGHLIGHT_END, extracted)

    def test_generates_real_pdf_and_docx_and_marks_missing_data(self):
        try:
            from pypdf import PdfReader
            from docx import Document
        except ImportError:
            self.skipTest("Dependencias de documentos no instaladas")
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = (
                "Hacer una propuesta tecnica para Estancias Demo. "
                "Objetivo: diagnosticar el sistema de agua. Todavia no se definieron honorarios."
            )
            event = {
                "id": "event-demo",
                "client_name": "Estancias Demo",
                "source": "telegram",
            }
            draft = heuristic_analysis(source, field_name="Campo Demo", source="telegram")
            crew_result = {
                "runs": [
                    {
                        "agent": "Comercial",
                        "output": {
                            "summary": "El cliente necesita diagnostico de agua.",
                            "findings": ["El objetivo fue expresamente informado."],
                            "recommendations": ["Definir alcance y honorarios antes de enviar."],
                            "risks": ["Alcance todavia abierto."],
                            "missing_data": ["Honorarios y forma de pago."],
                        },
                    }
                ],
                "decision": {
                    "summary": "Preparar una propuesta preliminar sin completar precios.",
                    "recommendation": "Confirmar alcance, honorarios y movilidad.",
                    "risks": ["No enviar sin revision de Lucas."],
                    "missing_data": ["Honorarios y forma de pago."],
                },
            }
            report = generate_consulting_report(
                event=event,
                draft=draft,
                crew_result=crew_result,
                source_text=source,
                assets=[],
                output_dir=root,
                logo_path=str(Path(__file__).parents[1] / "static" / "logo.png"),
                allow_fallback=True,
            )
            self.assertIsNotNone(report)
            self.assertEqual(report.playbook_key, "propuesta_comercial")
            self.assertEqual(report.status, "preliminar")
            self.assertTrue(Path(report.pdf_path).exists())
            self.assertTrue(Path(report.docx_path).exists())
            pdf_text = "\n".join(
                page.extract_text() or "" for page in PdfReader(report.pdf_path).pages
            )
            docx_text = "\n".join(
                paragraph.text for paragraph in Document(report.docx_path).paragraphs
            )
            # El entregable es para el cliente: nada de auditoria interna.
            self.assertNotIn("Datos y verificaciones pendientes", pdf_text)
            self.assertNotIn("Riesgos", pdf_text)
            self.assertNotIn("PRELIMINAR", pdf_text)
            self.assertNotIn("Contenido pendiente de completar", pdf_text)
            self.assertNotIn("Datos y verificaciones pendientes", docx_text)
            self.assertNotIn("$1.650.000", pdf_text)
            # Lo interno sigue disponible para Lucas en el payload.
            self.assertTrue(report.missing_data)

    def test_report_fails_closed_without_openai(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            with self.assertRaisesRegex(RuntimeError, "falta un cliente OpenAI"):
                generate_consulting_report(
                    event={"id": "event-fail", "client_name": "LMM"},
                    draft=heuristic_analysis("Hacer informe de recorrida LMM"),
                    crew_result={"runs": [], "decision": None},
                    source_text="Hacer informe de recorrida LMM",
                    assets=[],
                    output_dir=temp_dir,
                )

    def test_sol_is_used_for_final_report_with_reasoning_and_no_temperature(self):
        response = {
            "title": "Informe de recorrida - LMM",
            "subtitle": "Recorrida auditada",
            "executive_summary": "Resumen basado en la nota.",
            "recommendation": "Validar responsables.",
            "status": "preliminar",
            "sections": [],
            "calculations": [],
            "risks": [],
            "missing_data": ["Responsables"],
            "sources": ["Nota de recorrida"],
        }
        client = _FakeOpenAI(json.dumps(response))
        with tempfile.TemporaryDirectory() as temp_dir:
            report = generate_consulting_report(
                event={"id": "event-sol", "client_name": "LMM"},
                draft=heuristic_analysis("Hacer informe de recorrida LMM"),
                crew_result={"runs": [], "decision": None},
                source_text="Hacer informe de recorrida LMM",
                assets=[],
                output_dir=temp_dir,
                openai_client=client,
            )
        request = client.chat.completions.requests[0]
        self.assertEqual(report.model, "gpt-5.6-sol")
        self.assertEqual(request["model"], "gpt-5.6-sol")
        self.assertEqual(request["reasoning_effort"], "high")
        self.assertNotIn("temperature", request)

    def test_docx_image_is_sent_to_sol_and_kept_in_pdf_and_word(self):
        try:
            from docx import Document
            from pypdf import PdfReader
        except ImportError:
            self.skipTest("Dependencias de documentos no instaladas")
        response = {
            "title": "Informe de recorrida - LMM",
            "subtitle": "Recorrida auditada",
            "executive_summary": "Se interpreto la nota y la figura adjunta.",
            "recommendation": "Validar la lectura.",
            "status": "preliminar",
            "sections": [],
            "calculations": [],
            "risks": [],
            "missing_data": ["Validacion"],
            "sources": ["Word de recorrida"],
        }
        client = _FakeOpenAI(json.dumps(response))
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source_docx = root / "recorrida.docx"
            document = Document()
            document.add_paragraph("Informe de recorrida LMM")
            document.add_picture(str(Path(__file__).parents[1] / "static" / "logo.png"))
            document.save(source_docx)
            report = generate_consulting_report(
                event={"id": "event-image", "client_name": "LMM"},
                draft=heuristic_analysis("Hacer informe de recorrida LMM"),
                crew_result={"runs": [], "decision": None},
                source_text=extract_office_document(source_docx, source_docx.name),
                assets=[{"path": str(source_docx), "file_name": source_docx.name}],
                output_dir=root / "out",
                openai_client=client,
            )
            request_content = client.chat.completions.requests[0]["messages"][0]["content"]
            self.assertIsInstance(request_content, list)
            self.assertTrue(any(part.get("type") == "image_url" for part in request_content))
            pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(report.pdf_path).pages)
            docx_text = "\n".join(
                paragraph.text for paragraph in Document(report.docx_path).paragraphs
            )
            self.assertIn("Figuras del documento fuente", pdf_text)
            self.assertIn("Figuras del documento fuente", docx_text)

    def test_no_document_is_generated_for_a_simple_reminder(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = "La Susana: llamar mañana"
            report = generate_consulting_report(
                event={"id": "event-simple", "client_name": "La Susana"},
                draft=heuristic_analysis(source),
                crew_result={"runs": [], "decision": None},
                source_text=source,
                assets=[],
                output_dir=temp_dir,
            )
            self.assertIsNone(report)


if __name__ == "__main__":
    unittest.main()
