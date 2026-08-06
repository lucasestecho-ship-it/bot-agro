import os
import asyncio
import hashlib
import json
import logging
import tempfile
import base64
import html
import mimetypes
import shutil
import uuid
import re
import unicodedata
import secrets
from contextlib import asynccontextmanager
from pathlib import Path
import requests
from fastapi import BackgroundTasks, Body, FastAPI, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from telegram import Update
from telegram.ext import Application, MessageHandler, CommandHandler, filters, ContextTypes
from openai import OpenAI
import gspread
from google.oauth2.service_account import Credentials
from datetime import datetime, timedelta
import fitz
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
try:
    from PIL import Image
except ImportError:
    Image = None

from agent_crew import AgentCrew
from archive_manager import ArchiveManager
from capataz import (
    CapatazStore,
    PersistentStorageError,
    analyze_intake,
    argentina_now,
    iso_now,
    normalize_key,
)
from gmail_drafts import EmailDraftManager, GmailDraftService
from geospatial_worker import (
    GeoAsset,
    analyze_geospatial_package,
    cdse_configuration_status,
    is_geospatial_filename,
)
from push_notifications import PushNotifier
from consulting_reports import generate_consulting_report
from report_playbooks import public_report_catalog
from document_intake import extract_office_document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

TELEGRAM_TOKEN = os.environ.get("TELEGRAM_TOKEN")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
GOOGLE_SHEET_ID = os.environ.get("GOOGLE_SHEET_ID")
GOOGLE_CREDENTIALS_JSON = os.environ.get("GOOGLE_CREDENTIALS_JSON")
SUPABASE_URL = (os.environ.get("SUPABASE_URL") or "").rstrip("/")
SUPABASE_SERVICE_ROLE_KEY = os.environ.get("SUPABASE_SERVICE_ROLE_KEY")
SUPABASE_BUCKET = os.environ.get("SUPABASE_BUCKET")
FIELD_APP_TOKEN = str(os.environ.get("FIELD_APP_TOKEN") or "").strip()
PRIVATE_API_REQUIRES_TOKEN = bool(
    os.environ.get("RENDER") or SUPABASE_SERVICE_ROLE_KEY or OPENAI_API_KEY
)
MY_CHAT_ID = int(os.environ.get("MY_CHAT_ID", "1144480769"))
DATA_DIR = Path(os.environ.get("DATA_DIR", "/tmp/campo_bot")).resolve()
FIELD_ITEMS_DIR = DATA_DIR / "field_items"
FIELD_SESSIONS_DIR = DATA_DIR / "field_sessions"
STATIC_DIR = Path(__file__).resolve().parent / "static"
LEGACY_SESSION_PREFIX = "legacy:"
SAFE_RECORD_ID_RE = re.compile(r"^[A-Za-z0-9_-]{1,128}$")

capataz_store = CapatazStore(
    supabase_url=SUPABASE_URL,
    service_role_key=SUPABASE_SERVICE_ROLE_KEY,
    data_dir=DATA_DIR,
)

FIELD_ITEMS_COLUMNS = [
    "id", "tipo", "campo", "sector", "fecha_hora", "latitud", "longitud",
    "precision_gps", "nombre_archivo", "estado", "storage_status",
    "storage_provider", "storage_path", "storage_public_url", "storage_error",
    "session_id", "photo_label", "audio_label", "transcript_status", "transcript_text", "transcript_error",
    "transcript_model", "transcript_at", "created_at",
]
FIELD_SESSIONS_COLUMNS = [
    "id", "nombre", "campo", "sector", "estado", "started_at", "closed_at",
    "latitud_inicio", "longitud_inicio", "precision_gps_inicio", "notas",
    "created_at", "updated_at",
]
FIELD_REPORTS_COLUMNS = [
    "id", "session_id", "estado", "titulo", "resumen", "informe_markdown",
    "docx_storage_path", "docx_public_url", "pdf_storage_path", "pdf_public_url",
    "error", "progress_message",
    "started_at", "finished_at", "created_at", "updated_at",
]
MAX_REPORT_PHOTOS = 8
LIGHT_REPORT_ITEM_LIMIT = 30
PHOTO_PROMPT_LIMIT = 12
MAX_SOURCE_PHOTO_BYTES = 12 * 1024 * 1024
REPORT_DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
REPORT_PDF_CONTENT_TYPE = "application/pdf"
MAX_AUDIO_UPLOAD_BYTES = 30 * 1024 * 1024
MAX_PHOTO_UPLOAD_BYTES = 20 * 1024 * 1024
REPORT_SECTION_TITLES = [
    "Diagnostico de situacion",
    "Observaciones principales",
    "Analisis economico para la decision",
    "Recomendaciones",
]
NOISE_TRANSCRIPTS = {
    "bye",
    "bye.",
    "thank you",
    "thank you.",
    "thanks",
    "gracias",
    "sin transcripcion disponible",
    "sin transcripción disponible",
}

openai_client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
agent_crew = AgentCrew(capataz_store, openai_client=openai_client)
push_notifier = PushNotifier(capataz_store)
gmail_service = GmailDraftService()
email_draft_manager = EmailDraftManager(
    capataz_store,
    gmail_service=gmail_service,
    openai_client=openai_client,
)
archive_manager = ArchiveManager(
    supabase_url=SUPABASE_URL,
    service_role_key=SUPABASE_SERVICE_ROLE_KEY,
    bucket=SUPABASE_BUCKET,
)

def get_openai_client():
    if not openai_client:
        raise RuntimeError("OPENAI_API_KEY no configurado")
    return openai_client

# Estado en memoria de recorridas activas por chat_id
# { chat_id: {"campo": str, "inicio": datetime, "items": [ {tipo, texto, foto_path} ]} }
recorridas_activas = {}

# Paquetes geoespaciales que Telegram entrega como varios mensajes con el mismo
# media group. Se mantienen solo los segundos necesarios para agruparlos; los
# originales persistentes se suben a Supabase una vez creado el evento.
telegram_geo_batches = {}
# Documentos comunes (presupuestos, propuestas, informes) tambien pueden llegar
# en varios mensajes. Se agrupan para que Tero/Margen comparen el conjunto y no
# procesen cada oferta como si fuera un trabajo aislado.
telegram_document_batches = {}

TELEGRAM_WEBHOOK_PATH = "/telegram/webhook"


def telegram_webhook_base_url():
    """Return the public HTTPS origin used by Telegram to wake the web service."""
    configured = str(os.environ.get("TELEGRAM_WEBHOOK_URL") or "").strip().rstrip("/")
    if configured:
        return configured
    render_url = str(os.environ.get("RENDER_EXTERNAL_URL") or "").strip().rstrip("/")
    if render_url:
        return render_url
    render_hostname = str(os.environ.get("RENDER_EXTERNAL_HOSTNAME") or "").strip().strip("/")
    if render_hostname:
        return f"https://{render_hostname}"
    return ""


def telegram_webhook_secret():
    """Use an explicit secret or derive a stable one without exposing the bot token."""
    configured = str(os.environ.get("TELEGRAM_WEBHOOK_SECRET") or "").strip()
    if configured:
        return configured
    if not TELEGRAM_TOKEN:
        return ""
    return hashlib.sha256(f"capataz-campo:{TELEGRAM_TOKEN}".encode("utf-8")).hexdigest()

SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"
]

def get_google_creds():
    creds_dict = json.loads(GOOGLE_CREDENTIALS_JSON)
    return Credentials.from_service_account_info(creds_dict, scopes=SCOPES)

def get_google_sheet():
    creds = get_google_creds()
    gc = gspread.authorize(creds)
    return gc.open_by_key(GOOGLE_SHEET_ID)

def supabase_response_error(response, url):
    return f"Supabase Storage ERROR status={response.status_code} url={url} body={response.text}"

def upload_field_file_to_supabase(file_path, storage_path, content_type=None, upsert=False):
    if not (SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY and SUPABASE_BUCKET):
        return None

    logger.info(f"Subiendo a Supabase Storage: {storage_path}")
    upload_url = f"{SUPABASE_URL}/storage/v1/object/{SUPABASE_BUCKET}/{storage_path}"
    headers = {
        "apikey": SUPABASE_SERVICE_ROLE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
        "Content-Type": content_type or "application/octet-stream",
        "x-upsert": "true" if upsert else "false",
    }
    with open(file_path, "rb") as f:
        response = requests.post(upload_url, headers=headers, data=f, timeout=60)
    if not response.ok:
        raise RuntimeError(supabase_response_error(response, upload_url))
    public_url = f"{SUPABASE_URL}/storage/v1/object/public/{SUPABASE_BUCKET}/{storage_path}"
    return {"path": storage_path, "public_url": public_url}

def download_supabase_storage_file(storage_path, destination_path):
    if not (SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY and SUPABASE_BUCKET):
        raise RuntimeError("Supabase Storage no configurado")

    url = f"{SUPABASE_URL}/storage/v1/object/{SUPABASE_BUCKET}/{storage_path}"
    headers = {
        "apikey": SUPABASE_SERVICE_ROLE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
    }
    response = requests.get(url, headers=headers, timeout=60, stream=True)
    response.raise_for_status()
    destination_path.parent.mkdir(parents=True, exist_ok=True)
    with destination_path.open("wb") as output:
        for chunk in response.iter_content(chunk_size=1024 * 1024):
            if chunk:
                output.write(chunk)
    return destination_path

def safe_storage_segment(value, fallback):
    text = (value or "").strip().lower()
    text = re.sub(r"[^a-z0-9._-]+", "-", text)
    text = text.strip("-._")
    return text or fallback

def clean_report_filename_segment(value, fallback="Sin_Campo", max_length=60):
    text = str(value or "").strip()
    if not text:
        return fallback
    text = "".join(
        char for char in unicodedata.normalize("NFKD", text)
        if not unicodedata.combining(char)
    )
    text = re.sub(r"[^A-Za-z0-9]+", " ", text).strip()
    if not text:
        return fallback
    words = [word[:1].upper() + word[1:] for word in text.split()]
    cleaned = "_".join(words)
    return cleaned[:max_length].strip("_") or fallback

def encode_legacy_session_id(campo, sector):
    payload = json.dumps(
        {"campo": campo or "", "sector": sector or ""},
        ensure_ascii=False,
        separators=(",", ":"),
    ).encode("utf-8")
    encoded = base64.urlsafe_b64encode(payload).decode("ascii").rstrip("=")
    return f"{LEGACY_SESSION_PREFIX}{encoded}"

def decode_legacy_session_id(session_id):
    if not session_id.startswith(LEGACY_SESSION_PREFIX):
        return None
    encoded = session_id[len(LEGACY_SESSION_PREFIX):]
    padding = "=" * (-len(encoded) % 4)
    try:
        return json.loads(base64.urlsafe_b64decode((encoded + padding).encode("ascii")).decode("utf-8"))
    except Exception:
        return None

def is_legacy_session_id(session_id):
    return bool(decode_legacy_session_id(session_id or ""))

def fetch_supabase_rows(table, columns, order=None, limit=None):
    select = ",".join(columns)
    url = f"{SUPABASE_URL}/rest/v1/{table}?select={select}"
    if order:
        url += f"&order={order}"
    if limit is not None:
        url += f"&limit={limit}"
    response = requests.get(url, headers=supabase_headers(), timeout=30)
    if not response.ok:
        raise RuntimeError(response.text)
    return response.json()

def with_default_columns(row, columns):
    for column in columns:
        row.setdefault(column, None)
    return row

def build_legacy_sessions_from_items(items):
    groups = {}
    for item in items:
        if item.get("session_id"):
            continue
        campo = item.get("campo") or "Sin campo"
        sector = item.get("sector") or ""
        key = (campo, sector)
        groups.setdefault(key, []).append(item)

    sessions = []
    for (campo, sector), group_items in groups.items():
        dates = sorted([item.get("fecha_hora") for item in group_items if item.get("fecha_hora")])
        started_at = dates[0] if dates else None
        closed_at = dates[-1] if dates else None
        session_id = encode_legacy_session_id(campo, sector)
        sessions.append({
            "id": session_id,
            "nombre": f"Recorrida anterior - {campo}",
            "campo": campo,
            "sector": sector,
            "estado": "legacy",
            "started_at": started_at,
            "closed_at": closed_at,
            "latitud_inicio": group_items[0].get("latitud") if group_items else None,
            "longitud_inicio": group_items[0].get("longitud") if group_items else None,
            "precision_gps_inicio": group_items[0].get("precision_gps") if group_items else None,
            "notas": "Recorrida virtual creada con items viejos sin session_id.",
            "created_at": started_at,
            "updated_at": closed_at,
            "items_count": len(group_items),
            "has_items": bool(group_items),
            "legacy": True,
        })

    sessions.sort(key=lambda item: item.get("started_at") or "", reverse=True)
    return sessions

def get_legacy_session_and_items(session_id, items):
    legacy = decode_legacy_session_id(session_id)
    if not legacy:
        return None, []
    campo = legacy.get("campo") or "Sin campo"
    sector = legacy.get("sector") or ""
    matching_items = [
        item for item in items
        if not item.get("session_id")
        and (item.get("campo") or "Sin campo") == campo
        and (item.get("sector") or "") == sector
    ]
    sessions = build_legacy_sessions_from_items(matching_items)
    session = sessions[0] if sessions else {
        "id": session_id,
        "nombre": f"Recorrida anterior - {campo}",
        "campo": campo,
        "sector": sector,
        "estado": "legacy",
        "items_count": 0,
        "has_items": False,
        "legacy": True,
    }
    session["id"] = session_id
    return session, matching_items

def check_supabase_storage_health():
    result = {
        "configured": bool(SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY and SUPABASE_BUCKET),
        "can_write": False,
        "can_read": False,
        "error": "",
    }
    if not result["configured"]:
        result["error"] = "Faltan variables de Supabase Storage"
        return result

    health_dir = DATA_DIR / "_health"
    health_dir.mkdir(parents=True, exist_ok=True)
    local_path = health_dir / f"campo-health-{uuid.uuid4().hex}.txt"
    download_path = health_dir / f"campo-health-download-{uuid.uuid4().hex}.txt"
    storage_path = f"_health/{local_path.name}"
    content = f"campo health {datetime.utcnow().isoformat()}Z"
    try:
        local_path.write_text(content, encoding="utf-8")
        upload_field_file_to_supabase(local_path, storage_path, content_type="text/plain")
        result["can_write"] = True
        download_supabase_storage_file(storage_path, download_path)
        result["can_read"] = download_path.read_text(encoding="utf-8") == content
    except Exception as e:
        result["error"] = str(e)
    finally:
        try:
            local_path.unlink(missing_ok=True)
            download_path.unlink(missing_ok=True)
        except Exception:
            pass
        try:
            delete_url = f"{SUPABASE_URL}/storage/v1/object/{SUPABASE_BUCKET}"
            requests.delete(
                delete_url,
                headers=supabase_headers(),
                json={"prefixes": [storage_path]},
                timeout=30,
            )
        except Exception:
            pass
    return result

def purge_storage_health_leftovers(max_age_hours=24):
    """Borra restos de chequeos de salud en Supabase Storage."""
    return archive_manager.purge_health_leftovers(max_age_hours=max_age_hours)


def prune_local_data_dir(max_age_days=7):
    """Libera el disco efimero de Render: borra archivos locales viejos ya resguardados.

    Los originales viven en Supabase desde el momento de la subida y luego pasan a la
    PC de Lucas via el archivador; aca solo se limpia la copia temporal local.
    """
    from archive_manager import prune_directory

    return prune_directory(DATA_DIR, max_age_days=max_age_days)


def run_storage_cleanup():
    """Limpieza combinada: candidatos de archivado, restos de salud y disco local."""
    summary = {}
    try:
        summary["archive_candidates"] = (
            archive_manager.sync_candidates() if archive_manager.configured else 0
        )
    except Exception as exc:
        summary["archive_candidates_error"] = str(exc)
    summary["health_purge"] = purge_storage_health_leftovers()
    summary["local_prune"] = prune_local_data_dir(
        max_age_days=int(os.environ.get("LOCAL_PRUNE_MAX_AGE_DAYS", "7") or 7)
    )
    summary["archive_status"] = archive_manager.status()
    return summary


def check_supabase_table_health(table, columns):
    result = {
        "exists": False,
        "missing_columns": [],
        "error": "",
    }
    if not supabase_database_configured():
        result["error"] = "Faltan SUPABASE_URL o SUPABASE_SERVICE_ROLE_KEY"
        result["missing_columns"] = columns
        return result

    try:
        fetch_supabase_rows(table, ["id"], limit=1)
        result["exists"] = True
    except Exception as e:
        result["error"] = str(e)
        result["missing_columns"] = columns
        return result

    for column in columns:
        try:
            fetch_supabase_rows(table, [column], limit=1)
        except Exception:
            result["missing_columns"].append(column)

    return result

def field_metadata_to_item(metadata):
    stored_file = metadata.get("stored_file") or ""
    filename = Path(stored_file).name if stored_file else metadata.get("nombre_archivo", "")
    return {
        "id": metadata.get("id", ""),
        "tipo": metadata.get("item_type") or metadata.get("tipo", ""),
        "campo": metadata.get("campo", ""),
        "sector": metadata.get("sector", ""),
        "fecha_hora": metadata.get("captured_at") or metadata.get("fecha_hora", ""),
        "latitud": metadata.get("latitude") or metadata.get("latitud", ""),
        "longitud": metadata.get("longitude") or metadata.get("longitud", ""),
        "precision_gps": metadata.get("gps_accuracy") or metadata.get("precision_gps", ""),
        "nombre_archivo": filename,
        "estado": metadata.get("estado", "subido"),
        "storage_status": metadata.get("storage_status", ""),
        "storage_provider": metadata.get("storage_provider", ""),
        "storage_path": metadata.get("storage_path", ""),
        "storage_public_url": metadata.get("storage_public_url", ""),
        "storage_error": metadata.get("storage_error", ""),
        "session_id": metadata.get("session_id", ""),
        "session_nombre": metadata.get("session_nombre", ""),
        "photo_label": metadata.get("photo_label", ""),
        "audio_label": metadata.get("audio_label", ""),
        "transcript_status": metadata.get("transcript_status", ""),
        "transcript_text": metadata.get("transcript_text", ""),
        "transcript_error": metadata.get("transcript_error", ""),
        "transcript_model": metadata.get("transcript_model", ""),
        "transcript_at": metadata.get("transcript_at", ""),
        "drive_file_id": metadata.get("drive_file_id", ""),
        "drive_link": metadata.get("drive_web_link", ""),
        "drive_error": metadata.get("drive_error", ""),
    }

def supabase_database_configured():
    return bool(SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY)

def supabase_headers(prefer=None):
    headers = {
        "apikey": SUPABASE_SERVICE_ROLE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
        "Content-Type": "application/json",
    }
    if prefer:
        headers["Prefer"] = prefer
    return headers

def clean_db_value(value):
    return value if value not in ("", None) else None

def validate_record_id(value, label="id"):
    cleaned = str(value or "").strip()
    if label == "session_id" and len(cleaned) <= 512 and is_legacy_session_id(cleaned):
        return cleaned
    if not SAFE_RECORD_ID_RE.fullmatch(cleaned):
        raise HTTPException(
            status_code=400,
            detail=f"{label} invalido",
        )
    return cleaned

def local_session_path(session_id):
    safe_id = validate_record_id(session_id, "session_id")
    return FIELD_SESSIONS_DIR / f"{safe_id}.json"

def save_local_session(session):
    FIELD_SESSIONS_DIR.mkdir(parents=True, exist_ok=True)
    local_session_path(session["id"]).write_text(
        json.dumps(session, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

def normalize_field_session(data):
    now = datetime.utcnow().isoformat() + "Z"
    session_id = (data.get("id") or "").strip()
    if not session_id:
        raise HTTPException(status_code=400, detail="id es obligatorio")
    session_id = validate_record_id(session_id, "id")

    return {
        "id": session_id,
        "nombre": (data.get("nombre") or "").strip(),
        "campo": (data.get("campo") or "").strip(),
        "sector": (data.get("sector") or "").strip(),
        "estado": (data.get("estado") or "abierta").strip(),
        "started_at": clean_db_value(data.get("started_at")),
        "closed_at": clean_db_value(data.get("closed_at")),
        "latitud_inicio": clean_db_value(data.get("latitud_inicio")),
        "longitud_inicio": clean_db_value(data.get("longitud_inicio")),
        "precision_gps_inicio": clean_db_value(data.get("precision_gps_inicio")),
        "notas": (data.get("notas") or "").strip(),
        "created_at": clean_db_value(data.get("created_at")) or now,
        "updated_at": now,
    }

def load_local_sessions():
    sessions = []
    if FIELD_SESSIONS_DIR.exists():
        for session_path in FIELD_SESSIONS_DIR.glob("*.json"):
            try:
                sessions.append(json.loads(session_path.read_text(encoding="utf-8")))
            except Exception as e:
                logger.warning(f"No se pudo leer recorrida {session_path}: {e}")
    sessions.sort(key=lambda item: item.get("started_at") or "", reverse=True)
    return sessions

def load_local_session(session_id):
    path = local_session_path(session_id)
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))

def load_local_field_items():
    items = []
    session_names = {session.get("id"): session.get("nombre", "") for session in load_local_sessions()}
    if FIELD_ITEMS_DIR.exists():
        for metadata_path in FIELD_ITEMS_DIR.rglob("*.json"):
            try:
                metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
            except Exception as e:
                logger.warning(f"No se pudo leer metadata de campo {metadata_path}: {e}")
                continue

            item = field_metadata_to_item(metadata)
            if not item["nombre_archivo"]:
                item["nombre_archivo"] = metadata_path.name.replace(".json", "")
            item["session_nombre"] = item.get("session_nombre") or session_names.get(item.get("session_id"), "")
            items.append(item)

    items.sort(key=lambda item: item.get("fecha_hora") or "", reverse=True)
    return items

def upsert_field_item_metadata(metadata):
    if not supabase_database_configured():
        return False

    item = field_metadata_to_item(metadata)
    payload = {
        "id": item["id"],
        "tipo": item["tipo"],
        "campo": item["campo"],
        "sector": item["sector"],
        "fecha_hora": clean_db_value(item["fecha_hora"]),
        "latitud": clean_db_value(item["latitud"]),
        "longitud": clean_db_value(item["longitud"]),
        "precision_gps": clean_db_value(item["precision_gps"]),
        "nombre_archivo": item["nombre_archivo"],
        "estado": item["estado"],
        "storage_status": item["storage_status"],
        "storage_provider": item["storage_provider"],
        "storage_path": item["storage_path"],
        "storage_public_url": item["storage_public_url"],
        "storage_error": item["storage_error"],
        "session_id": clean_db_value(item.get("session_id")),
        "photo_label": clean_db_value(item.get("photo_label")),
        "audio_label": clean_db_value(item.get("audio_label")),
        "transcript_status": clean_db_value(item.get("transcript_status")),
        "transcript_text": clean_db_value(item.get("transcript_text")),
        "transcript_error": clean_db_value(item.get("transcript_error")),
        "transcript_model": clean_db_value(item.get("transcript_model")),
        "transcript_at": clean_db_value(item.get("transcript_at")),
        "created_at": clean_db_value(metadata.get("received_at")),
    }
    optional_ai_columns = {
        "photo_label", "audio_label", "transcript_status", "transcript_text",
        "transcript_error", "transcript_model", "transcript_at",
    }
    optional_storage_columns = {
        "storage_status", "storage_provider", "storage_path",
        "storage_public_url", "storage_error",
    }
    payload_options = [
        payload,
        {key: value for key, value in payload.items() if key not in optional_ai_columns},
        {
            key: value for key, value in payload.items()
            if key not in optional_ai_columns | optional_storage_columns
        },
    ]
    url = f"{SUPABASE_URL}/rest/v1/field_items?on_conflict=id"
    logger.info("Guardando metadata en Supabase")
    last_error = ""
    for index, candidate in enumerate(payload_options):
        response = requests.post(
            url,
            headers=supabase_headers("resolution=merge-duplicates,return=minimal"),
            json=candidate,
            timeout=30,
        )
        if response.ok:
            if index:
                logger.warning(
                    "Metadata guardada con columnas compatibles; falta aplicar la migracion completa de Supabase"
                )
            logger.info(f"Metadata OK: {item['id']} session_id={item.get('session_id')}")
            return True
        last_error = response.text
    raise RuntimeError(last_error)

def verify_field_item_session_assignment(item_id, session_id):
    if not supabase_database_configured():
        return True
    response = requests.get(
        f"{SUPABASE_URL}/rest/v1/field_items",
        headers=supabase_headers(),
        params={"id": f"eq.{item_id}", "select": "id,session_id", "limit": 1},
        timeout=30,
    )
    if not response.ok:
        raise RuntimeError(f"No se pudo verificar session_id en Supabase: {response.text}")
    rows = response.json()
    if not rows:
        raise RuntimeError("Supabase no devolvio el item recien guardado")
    assigned = str(rows[0].get("session_id") or "")
    if assigned != str(session_id or ""):
        raise RuntimeError(
            f"Supabase guardo session_id={assigned or 'vacio'} y se esperaba {session_id}"
        )
    return True

def list_field_items_from_supabase():
    if not supabase_database_configured():
        return None

    fallback_columns = [
        "id", "tipo", "campo", "sector", "fecha_hora", "latitud", "longitud",
        "precision_gps", "nombre_archivo", "estado", "storage_status",
        "storage_provider", "storage_path", "storage_public_url", "storage_error",
        "session_id", "photo_label", "audio_label", "created_at",
    ]
    legacy_columns = [
        "id", "tipo", "campo", "sector", "fecha_hora", "latitud", "longitud",
        "precision_gps", "nombre_archivo", "estado", "storage_status",
        "storage_provider", "storage_path", "storage_public_url", "storage_error",
        "created_at",
    ]
    minimal_columns = [
        "id", "tipo", "campo", "sector", "fecha_hora", "latitud", "longitud",
        "precision_gps", "nombre_archivo", "estado", "created_at",
    ]
    last_error = None
    for columns in (FIELD_ITEMS_COLUMNS, fallback_columns, legacy_columns, minimal_columns):
        try:
            items = fetch_supabase_rows(
                "field_items",
                columns,
                order="fecha_hora.desc.nullslast",
            )
            break
        except Exception as e:
            last_error = e
    else:
        raise RuntimeError(str(last_error))

    session_names = get_session_name_map_from_supabase()
    for item in items:
        with_default_columns(item, FIELD_ITEMS_COLUMNS)
        item.setdefault("drive_file_id", "")
        item.setdefault("drive_link", "")
        item.setdefault("drive_error", "")
        item.setdefault("session_id", "")
        item["session_nombre"] = session_names.get(item.get("session_id") or "", "")
    return items

def patch_field_item_transcript(item_id, payload):
    if not supabase_database_configured():
        return False

    url = f"{SUPABASE_URL}/rest/v1/field_items?id=eq.{item_id}"
    response = requests.patch(
        url,
        headers=supabase_headers("return=minimal"),
        json=payload,
        timeout=30,
    )
    if not response.ok:
        raise RuntimeError(response.text)
    return True

def upsert_field_report(report):
    if not supabase_database_configured():
        return False

    url = f"{SUPABASE_URL}/rest/v1/field_reports?on_conflict=id"
    response = requests.post(
        url,
        headers=supabase_headers("resolution=merge-duplicates,return=minimal"),
        json=report,
        timeout=30,
    )
    if not response.ok:
        raise RuntimeError(response.text)
    return True

def get_field_report_from_supabase(session_id):
    if not supabase_database_configured():
        return None

    select_options = [
        FIELD_REPORTS_COLUMNS,
        [
            "id", "session_id", "estado", "titulo", "resumen", "informe_markdown",
            "docx_storage_path", "docx_public_url", "error", "created_at", "updated_at",
        ],
    ]
    last_error = None
    for columns in select_options:
        try:
            select = ",".join(columns)
            url = (
                f"{SUPABASE_URL}/rest/v1/field_reports"
                f"?session_id=eq.{session_id}"
                f"&select={select}"
                "&order=created_at.desc.nullslast&limit=1"
            )
            response = requests.get(url, headers=supabase_headers(), timeout=30)
            if not response.ok:
                raise RuntimeError(response.text)
            rows = response.json()
            break
        except Exception as e:
            last_error = e
    else:
        raise RuntimeError(str(last_error))

    for row in rows:
        with_default_columns(row, FIELD_REPORTS_COLUMNS)
    return rows[0] if rows else None


def list_recent_field_reports_from_supabase(limit=10):
    if not supabase_database_configured():
        return []
    select = ",".join(FIELD_REPORTS_COLUMNS)
    response = requests.get(
        f"{SUPABASE_URL}/rest/v1/field_reports",
        headers=supabase_headers(),
        params={
            "select": select,
            "order": "created_at.desc.nullslast",
            "limit": max(1, min(int(limit or 10), 30)),
        },
        timeout=30,
    )
    if not response.ok:
        raise RuntimeError(response.text)
    rows = response.json()
    for row in rows:
        with_default_columns(row, FIELD_REPORTS_COLUMNS)
    return rows

def update_field_report(report_id, payload):
    if not supabase_database_configured():
        return False

    url = f"{SUPABASE_URL}/rest/v1/field_reports?id=eq.{report_id}"
    response = requests.patch(
        url,
        headers=supabase_headers("return=minimal"),
        json=payload,
        timeout=30,
    )
    if not response.ok:
        raise RuntimeError(response.text)
    return True

def get_session_name_map_from_supabase():
    if not supabase_database_configured():
        return {}

    url = f"{SUPABASE_URL}/rest/v1/field_sessions?select=id,nombre"
    response = requests.get(url, headers=supabase_headers(), timeout=30)
    if not response.ok:
        return {}
    return {row.get("id"): row.get("nombre", "") for row in response.json()}

def upsert_field_session_to_supabase(session):
    if not supabase_database_configured():
        return False

    url = f"{SUPABASE_URL}/rest/v1/field_sessions?on_conflict=id"
    logger.info(f"Creando recorrida: {session['id']}")
    response = requests.post(
        url,
        headers=supabase_headers("resolution=merge-duplicates,return=minimal"),
        json=session,
        timeout=30,
    )
    if not response.ok:
        raise RuntimeError(response.text)
    logger.info(f"Recorrida OK: {session['id']}")
    return True

def close_field_session_in_supabase(session_id, closed_at):
    if not supabase_database_configured():
        return False

    url = f"{SUPABASE_URL}/rest/v1/field_sessions?id=eq.{session_id}"
    logger.info(f"Cerrando recorrida: {session_id}")
    response = requests.patch(
        url,
        headers=supabase_headers("return=minimal"),
        json={
            "estado": "cerrada",
            "closed_at": closed_at,
            "updated_at": datetime.utcnow().isoformat() + "Z",
        },
        timeout=30,
    )
    if not response.ok:
        raise RuntimeError(response.text)
    logger.info(f"Recorrida OK: {session_id}")
    return True

def count_items_by_session(items):
    counts = {}
    for item in items:
        session_id = item.get("session_id") or ""
        if session_id:
            counts[session_id] = counts.get(session_id, 0) + 1
    return counts

def list_field_sessions_from_supabase():
    if not supabase_database_configured():
        return None

    sessions = []
    session_error = None
    try:
        sessions = fetch_supabase_rows(
            "field_sessions",
            FIELD_SESSIONS_COLUMNS,
            order="started_at.desc.nullslast",
        )
    except Exception as e:
        session_error = e
        logger.error(f"Recorrida ERROR: {e}")

    item_error = None
    try:
        items = list_field_items_from_supabase() or []
    except Exception as e:
        item_error = e
        logger.error(f"Metadata ERROR: {e}")
        items = []
    counts = count_items_by_session(items)
    for session in sessions:
        with_default_columns(session, FIELD_SESSIONS_COLUMNS)
        session["items_count"] = counts.get(session.get("id"), 0)
        session["has_items"] = session["items_count"] > 0
        session["legacy"] = False
        if item_error:
            session["items_error"] = str(item_error)

    sessions.extend(build_legacy_sessions_from_items(items))
    if session_error and not sessions:
        raise RuntimeError(str(session_error))
    return sessions

def get_items_for_session_from_supabase(session_id):
    items = list_field_items_from_supabase() or []
    if is_legacy_session_id(session_id):
        return get_legacy_session_and_items(session_id, items)

    session = get_field_session_from_supabase(session_id)
    session_items = [item for item in items if item.get("session_id") == session_id]
    if session:
        session["items_count"] = len(session_items)
        session["has_items"] = bool(session_items)
        session["legacy"] = False
    return session, session_items

def get_field_session_from_supabase(session_id):
    if not supabase_database_configured():
        return None

    if is_legacy_session_id(session_id):
        session, _items = get_items_for_session_from_supabase(session_id)
        return session

    select_options = [
        FIELD_SESSIONS_COLUMNS,
        ["id", "nombre", "campo", "sector", "estado", "started_at", "closed_at", "created_at"],
    ]
    last_error = None
    for columns in select_options:
        try:
            select = ",".join(columns)
            url = (
                f"{SUPABASE_URL}/rest/v1/field_sessions"
                f"?id=eq.{session_id}"
                f"&select={select}"
                "&limit=1"
            )
            response = requests.get(url, headers=supabase_headers(), timeout=30)
            if not response.ok:
                raise RuntimeError(response.text)
            rows = response.json()
            break
        except Exception as e:
            last_error = e
    else:
        raise RuntimeError(str(last_error))

    for row in rows:
        with_default_columns(row, FIELD_SESSIONS_COLUMNS)
    return rows[0] if rows else None

def item_file_extension(item, fallback):
    filename = item.get("nombre_archivo") or ""
    suffix = Path(filename).suffix
    return suffix or fallback

def download_field_item_file(item, work_dir, fallback_extension):
    item_id = item.get("id") or uuid.uuid4().hex
    destination = work_dir / f"{item_id}{item_file_extension(item, fallback_extension)}"
    public_url = item.get("storage_public_url") or ""
    if public_url:
        try:
            response = requests.get(public_url, timeout=60, stream=True)
            response.raise_for_status()
            destination.parent.mkdir(parents=True, exist_ok=True)
            with destination.open("wb") as output:
                for chunk in response.iter_content(chunk_size=1024 * 1024):
                    if chunk:
                        output.write(chunk)
            return destination
        except Exception:
            if not item.get("storage_path"):
                raise

    storage_path = item.get("storage_path") or ""
    if storage_path:
        return download_supabase_storage_file(storage_path, destination)

    raise RuntimeError(f"Item {item_id} no tiene storage_public_url ni storage_path")

def transcribe_field_audio(item, work_dir):
    if item.get("transcript_status") == "done" and item.get("transcript_text"):
        return item["transcript_text"]

    item_id = item.get("id", "")
    filename = item.get("nombre_archivo") or item_id
    logger.info(f"Transcribiendo audio: {item_id}")
    try:
        audio_path = download_field_item_file(item, work_dir, ".webm")
    except Exception as e:
        error_payload = {
            "transcript_status": "error",
            "transcript_error": f"No se pudo descargar audio: {e}",
            "transcript_at": datetime.utcnow().isoformat() + "Z",
        }
        try:
            patch_field_item_transcript(item_id, error_payload)
        except Exception:
            pass
        item.update(error_payload)
        logger.error(f"Audio ERROR: {item_id}: {error_payload['transcript_error']}")
        return None

    size_bytes = audio_path.stat().st_size if audio_path.exists() else 0
    logger.info(f"Audio descargado: nombre_archivo={filename} path={audio_path} bytes={size_bytes}")
    if size_bytes < 5 * 1024:
        error_payload = {
            "transcript_status": "error",
            "transcript_error": "Audio demasiado corto o vacío",
            "transcript_at": datetime.utcnow().isoformat() + "Z",
        }
        try:
            patch_field_item_transcript(item_id, error_payload)
        except Exception:
            pass
        item.update(error_payload)
        logger.error(f"Audio ERROR: {item_id}: {error_payload['transcript_error']}")
        return None

    try:
        transcript_text = transcribe_audio(audio_path)
        transcript_payload = {
            "transcript_status": "done",
            "transcript_text": transcript_text,
            "transcript_model": "whisper-1",
            "transcript_at": datetime.utcnow().isoformat() + "Z",
            "transcript_error": None,
        }
        patch_field_item_transcript(item_id, transcript_payload)
        item.update(transcript_payload)
        logger.info(f"Audio transcripto OK: {item_id}")
        return transcript_text
    except Exception as e:
        message = str(e)
        error_payload = {
            "transcript_status": "error",
            "transcript_error": message,
            "transcript_at": datetime.utcnow().isoformat() + "Z",
        }
        try:
            patch_field_item_transcript(item_id, error_payload)
        except Exception:
            pass
        item.update(error_payload)
        logger.error(f"Audio ERROR: {item_id}: {message}")
        return None

def format_item_line(item):
    return (
        f"- {item.get('fecha_hora') or 'sin fecha'} | {item.get('tipo')} | "
        f"{item.get('campo') or 'sin campo'} | {item.get('sector') or 'sin sector'} | "
        f"GPS {item.get('latitud') or '-'}, {item.get('longitud') or '-'} "
        f"+/- {item.get('precision_gps') or '-'} m | archivo: {item.get('nombre_archivo') or '-'}"
    )

def clean_inline_markdown(text):
    cleaned = str(text or "")
    cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"__(.*?)__", r"\1", cleaned)
    cleaned = re.sub(r"`([^`]*)`", r"\1", cleaned)
    cleaned = re.sub(r"https?://\S+", "Ver archivo original", cleaned)
    return cleaned.strip()

def normalize_heading_key(text):
    cleaned = clean_inline_markdown(text).lower().strip()
    cleaned = "".join(
        char for char in unicodedata.normalize("NFKD", cleaned)
        if not unicodedata.combining(char)
    )
    return cleaned

def normalize_transcript_text(text):
    cleaned = clean_inline_markdown(text)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned

def is_noise_transcript(text):
    cleaned = normalize_transcript_text(text)
    normalized = cleaned.lower().strip(" .,!¡¿?;:")
    if not normalized:
        return True
    return normalized in NOISE_TRANSCRIPTS

def valid_transcript_blocks(audios):
    blocks = []
    for audio in audios:
        transcript = normalize_transcript_text(audio.get("transcript_text"))
        if is_noise_transcript(transcript):
            continue
        blocks.append(
            "Fecha {fecha}:\n{texto}".format(
                fecha=audio.get("fecha_hora") or "sin fecha",
                texto=transcript,
            )
        )
    return blocks

def transcript_report_markdown(audios):
    transcript_blocks = valid_transcript_blocks(audios)
    if not transcript_blocks:
        if audios:
            return "No se pudieron incorporar transcripciones válidas de esta recorrida."
        return "No se registraron notas de voz en esta recorrida."
    return "\n\n".join(
        f"### Nota de voz {index}\n{block}"
        for index, block in enumerate(transcript_blocks, start=1)
    )

def audio_error_lines(audios_con_error):
    return [
        "- {archivo} | {fecha} | {error}".format(
            archivo=audio.get("nombre_archivo") or audio.get("id") or "audio",
            fecha=audio.get("fecha_hora") or "sin fecha",
            error=audio.get("transcript_error") or "Error no registrado",
        )
        for audio in audios_con_error
    ]

def summarized_photo_lines(photos, limit=PHOTO_PROMPT_LIMIT):
    lines = [
        format_item_line(photo) + f" | comentario: {photo.get('photo_label') or 'sin comentario'}"
        for photo in photos[:limit]
    ]
    remaining = len(photos) - len(lines)
    if remaining > 0:
        lines.append(f"- {remaining} fotos adicionales listadas solo como evidencia en el DOCX.")
    return lines

def summarized_item_lines(items, limit=LIGHT_REPORT_ITEM_LIMIT):
    lines = [format_item_line(item) for item in items[:limit]]
    remaining = len(items) - len(lines)
    if remaining > 0:
        lines.append(f"- {remaining} items adicionales omitidos del prompt para modo liviano.")
    return lines

def build_basic_report_markdown(session, audios, photos, items, audios_con_error=None):
    transcript_blocks = valid_transcript_blocks(audios)
    transcript_section = transcript_report_markdown(audios)
    audio_error_count = len(audios_con_error or [])
    summary = (
        "Informe generado con las notas de voz y la metadata registradas durante la recorrida."
        if transcript_blocks
        else "Informe elaborado con la metadata disponible. No se agregan conclusiones fuera de los datos registrados."
    )
    return "\n\n".join([
        f"# Informe de recorrida - {session.get('campo') or 'Campo'} - {(session.get('started_at') or '')[:10]}",
        "## Datos generales\n"
        f"- Campo: {session.get('campo') or 'No registrado en la recorrida'}\n"
        f"- Sector: {session.get('sector') or 'No registrado en la recorrida'}\n"
        f"- Fecha de inicio: {session.get('started_at') or 'No registrado en la recorrida'}\n"
        f"- Fecha de cierre: {session.get('closed_at') or 'No registrado en la recorrida'}\n"
        f"- Cantidad de notas de voz: {len(audios)}\n"
        f"- Notas con error de transcripción: {audio_error_count}\n"
        f"- Cantidad de fotos: {len(photos)}",
        f"## Resumen ejecutivo\n{summary}",
        "## Diagnostico de situacion\nNo registrado en la recorrida.",
        "## Observaciones principales\n" + (
            "Las observaciones de la recorrida están preservadas en la sección de notas de voz."
            if transcript_blocks
            else "No registrado en la recorrida."
        ),
        "## Analisis economico para la decision\nNo se registraron costos ni beneficios cuantificables. Completar precios, cantidades y horizonte antes de decidir una inversion.",
        "## Recomendaciones\nRevisar las evidencias disponibles y completar observaciones manuales si corresponde.",
        f"## Notas de voz registradas\n{transcript_section}",
    ])

def build_report_markdown(session, audios, photos, items, audios_con_error=None):
    logger.info("Generando texto del informe")
    transcript_blocks = valid_transcript_blocks(audios)
    light_mode = len(items) > LIGHT_REPORT_ITEM_LIMIT
    photo_lines = summarized_photo_lines(photos)
    item_lines = summarized_item_lines(items) if light_mode else [format_item_line(item) for item in items]
    prompt = f"""
Redacta un informe profesional, limpio y corto para cliente.
No inventes datos. Si falta informacion, escribir "No registrado en la recorrida".
No interpretes fotos ni describas su contenido visual; las fotos son solo evidencia documental.
Las recomendaciones deben basarse solo en las notas de voz transcriptas y la metadata de campo.
Integra la dimension economica a la decision: alternativas, costos y beneficios a cuantificar, horizonte y riesgos. No inventes precios ni montos.
Integra todas las notas de voz válidas en el resumen, el diagnóstico, las observaciones o las recomendaciones; no omitas notas por ser breves.
No muestres IDs de audios, nombres de archivos ni errores internos.
No escribas secciones llamadas Audios transcriptos, Audios no transcriptos, Anexo tecnico o Informe Tecnico de Consultoria Agronomica.
No uses markdown de negritas con **.

Estructura requerida:
1. Resumen ejecutivo
2. Diagnostico de situacion
3. Observaciones principales
4. Analisis economico para la decision
5. Recomendaciones

Recorrida:
{json.dumps(session, ensure_ascii=False, indent=2)}

Notas de voz transcriptas que deben incorporarse al análisis:
{chr(10).join(transcript_blocks) or "No registrado en la recorrida"}

Fotos como evidencia, solo metadata:
{chr(10).join(photo_lines) or "No registrado en la recorrida"}

Items relevados, solo metadata resumida:
{chr(10).join(item_lines) or "No registrado en la recorrida"}
"""
    response = get_openai_client().chat.completions.create(
        model=os.environ.get("FIELD_REPORT_MODEL", "gpt-4o-mini"),
        messages=[
            {"role": "system", "content": "Sos un asesor ganadero profesional. Escribis claro, breve, util y sin inventar datos."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )
    report_markdown = response.choices[0].message.content.strip()
    return f"{report_markdown}\n\n## Notas de voz registradas\n{transcript_report_markdown(audios)}"

def markdown_summary(markdown_text):
    for line in markdown_text.splitlines():
        cleaned = clean_inline_markdown(line).strip().lstrip("#").strip()
        if cleaned and not cleaned.lower().startswith("informe de recorrida"):
            return cleaned[:1000]
    return ""

BRAND_DARK = "1F4A36"
BRAND_MID = "4B8058"
BRAND_BEIGE = "F1EBDD"
BRAND_BROWN = "886844"
BRAND_TEXT = "24342A"

def logo_path():
    return STATIC_DIR / "logo.png"

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)

def set_cell_border(cell, color=BRAND_BROWN, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)

def set_cell_text(cell, text, bold=False, color=BRAND_TEXT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def style_report_document(document):
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.different_first_page_header_footer = True

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = RGBColor.from_string(BRAND_TEXT)
    for style_name, color, size in [
        ("Heading 1", BRAND_DARK, 18),
        ("Heading 2", BRAND_MID, 14),
        ("Heading 3", BRAND_BROWN, 11),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

def add_logo_to_paragraph(paragraph, width):
    path = logo_path()
    if not path.exists():
        logger.warning("logo no encontrado")
        return False
    paragraph.add_run().add_picture(str(path), width=width)
    return True

def add_report_header(document):
    header = document.sections[0].header
    table = header.add_table(rows=1, cols=2, width=Inches(6.7))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]
    paragraph = left.paragraphs[0]
    add_logo_to_paragraph(paragraph, Inches(0.65))
    right_paragraph = right.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right_paragraph.add_run("Lucas Estecho – Asesor Ganadero")
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(BRAND_DARK)

def add_cover_page(document, session, title):
    logo_paragraph = document.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_logo_to_paragraph(logo_paragraph, Inches(2.4))

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run("Informe de recorrida")
    title_run.font.name = "Arial"
    title_run.font.bold = True
    title_run.font.size = Pt(30)
    title_run.font.color.rgb = RGBColor.from_string(BRAND_DARK)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(session.get("nombre") or title)
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(15)
    subtitle_run.font.color.rgb = RGBColor.from_string(BRAND_MID)

    info = document.add_table(rows=0, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in [
        ("Campo", session.get("campo") or "No registrado en la recorrida"),
        ("Sector", session.get("sector") or "No registrado en la recorrida"),
        ("Fecha", (session.get("started_at") or "No registrado en la recorrida")[:10]),
    ]:
        row = info.add_row().cells
        set_cell_shading(row[0], BRAND_BEIGE)
        set_cell_text(row[0], label, bold=True, color=BRAND_DARK)
        set_cell_text(row[1], value)
        set_cell_border(row[0], color="D6C7A8", size="6")
        set_cell_border(row[1], color="D6C7A8", size="6")

    document.add_paragraph()
    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("Ing. Agr. Lucas Estecho\nAsesor Ganadero")
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(BRAND_BROWN)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def add_divider(document):
    table = document.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, BRAND_BROWN)
    cell.text = ""

def extract_markdown_section(markdown_text, section_title):
    lines = markdown_text.splitlines()
    captured = []
    capture = False
    wanted = normalize_heading_key(section_title)
    for line in lines:
        stripped = line.strip().lstrip("#").strip()
        if normalize_heading_key(stripped).startswith(wanted):
            capture = True
            continue
        if capture and line.strip().startswith("#"):
            break
        if capture:
            captured.append(line.strip())
    text = "\n".join(clean_inline_markdown(line) for line in captured if line).strip()
    return text or "No registrado en la recorrida"

def add_highlight_box(document, title, text):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, BRAND_BEIGE)
    set_cell_border(cell, color="D6C7A8", size="8")
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(title + "\n")
    title_run.font.name = "Arial"
    title_run.font.bold = True
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = RGBColor.from_string(BRAND_DARK)
    body_run = paragraph.add_run(clean_inline_markdown(text))
    body_run.font.name = "Arial"
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = RGBColor.from_string(BRAND_TEXT)

def add_markdown_to_doc(document, markdown_text):
    banned_sections = (
        "audios transcriptos",
        "audios transcritos",
        "audios no transcriptos",
        "audios no transcritos",
        "anexo tecnico",
        "anexo técnico",
        "informe tecnico",
        "informe técnico",
        "informe tecnico de consultoria",
        "informe técnico de consultoría",
    )
    for line in markdown_text.splitlines():
        text = clean_inline_markdown(line).strip()
        if not text:
            continue
        normalized = normalize_heading_key(text.lstrip("#").strip())
        if normalized.startswith(banned_sections):
            continue
        if "sin transcripcion disponible" in normalized or "sin transcripción disponible" in normalized:
            continue
        if text.startswith("# "):
            document.add_heading(text[2:].strip(), level=1)
        elif text.startswith("## "):
            document.add_heading(text[3:].strip(), level=2)
        elif text.startswith("### "):
            document.add_heading(text[4:].strip(), level=3)
        elif text.startswith(("- ", "* ")):
            document.add_paragraph(text[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(text)

def prepare_photo_for_docx(photo_path, work_dir, index):
    size_bytes = photo_path.stat().st_size if photo_path.exists() else 0
    if size_bytes > MAX_SOURCE_PHOTO_BYTES:
        raise RuntimeError(f"foto demasiado pesada para insertar ({size_bytes} bytes)")

    if not Image:
        logger.warning("Pillow no disponible; se inserta la foto sin comprimir")
        return photo_path

    output_path = work_dir / f"docx_photo_{index}.jpg"
    with Image.open(photo_path) as image:
        image.verify()
    with Image.open(photo_path) as image:
        image.thumbnail((1280, 1280))
        if image.mode not in ("RGB", "L"):
            image = image.convert("RGB")
        elif image.mode == "L":
            image = image.convert("RGB")
        image.save(output_path, format="JPEG", quality=68, optimize=True)
    return output_path

def add_photo_metadata_table(document, photos, title):
    if not photos:
        return

    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["fecha", "sector", "comentario", "GPS", "archivo"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BRAND_BEIGE)
        set_cell_text(cell, header, bold=True, color=BRAND_DARK)
        set_cell_border(cell, color="D6C7A8", size="4")

    for photo in photos:
        row = table.add_row().cells
        values = [
            photo.get("fecha_hora") or "",
            photo.get("sector") or "",
            photo.get("photo_label") or "",
            f"{photo.get('latitud') or '-'}, {photo.get('longitud') or '-'}",
            photo.get("nombre_archivo") or "",
        ]
        for idx, value in enumerate(values):
            set_cell_text(row[idx], value)
            set_cell_border(row[idx], color="D6C7A8", size="4")

def add_photo_evidence_to_doc(document, photos, work_dir):
    if not photos:
        document.add_paragraph("No registrado en la recorrida.")
        return

    inserted_photos = photos[:MAX_REPORT_PHOTOS]
    not_inserted_photos = photos[MAX_REPORT_PHOTOS:]
    omitted_count = len(not_inserted_photos)
    if omitted_count:
        document.add_paragraph(f"Se muestran {MAX_REPORT_PHOTOS} fotos de {len(photos)} para mantener el informe liviano.")
        logger.info(f"Fotos limitadas en DOCX: {len(inserted_photos)} de {len(photos)}")

    for index, photo in enumerate(inserted_photos, start=1):
        photo_path = None
        prepared_photo_path = None
        try:
            photo_path = download_field_item_file(photo, work_dir, ".jpg")
            prepared_photo_path = prepare_photo_for_docx(photo_path, work_dir, index)
        except Exception as e:
            logger.error(f"Foto omitida en DOCX: {photo.get('nombre_archivo') or photo.get('id')}: {e}")
            not_inserted_photos.append(photo)
            continue

        label = clean_inline_markdown(photo.get("photo_label") or "")
        heading = f"Foto {index} — {label}" if label else f"Foto {index}"
        document.add_heading(heading, level=3)
        frame = document.add_table(rows=1, cols=1)
        frame.alignment = WD_TABLE_ALIGNMENT.CENTER
        frame_cell = frame.rows[0].cells[0]
        set_cell_border(frame_cell, color="D6C7A8", size="10")
        paragraph = frame_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(prepared_photo_path), width=Inches(5.0))
        try:
            if prepared_photo_path and prepared_photo_path != photo_path:
                prepared_photo_path.unlink(missing_ok=True)
            if photo_path:
                photo_path.unlink(missing_ok=True)
        except Exception as e:
            logger.warning(f"No se pudo borrar temporal de foto: {e}")

        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_text = (
            f"{photo.get('fecha_hora') or 'No registrado en la recorrida'} | "
            f"Campo: {photo.get('campo') or 'No registrado en la recorrida'} | "
            f"Sector: {photo.get('sector') or 'No registrado en la recorrida'} | "
            f"Lat: {photo.get('latitud') or '-'} | Long: {photo.get('longitud') or '-'} | "
            f"Precision GPS: {photo.get('precision_gps') or '-'}"
        )
        run = caption.add_run(caption_text)
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string(BRAND_BROWN)

        evidence = document.add_table(rows=0, cols=2)
        for label, value in [
            ("Archivo", photo.get("nombre_archivo") or "No registrado en la recorrida"),
            ("Comentario", photo.get("photo_label") or "Sin comentario"),
            ("Original", "Ver archivo original" if photo.get("storage_public_url") else "No registrado en la recorrida"),
        ]:
            row = evidence.add_row().cells
            set_cell_shading(row[0], BRAND_BEIGE)
            set_cell_text(row[0], label, bold=True, color=BRAND_DARK)
            set_cell_text(row[1], value)
            set_cell_border(row[0], color="D6C7A8", size="4")
            set_cell_border(row[1], color="D6C7A8", size="4")

    add_photo_metadata_table(document, not_inserted_photos, "Evidencias fotograficas no insertadas")

def create_report_docx(session, items, audios, photos, markdown_text, output_path, work_dir):
    logger.info("Creando DOCX")
    document = DocxDocument()
    title = f"Informe de recorrida - {session.get('campo') or 'Campo'} - {(session.get('started_at') or '')[:10]}"
    style_report_document(document)
    add_report_header(document)
    add_cover_page(document, session, title)

    document.add_heading("Datos generales", level=1)
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in [
        ("Campo", session.get("campo") or "No registrado en la recorrida"),
        ("Sector", session.get("sector") or "No registrado en la recorrida"),
        ("Nombre de recorrida", session.get("nombre") or "No registrado en la recorrida"),
        ("Fecha de inicio", session.get("started_at") or "No registrado en la recorrida"),
        ("Fecha de cierre", session.get("closed_at") or "No registrado en la recorrida"),
        ("Cantidad de audios", len(audios)),
        ("Cantidad de fotos", len(photos)),
    ]:
        row = table.add_row().cells
        set_cell_shading(row[0], BRAND_BEIGE)
        set_cell_text(row[0], label, bold=True, color=BRAND_DARK)
        set_cell_text(row[1], value)
        set_cell_border(row[0], color="D6C7A8", size="6")
        set_cell_border(row[1], color="D6C7A8", size="6")

    add_divider(document)
    document.add_heading("Resumen ejecutivo", level=1)
    add_highlight_box(document, "Resumen ejecutivo", extract_markdown_section(markdown_text, "Resumen ejecutivo"))

    for section_title in REPORT_SECTION_TITLES:
        document.add_paragraph()
        document.add_heading(section_title, level=1)
        add_markdown_to_doc(document, extract_markdown_section(markdown_text, section_title))

    document.add_heading("Evidencias fotograficas", level=1)
    document.add_paragraph("Las fotos se incluyen solo como evidencia. No fueron interpretadas por IA.")
    add_photo_evidence_to_doc(document, photos, work_dir)

    document.add_paragraph()
    method_note = document.add_paragraph()
    method_note.add_run(
        "Nota metodologica: El informe fue elaborado a partir de audios, fotos y coordenadas relevadas durante la recorrida. "
        "Las fotografias se incorporan como evidencia documental. No fueron interpretadas automaticamente por IA."
    )

    document.add_paragraph()
    closing = document.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = closing.add_run("Elaborado por Ing. Agr. Lucas Estecho\nAsesor Ganadero")
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(BRAND_DARK)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path

def create_report_pdf(session, items, audios, photos, markdown_text, output_path, work_dir):
    """Crea el PDF entregable sin depender de LibreOffice en Render."""
    logger.info("Creando PDF")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    dark = colors.HexColor(f"#{BRAND_DARK}")
    beige = colors.HexColor(f"#{BRAND_BEIGE}")
    brown = colors.HexColor(f"#{BRAND_BROWN}")
    styles.add(ParagraphStyle(
        name="CapatazTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=dark,
        alignment=TA_CENTER,
        spaceAfter=18,
    ))
    styles.add(ParagraphStyle(
        name="CapatazHeading",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=17,
        textColor=dark,
        spaceBefore=12,
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="CapatazBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        textColor=colors.HexColor(f"#{BRAND_TEXT}"),
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="CapatazCaption",
        parent=styles["BodyText"],
        fontName="Helvetica-Oblique",
        fontSize=8,
        leading=10,
        textColor=brown,
        alignment=TA_CENTER,
        spaceAfter=10,
    ))
    styles.add(ParagraphStyle(
        name="CapatazSignature",
        parent=styles["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=13,
        textColor=dark,
        alignment=TA_RIGHT,
        spaceBefore=18,
    ))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=1.7 * cm,
        leftMargin=1.7 * cm,
        topMargin=1.7 * cm,
        bottomMargin=1.7 * cm,
        title=f"Informe de recorrida - {session.get('campo') or 'Campo'}",
        author="Ing. Agr. Lucas Estecho",
    )
    story = []
    title = f"Informe de recorrida<br/>{html.escape(str(session.get('campo') or 'Campo'))}"
    story.append(Paragraph(title, styles["CapatazTitle"]))
    details = [
        ["Campo", session.get("campo") or "No registrado"],
        ["Sector", session.get("sector") or "No registrado"],
        ["Recorrida", session.get("nombre") or "No registrado"],
        ["Inicio", session.get("started_at") or "No registrado"],
        ["Cierre", session.get("closed_at") or "No registrado"],
        ["Evidencias", f"{len(audios)} audio(s) y {len(photos)} foto(s)"],
    ]
    details_table = Table(
        [[Paragraph(f"<b>{html.escape(str(label))}</b>", styles["CapatazBody"]),
          Paragraph(html.escape(str(value)), styles["CapatazBody"])] for label, value in details],
        colWidths=[4 * cm, 12 * cm],
    )
    details_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), beige),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#D6C7A8")),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D6C7A8")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.extend([details_table, Spacer(1, 12)])

    for section_title in ["Resumen ejecutivo", *REPORT_SECTION_TITLES]:
        story.append(Paragraph(html.escape(section_title), styles["CapatazHeading"]))
        section_text = extract_markdown_section(markdown_text, section_title)
        paragraphs = [line.strip() for line in section_text.splitlines() if line.strip()]
        if not paragraphs:
            paragraphs = ["No registrado en la recorrida"]
        for paragraph in paragraphs:
            story.append(Paragraph(html.escape(clean_inline_markdown(paragraph)), styles["CapatazBody"]))

    story.append(Paragraph("Evidencias fotograficas", styles["CapatazHeading"]))
    story.append(Paragraph(
        "Las fotografias se incorporan como evidencia documental y no fueron interpretadas automaticamente por IA.",
        styles["CapatazBody"],
    ))
    temporary_photos = []
    for index, photo in enumerate(photos[:MAX_REPORT_PHOTOS], start=1):
        try:
            downloaded = download_field_item_file(photo, work_dir, ".jpg")
            prepared = prepare_photo_for_docx(downloaded, work_dir, index)
            temporary_photos.extend([downloaded, prepared])
            image = PdfImage(str(prepared))
            max_width = 15.5 * cm
            max_height = 16 * cm
            scale = min(max_width / image.drawWidth, max_height / image.drawHeight, 1)
            image.drawWidth *= scale
            image.drawHeight *= scale
            label = photo.get("photo_label") or f"Foto {index}"
            caption = (
                f"{label} | {photo.get('fecha_hora') or 'sin fecha'} | "
                f"Sector: {photo.get('sector') or 'sin sector'} | "
                f"GPS: {photo.get('latitud') or '-'}, {photo.get('longitud') or '-'}"
            )
            story.append(KeepTogether([
                Paragraph(html.escape(str(label)), styles["CapatazHeading"]),
                image,
                Paragraph(html.escape(caption), styles["CapatazCaption"]),
            ]))
        except Exception as exc:
            logger.error(f"Foto omitida en PDF: {photo.get('id')}: {exc}")

    story.append(Paragraph(
        "Nota metodologica: informe elaborado a partir de audios, fotos, coordenadas y datos registrados durante la recorrida.",
        styles["CapatazBody"],
    ))
    story.append(Paragraph(
        "Elaborado por Ing. Agr. Lucas Estecho<br/>Asesor Ganadero",
        styles["CapatazSignature"],
    ))
    doc.build(story)
    for path in temporary_photos:
        try:
            if path:
                Path(path).unlink(missing_ok=True)
        except Exception as exc:
            logger.warning(f"No se pudo borrar temporal de PDF: {exc}")
    return output_path

def parse_iso_datetime(value):
    if not value:
        return None
    try:
        return datetime.fromisoformat(str(value).replace("Z", "+00:00")).replace(tzinfo=None)
    except Exception:
        return None

def recent_generating_report(report):
    if not report or report.get("estado") != "generando":
        return False
    updated = parse_iso_datetime(report.get("updated_at") or report.get("started_at"))
    if not updated:
        return False
    return (datetime.utcnow() - updated).total_seconds() < 600

def update_report_progress(report_id, message, estado="generando"):
    payload = {
        "estado": estado,
        "progress_message": message,
        "updated_at": datetime.utcnow().isoformat() + "Z",
    }
    update_field_report(report_id, payload)

def generate_field_report(session_id, force=False):
    if not supabase_database_configured():
        raise RuntimeError("Supabase Database no configurado")

    logger.info(f"Iniciando generacion de informe: {session_id}")
    existing_report = get_field_report_from_supabase(session_id)
    if existing_report and existing_report.get("estado") == "done" and not force:
        logger.info(f"Informe ya existente para {session_id}: {existing_report.get('id')}")
        return existing_report
    if existing_report and recent_generating_report(existing_report) and not force:
        logger.info(f"Informe en generacion reciente para {session_id}: {existing_report.get('id')}")
        return existing_report

    session, items = get_items_for_session_from_supabase(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="recorrida no encontrada")

    items.sort(key=lambda item: item.get("fecha_hora") or "")
    audios = [item for item in items if item.get("tipo") == "audio"]
    photos = [item for item in items if item.get("tipo") == "foto"]
    light_mode = len(items) > LIGHT_REPORT_ITEM_LIMIT
    logger.info(f"Audios encontrados: {len(audios)}")
    logger.info(f"Fotos encontradas: {len(photos)}")
    if light_mode:
        logger.info(f"Informe en modo liviano por cantidad de items: {len(items)}")

    report_id = uuid.uuid4().hex
    now = datetime.utcnow().isoformat() + "Z"
    title = f"Informe de recorrida - {session.get('campo') or 'Campo'} - {(session.get('started_at') or now)[:10]}"
    work_dir = DATA_DIR / "field_reports" / safe_storage_segment(session_id, "session") / report_id
    work_dir.mkdir(parents=True, exist_ok=True)
    docx_path = None
    pdf_path = None
    upload_started = False

    try:
        upsert_field_report({
            "id": report_id,
            "session_id": session_id,
            "estado": "generando",
            "titulo": title,
            "progress_message": "Transcribiendo audios",
            "started_at": now,
            "finished_at": None,
            "error": None,
            "created_at": now,
            "updated_at": now,
        })

        for audio in audios:
            transcribe_field_audio(audio, work_dir)

        audios_con_error = [audio for audio in audios if audio.get("transcript_status") == "error"]
        valid_audio_count = len(valid_transcript_blocks(audios))
        logger.info(f"Transcripciones incorporadas al informe: {valid_audio_count}")
        if not valid_audio_count:
            logger.info("No se pudieron transcribir audios validos de esta recorrida.")

        update_report_progress(report_id, "Generando texto")
        try:
            markdown_text = build_report_markdown(session, audios, photos, items, audios_con_error)
        except Exception as e:
            logger.error(f"Informe texto IA ERROR: {e}")
            markdown_text = build_basic_report_markdown(session, audios, photos, items, audios_con_error)
        summary = markdown_summary(markdown_text)
        report_date = (session.get("started_at") or now)[:10]
        campo_filename = clean_report_filename_segment(session.get("campo"), "Sin_Campo")
        docx_name = f"Informe_{campo_filename}_{report_date}.docx"
        pdf_name = f"Informe_{campo_filename}_{report_date}.pdf"
        docx_path = work_dir / docx_name
        pdf_path = work_dir / pdf_name
        update_report_progress(report_id, "Creando DOCX")
        try:
            create_report_docx(session, items, audios, photos, markdown_text, docx_path, work_dir)
        except MemoryError:
            raise RuntimeError("Memoria insuficiente generando DOCX; informe demasiado pesado")
        update_report_progress(report_id, "Creando PDF")
        try:
            create_report_pdf(session, items, audios, photos, markdown_text, pdf_path, work_dir)
        except MemoryError:
            raise RuntimeError("Memoria insuficiente generando PDF; informe demasiado pesado")

        campo_segment = safe_storage_segment(session.get("campo"), "sin-campo")
        session_segment = safe_storage_segment(session_id, "session")
        docx_storage_path = f"reports/{campo_segment}/{session_segment}/{docx_name}"
        pdf_storage_path = f"reports/{campo_segment}/{session_segment}/{pdf_name}"
        logger.info(f"Subiendo informe a Supabase: {docx_storage_path} y {pdf_storage_path}")
        update_report_progress(report_id, "Subiendo DOCX y PDF")
        upload_started = True
        try:
            docx_file = upload_field_file_to_supabase(
                docx_path,
                docx_storage_path,
                content_type=REPORT_DOCX_CONTENT_TYPE,
                upsert=True,
            )
            pdf_file = upload_field_file_to_supabase(
                pdf_path,
                pdf_storage_path,
                content_type=REPORT_PDF_CONTENT_TYPE,
                upsert=True,
            )
        except Exception as e:
            logger.error(f"Informe creado, falló subida a Supabase. docx={docx_path} pdf={pdf_path}")
            update_report_progress(report_id, "Informe creado, falló subida a Supabase", estado="error")
            raise RuntimeError(f"{e} | docx_path={docx_path} | pdf_path={pdf_path}")
        report = {
            "id": report_id,
            "session_id": session_id,
            "estado": "done",
            "titulo": title,
            "resumen": summary,
            "informe_markdown": markdown_text,
            "docx_storage_path": docx_file.get("path", ""),
            "docx_public_url": docx_file.get("public_url", ""),
            "pdf_storage_path": pdf_file.get("path", ""),
            "pdf_public_url": pdf_file.get("public_url", ""),
            "error": None,
            "created_at": now,
            "progress_message": "Informe listo",
            "finished_at": datetime.utcnow().isoformat() + "Z",
            "updated_at": datetime.utcnow().isoformat() + "Z",
        }
        upsert_field_report(report)
        logger.info(f"Informe OK: {report_id}")
        return report
    except Exception as e:
        progress_message = "Informe creado, falló subida a Supabase" if upload_started and docx_path else "Error al generar informe"
        error_message = str(e) or "Error desconocido generando informe"
        if isinstance(e, MemoryError):
            error_message = "Memoria insuficiente generando DOCX; informe demasiado pesado"
        if docx_path and docx_path.exists():
            logger.error(f"Copia local temporal del DOCX: {docx_path}")
        if pdf_path and pdf_path.exists():
            logger.error(f"Copia local temporal del PDF: {pdf_path}")
        error_report = {
            "id": report_id,
            "session_id": session_id,
            "estado": "error",
            "titulo": title,
            "error": error_message,
            "progress_message": progress_message,
            "started_at": now,
            "finished_at": datetime.utcnow().isoformat() + "Z",
            "created_at": now,
            "updated_at": datetime.utcnow().isoformat() + "Z",
        }
        try:
            upsert_field_report(error_report)
        except Exception:
            pass
        logger.error(f"Informe ERROR: {error_message}")
        raise RuntimeError(error_message)

def get_superficie_from_hoja2(lote):
    try:
        hoja2 = get_google_sheet().worksheet("Hoja 2")
        lotes = hoja2.col_values(1)
        superficies = hoja2.col_values(2)
        lote_normalizado = lote.strip().lower()
        for i, nombre in enumerate(lotes):
            if nombre.strip().lower() == lote_normalizado:
                if i < len(superficies):
                    return superficies[i]
        return None
    except Exception as e:
        logger.warning(f"No se pudo obtener superficie de Hoja 2: {e}")
        return None

def get_next_receta_number(worksheet):
    values = worksheet.col_values(11)
    nums = []
    for v in values[1:]:
        try:
            n = int(v)
            if n > 0:
                nums.append(n)
        except:
            pass
    return max(nums) + 1 if nums else 1

def transcribe_audio(file_path):
    with open(file_path, "rb") as audio_file:
        transcript = get_openai_client().audio.transcriptions.create(
            model="whisper-1",
            file=audio_file
        )
    return transcript.text

def image_to_base64(image_path):
    with open(image_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")

def transcribe_image_base64(image_base64):
    response = get_openai_client().chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/jpeg;base64,{image_base64}"
                        }
                    },
                    {
                        "type": "text",
                        "text": (
                            "Transcribi todo el texto que ves en esta imagen de forma completa y ordenada. "
                            "Si es un presupuesto, factura, nota o documento, extrae todos los datos visibles: "
                            "cliente, montos, conceptos, fechas, condiciones, totales, observaciones. "
                            "Devuelve el texto plano sin formato especial."
                        )
                    }
                ]
            }
        ],
        max_tokens=1500
    )
    return response.choices[0].message.content.strip()

def transcribe_image(image_path):
    image_base64 = image_to_base64(image_path)
    return transcribe_image_base64(image_base64)

def describir_imagen_recorrida(image_path):
    """Describe una foto de campo en contexto de recorrida tecnica."""
    image_base64 = image_to_base64(image_path)
    response = get_openai_client().chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/jpeg;base64,{image_base64}"
                        }
                    },
                    {
                        "type": "text",
                        "text": (
                            "Sos un asesor agronomo. Describi brevemente (2-4 oraciones) lo que ves en esta foto "
                            "tomada durante una recorrida de campo: cultivo, estado fenologico, malezas, plagas, "
                            "enfermedades, suelo, infraestructura o lo que sea relevante. Se conciso y tecnico."
                        )
                    }
                ]
            }
        ],
        max_tokens=400
    )
    return response.choices[0].message.content.strip()

def transcribe_pdf(pdf_path):
    document = fitz.open(pdf_path)
    if len(document) > 40:
        document.close()
        raise ValueError("El PDF supera el limite de 40 paginas")
    extracted = []
    for index, page in enumerate(document):
        page_text = page.get_text("text").strip()
        if page_text:
            extracted.append(f"--- Pagina {index + 1} ---\n{page_text}")
    plain_text = "\n\n".join(extracted).strip()
    if len(plain_text) >= 100:
        document.close()
        return plain_text[:100000]

    if len(document) > 12:
        document.close()
        raise ValueError("El PDF escaneado supera el limite de 12 paginas")

    scanned_pages = []
    try:
        for index, page in enumerate(document):
            with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as temp:
                image_path = temp.name
            try:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                pixmap.save(image_path)
                page_text = transcribe_image(image_path)
                scanned_pages.append(f"--- Pagina {index + 1} ---\n{page_text}")
            finally:
                Path(image_path).unlink(missing_ok=True)
    finally:
        document.close()
    return "\n\n".join(scanned_pages)[:100000]


def clasificar_mensaje(text):
    prompt = (
        "Clasifica el siguiente mensaje en UNA de estas categorias:\n"
        "- receta: aplicacion fitosanitaria, agroquimicos, pulverizacion, lotes, cultivos, productos como roundup, harrier, etc.\n"
        "- cliente_nuevo: registrar nuevo cliente, oportunidad comercial, contacto nuevo, posible trabajo\n"
        "- cliente_consulta: consultar estado de clientes, ver pendientes, ver seguimientos\n"
        "- cliente_update: actualizar estado de cliente existente, cambiar fecha, marcar cerrado, perdido, etc.\n"
        "- tarea: algo para hacer, pendiente, recordatorio de accion\n"
        "- recorrida: visita a campo, recorrida tecnica, inspeccion, reporte de visita\n"
        "- presupuesto: presupuesto enviado, cotizacion, propuesta economica, factura, nota de pedido\n"
        "- compra: comprar material, insumo, herramienta\n"
        "- idea: idea de negocio, contenido, post, mejora, proyecto futuro\n\n"
        "Responde UNICAMENTE con una de estas palabras: receta, cliente_nuevo, cliente_consulta, cliente_update, tarea, recorrida, presupuesto, compra, idea\n\n"
        f"Mensaje: {text}"
    )
    response = get_openai_client().chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip().lower()

def extract_receta(text):
    today = datetime.now().strftime("%d/%m/%Y")
    prompt = (
        "Sos un asistente agronomo. Extrae datos de aplicacion fitosanitaria y responde SOLO JSON puro:\n"
        "{\n"
        f'  "fecha": "usa {today} si no menciona",\n'
        '  "campo": "nombre del campo",\n'
        '  "cultivo": "cultivo",\n'
        '  "lote": "nombre del lote",\n'
        '  "labor": "Pulverizacion",\n'
        '  "superficie": "numero o null",\n'
        '  "productos": [\n'
        '    {"producto": "nombre", "dosis": "solo numero", "unidad": "kg/ha o L/ha o cc/ha", "orden_carga": "numero o null"}\n'
        '  ]\n'
        '}\n'
        'REGLAS: orden de carga: primero=1, segundo=2, etc. Dosis solo numero. Solo JSON sin markdown.\n'
        f'Mensaje: {text}'
    )
    response = get_openai_client().chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    raw = response.choices[0].message.content.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)

def extract_cliente_nuevo(text):
    today = datetime.now().strftime("%d/%m/%Y")
    fecha_default = (datetime.now() + timedelta(days=3)).strftime("%d/%m/%Y")
    prompt = (
        "Extrae datos de este nuevo cliente y responde SOLO JSON puro:\n"
        "{\n"
        f'  "fecha": "{today}",\n'
        '  "cliente": "nombre del cliente",\n'
        '  "empresa": "empresa o establecimiento o null",\n'
        '  "zona": "zona o localidad o null",\n'
        '  "provincia": "provincia o null",\n'
        '  "contacto": "nombre de contacto o null",\n'
        '  "telefono": "telefono o null",\n'
        '  "email": "email o null",\n'
        '  "origen": "como llego: recomendacion, instagram, web, conocido, etc. o null",\n'
        '  "necesidad": "que necesita o null",\n'
        '  "tipo_trabajo": "aguadas, caminos, apotreramiento, topografia, pasturas, asesoramiento integral, otro o null",\n'
        '  "estado": "nuevo, contactado, reunion pendiente, presupuesto pendiente, presupuesto enviado, en seguimiento, cerrado, perdido",\n'
        '  "proxima_accion": "que hay que hacer o null",\n'
        f'  "fecha_seguimiento": "fecha DD/MM/YYYY, si no menciona usa {fecha_default}",\n'
        '  "presupuesto": "monto o pendiente",\n'
        '  "probabilidad_cierre": "alta, media, baja o null",\n'
        '  "prioridad": "alta, media, baja",\n'
        '  "observaciones": "cualquier dato extra o null"\n'
        '}\n'
        f'Hoy es {today}. Mensaje: {text}'
    )
    response = get_openai_client().chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    raw = response.choices[0].message.content.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)

def extract_cliente_update(text):
    prompt = (
        "El usuario quiere actualizar un cliente. Responde SOLO JSON puro:\n"
        "{\n"
        '  "cliente": "nombre del cliente a actualizar",\n'
        '  "nuevo_estado": "nuevo estado o null",\n'
        '  "proxima_accion": "nueva accion o null",\n'
        '  "fecha_seguimiento": "nueva fecha DD/MM/YYYY o null",\n'
        '  "observaciones": "nueva observacion o null"\n'
        '}\n'
        f'Mensaje: {text}'
    )
    response = get_openai_client().chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    raw = response.choices[0].message.content.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)

def extract_tarea(text):
    today = datetime.now().strftime("%d/%m/%Y")
    fecha_default = (datetime.now() + timedelta(days=1)).strftime("%d/%m/%Y")
    prompt = (
        "Extrae datos de esta tarea y responde SOLO JSON puro:\n"
        "{\n"
        f'  "fecha": "{today}",\n'
        '  "tarea": "descripcion de la tarea",\n'
        '  "cliente": "cliente relacionado o null",\n'
        '  "categoria": "comercial, tecnico, compra, administrativo, contenido",\n'
        '  "responsable": "nombre o Lucas si no menciona",\n'
        '  "estado": "pendiente",\n'
        '  "prioridad": "alta, media, baja",\n'
        f'  "fecha_limite": "fecha DD/MM/YYYY, si no menciona usa {fecha_default}",\n'
        '  "observaciones": "extra o null"\n'
        '}\n'
        f'Mensaje: {text}'
    )
    response = get_openai_client().chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    raw = response.choices[0].message.content.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)

def extract_recorrida(text):
    today = datetime.now().strftime("%d/%m/%Y")
    fecha_default = (datetime.now() + timedelta(days=14)).strftime("%d/%m/%Y")
    prompt = (
        "Extrae datos de esta recorrida de campo y responde SOLO JSON puro:\n"
        "{\n"
        f'  "fecha": "{today}",\n'
        '  "cliente": "nombre del cliente o campo",\n'
        '  "campo": "nombre del campo o lote o null",\n'
        '  "zona": "zona o localidad o null",\n'
        '  "resumen": "resumen general de la visita",\n'
        '  "problemas": "problemas detectados o null",\n'
        '  "recomendaciones": "recomendaciones o null",\n'
        '  "urgencia": "alta, media, baja",\n'
        f'  "proxima_visita": "fecha DD/MM/YYYY, si no menciona usa {fecha_default}",\n'
        '  "observaciones": "extra o null"\n'
        '}\n'
        f'Hoy es {today}. Mensaje: {text}'
    )
    response = get_openai_client().chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    raw = response.choices[0].message.content.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)

def extract_presupuesto(text):
    today = datetime.now().strftime("%d/%m/%Y")
    prompt = (
        "Extrae datos de este presupuesto y responde SOLO JSON puro:\n"
        "{\n"
        f'  "fecha": "{today}",\n'
        '  "cliente": "nombre del cliente",\n'
        '  "trabajo": "tipo de trabajo",\n'
        '  "descripcion": "descripcion del trabajo o null",\n'
        '  "honorarios": "monto o 0",\n'
        '  "viaticos": "monto o 0",\n'
        '  "total": "monto total o 0",\n'
        '  "estado": "borrador, enviado, aprobado, rechazado",\n'
        f'  "fecha_envio": "fecha DD/MM/YYYY o {today}",\n'
        '  "fecha_respuesta": "fecha esperada o pendiente",\n'
        '  "observaciones": "extra o null"\n'
        '}\n'
        f'Hoy es {today}. Mensaje: {text}'
    )
    response = get_openai_client().chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    raw = response.choices[0].message.content.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)

def extract_compra(text):
    today = datetime.now().strftime("%d/%m/%Y")
    prompt = (
        "Extrae datos de esta compra y responde SOLO JSON puro:\n"
        "{\n"
        f'  "fecha": "{today}",\n'
        '  "cliente_obra": "cliente u obra relacionada o null",\n'
        '  "material": "nombre del material o producto",\n'
        '  "cantidad": "cantidad o null",\n'
        '  "unidad": "unidad de medida o null",\n'
        '  "proveedor": "proveedor o a definir",\n'
        '  "precio_unitario": "precio o 0",\n'
        '  "total": "total o 0",\n'
        '  "estado": "a cotizar, cotizado, pedido, recibido",\n'
        '  "observaciones": "extra o null"\n'
        '}\n'
        f'Mensaje: {text}'
    )
    response = get_openai_client().chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    raw = response.choices[0].message.content.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)

def extract_idea(text):
    today = datetime.now().strftime("%d/%m/%Y")
    prompt = (
        "Extrae datos de esta idea y responde SOLO JSON puro:\n"
        "{\n"
        f'  "fecha": "{today}",\n'
        '  "tipo": "contenido, negocio, mejora, producto, proceso u otro",\n'
        '  "idea": "descripcion de la idea",\n'
        '  "cliente_tema": "cliente o tema relacionado o general",\n'
        '  "estado": "nueva",\n'
        '  "observaciones": "extra o null"\n'
        '}\n'
        f'Mensaje: {text}'
    )
    response = get_openai_client().chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    raw = response.choices[0].message.content.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)

def get_clientes_activos():
    try:
        hoja = get_google_sheet().worksheet("clientes")
        return hoja.get_all_records()
    except Exception as e:
        logger.error(f"Error obteniendo clientes: {e}")
        return []

def buscar_y_actualizar_cliente(nombre, nuevo_estado=None, proxima_accion=None, fecha_seguimiento=None, observaciones=None):
    try:
        hoja = get_google_sheet().worksheet("clientes")
        clientes = hoja.col_values(2)
        nombre_lower = nombre.strip().lower()
        for i, c in enumerate(clientes):
            if c.strip().lower() == nombre_lower:
                fila = i + 1
                if nuevo_estado:
                    hoja.update_cell(fila, 12, nuevo_estado)
                if proxima_accion:
                    hoja.update_cell(fila, 13, proxima_accion)
                if fecha_seguimiento:
                    hoja.update_cell(fila, 14, fecha_seguimiento)
                if observaciones:
                    hoja.update_cell(fila, 18, observaciones)
                return True
        return False
    except Exception as e:
        logger.error(f"Error actualizando cliente: {e}")
        return False

def calcular_consumo(dosis, superficie):
    try:
        return round(float(str(dosis).replace(",", ".")) * float(str(superficie).replace(",", ".")), 2)
    except:
        return ""

def save_receta(worksheet, data, receta_num):
    rows = []
    for producto in data["productos"]:
        superficie = data.get("superficie", "")
        dosis = producto.get("dosis", "")
        consumo = calcular_consumo(dosis, superficie)
        row = [
            data.get("fecha", ""),
            data.get("campo", ""),
            data.get("cultivo", ""),
            data.get("lote", ""),
            data.get("labor", "Pulverizacion"),
            superficie,
            producto.get("producto", ""),
            dosis,
            producto.get("unidad", ""),
            producto.get("orden_carga", ""),
            receta_num,
            "",
            consumo
        ]
        rows.append(row)
    for row in rows:
        worksheet.append_row(row)
    return rows

# ============================================================
# RECORRIDAS: resumen con GPT y generacion de Google Doc
# ============================================================

def generar_resumen_recorrida(campo, items):
    """Usa GPT para generar resumen, problemas y recomendaciones."""
    bloques = []
    for i, item in enumerate(items, 1):
        if item["tipo"] == "texto":
            bloques.append(f"[Nota {i}] {item['texto']}")
        elif item["tipo"] == "audio":
            bloques.append(f"[Audio {i} transcripto] {item['texto']}")
        elif item["tipo"] == "foto":
            bloques.append(f"[Foto {i}] {item['texto']}")
    contenido = "\n\n".join(bloques)

    prompt = (
        "Sos un asesor agronomo. Con las siguientes notas, audios y descripciones de fotos "
        f"tomadas durante una recorrida en el campo '{campo}', generá un informe tecnico claro. "
        "Responde SOLO JSON puro:\n"
        "{\n"
        '  "resumen": "parrafo narrativo con lo observado en la recorrida",\n'
        '  "problemas": ["lista", "de problemas", "detectados"],\n'
        '  "recomendaciones": ["lista", "de recomendaciones", "accionables"],\n'
        '  "urgencia": "alta, media o baja"\n'
        "}\n\n"
        f"Notas:\n{contenido}"
    )
    response = get_openai_client().chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    raw = response.choices[0].message.content.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)

def crear_docx_recorrida(campo, fecha_str, resumen_data, items, output_path):
    """Crea un archivo .docx con el informe completo y fotos embebidas."""
    doc = DocxDocument()

    # Titulo principal
    h = doc.add_heading("RECORRIDA DE CAMPO", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Datos generales
    p = doc.add_paragraph()
    p.add_run("Campo: ").bold = True
    p.add_run(campo)

    p = doc.add_paragraph()
    p.add_run("Fecha: ").bold = True
    p.add_run(fecha_str)

    p = doc.add_paragraph()
    p.add_run("Urgencia: ").bold = True
    p.add_run(str(resumen_data.get("urgencia", "media")).upper())

    doc.add_paragraph()  # espacio

    # Resumen
    doc.add_heading("RESUMEN", level=1)
    doc.add_paragraph(resumen_data.get("resumen", ""))

    # Problemas
    doc.add_heading("PROBLEMAS DETECTADOS", level=1)
    problemas = resumen_data.get("problemas", []) or []
    if problemas:
        for p_item in problemas:
            doc.add_paragraph(str(p_item), style="List Bullet")
    else:
        doc.add_paragraph("Sin problemas destacados", style="List Bullet")

    # Recomendaciones
    doc.add_heading("RECOMENDACIONES", level=1)
    recos = resumen_data.get("recomendaciones", []) or []
    if recos:
        for r in recos:
            doc.add_paragraph(str(r), style="List Bullet")
    else:
        doc.add_paragraph("Sin recomendaciones", style="List Bullet")

    # Notas y audios
    doc.add_heading("NOTAS Y AUDIOS", level=1)
    hay_notas = False
    for i, item in enumerate(items, 1):
        if item["tipo"] == "texto":
            p = doc.add_paragraph()
            p.add_run(f"Nota {i}: ").bold = True
            p.add_run(item["texto"])
            hay_notas = True
        elif item["tipo"] == "audio":
            p = doc.add_paragraph()
            p.add_run(f"Audio {i}: ").bold = True
            p.add_run(item["texto"])
            hay_notas = True
    if not hay_notas:
        doc.add_paragraph("(Sin notas ni audios)")

    # Fotos con descripcion
    fotos = [it for it in items if it["tipo"] == "foto"]
    doc.add_heading("FOTOS", level=1)
    if fotos:
        for i, foto in enumerate(fotos, 1):
            p = doc.add_paragraph()
            p.add_run(f"Foto {i}: ").bold = True
            p.add_run(foto.get("texto", ""))
            try:
                if foto.get("foto_path") and os.path.exists(foto["foto_path"]):
                    doc.add_picture(foto["foto_path"], width=Inches(5.0))
            except Exception as e:
                logger.warning(f"No se pudo insertar foto {i}: {e}")
            doc.add_paragraph()  # espacio entre fotos
    else:
        doc.add_paragraph("(Sin fotos)")

    # Firma al pie del documento
    doc.add_paragraph()
    doc.add_paragraph()
    firma1 = doc.add_paragraph()
    firma1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run1 = firma1.add_run("Informe preparado por")
    run1.italic = True

    firma2 = doc.add_paragraph()
    firma2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = firma2.add_run("Ing. Agr. Lucas Estecho")
    run2.bold = True

    firma3 = doc.add_paragraph()
    firma3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    firma3.add_run("M.P. 2009")

    doc.save(output_path)
    return output_path

# ============================================================
# Recordatorios
# ============================================================

async def enviar_recordatorios(context):
    try:
        today = datetime.now().date()
        clientes = get_clientes_activos()
        pendientes_hoy = []
        atrasados = []

        for c in clientes:
            estado = str(c.get("Estado", "")).lower()
            if estado in ["cerrado", "perdido"]:
                continue
            fecha_seg = c.get("Fecha seguimiento", "")
            cliente = c.get("Cliente", "")
            proxima = c.get("Proxima accion", "")
            if fecha_seg:
                try:
                    fecha = datetime.strptime(fecha_seg, "%d/%m/%Y").date()
                    dias_diff = (today - fecha).days
                    if fecha == today:
                        pendientes_hoy.append("- " + cliente + ": " + proxima)
                    elif dias_diff > 0:
                        atrasados.append("- " + cliente + " (hace " + str(dias_diff) + " dias): " + proxima)
                except:
                    pass

        if not pendientes_hoy and not atrasados:
            return

        msg = "Buenos dias! Resumen comercial:\n\n"
        if pendientes_hoy:
            msg += "PARA HOY:\n" + "\n".join(pendientes_hoy) + "\n\n"
        if atrasados:
            msg += "ATRASADOS:\n" + "\n".join(atrasados) + "\n"

        await context.bot.send_message(chat_id=MY_CHAT_ID, text=msg)

    except Exception as e:
        logger.error(f"Error en recordatorios: {e}")

# ============================================================
# Comandos de recorrida
# ============================================================

async def cmd_recorrida_inicio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.message.chat_id
    args = context.args
    if not args:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Usa: /recorrida_inicio <nombre del campo>\nEjemplo: /recorrida_inicio La Esperanza"
        )
        return
    campo = " ".join(args).strip()
    recorridas_activas[chat_id] = {
        "campo": campo,
        "inicio": datetime.now(),
        "items": []
    }
    await context.bot.send_message(
        chat_id=chat_id,
        text=(
            f"Recorrida iniciada en: {campo}\n\n"
            "Mandame notas, audios o fotos durante la recorrida. "
            "Cuando termines enviá /cerrar_recorrida y te armo el Google Doc."
        )
    )

async def cmd_recorrida_cancelar(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.message.chat_id
    if chat_id in recorridas_activas:
        del recorridas_activas[chat_id]
        await context.bot.send_message(chat_id=chat_id, text="Recorrida cancelada. No se guardo nada.")
    else:
        await context.bot.send_message(chat_id=chat_id, text="No hay recorrida activa.")

async def cmd_cerrar_recorrida(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.message.chat_id
    if chat_id not in recorridas_activas:
        await context.bot.send_message(chat_id=chat_id, text="No hay recorrida activa. Iniciá una con /recorrida_inicio <campo>")
        return

    sesion = recorridas_activas[chat_id]
    campo = sesion["campo"]
    items = sesion["items"]

    if not items:
        await context.bot.send_message(chat_id=chat_id, text="No hay items cargados en la recorrida. Cancelada.")
        del recorridas_activas[chat_id]
        return

    await context.bot.send_message(chat_id=chat_id, text="Generando informe...")

    docx_path = None
    try:
        fecha_str = sesion["inicio"].strftime("%d/%m/%Y")
        fecha_archivo = sesion["inicio"].strftime("%Y-%m-%d")
        resumen_data = generar_resumen_recorrida(campo, items)

        # Crear el .docx en un archivo temporal
        campo_safe = "".join(c if c.isalnum() or c in " _-" else "_" for c in campo).strip().replace(" ", "_")
        nombre_archivo = f"Recorrida_{campo_safe}_{fecha_archivo}.docx"
        tmp_dir = tempfile.gettempdir()
        docx_path = os.path.join(tmp_dir, nombre_archivo)
        crear_docx_recorrida(campo, fecha_str, resumen_data, items, docx_path)

        # Guardar fila resumen en la hoja 'recorridas'
        try:
            sh = get_google_sheet()
            sh.worksheet("recorridas").append_row([
                fecha_str,
                "",  # cliente
                campo,
                "",  # zona
                resumen_data.get("resumen", ""),
                " | ".join(resumen_data.get("problemas", []) or []),
                " | ".join(resumen_data.get("recomendaciones", []) or []),
                resumen_data.get("urgencia", ""),
                "",  # proxima visita
                nombre_archivo  # referencia al archivo generado
            ])
        except Exception as e:
            logger.error(f"No se pudo guardar en hoja recorridas: {e}")

        # Enviar el archivo docx por Telegram
        with open(docx_path, "rb") as f:
            await context.bot.send_document(
                chat_id=chat_id,
                document=f,
                filename=nombre_archivo,
                caption=f"Informe de recorrida: {campo} - {fecha_str}"
            )

        # Enviar el resumen formateado como mensaje
        problemas_txt = "\n".join([f"• {p}" for p in resumen_data.get("problemas", [])]) or "• Ninguno"
        recos_txt = "\n".join([f"• {r}" for r in resumen_data.get("recomendaciones", [])]) or "• Ninguna"

        respuesta = (
            f"✅ Recorrida cerrada!\n\n"
            f"📍 Campo: {campo}\n"
            f"📅 Fecha: {fecha_str}\n"
            f"📝 Items: {len(items)}\n"
            f"⚠️ Urgencia: {str(resumen_data.get('urgencia', 'media')).upper()}\n\n"
            f"📋 RESUMEN\n{resumen_data.get('resumen', '')}\n\n"
            f"❗ PROBLEMAS\n{problemas_txt}\n\n"
            f"💡 RECOMENDACIONES\n{recos_txt}"
        )
        await context.bot.send_message(chat_id=chat_id, text=respuesta)

        # Limpiar archivos temporales (fotos + docx)
        for it in items:
            if it["tipo"] == "foto" and it.get("foto_path") and os.path.exists(it["foto_path"]):
                try:
                    os.unlink(it["foto_path"])
                except:
                    pass
        if docx_path and os.path.exists(docx_path):
            try:
                os.unlink(docx_path)
            except:
                pass

        del recorridas_activas[chat_id]

    except Exception as e:
        logger.error(f"Error cerrando recorrida: {e}")
        await context.bot.send_message(chat_id=chat_id, text=f"Error al cerrar recorrida: {e}")
        # Limpiar el docx si quedo creado
        if docx_path and os.path.exists(docx_path):
            try:
                os.unlink(docx_path)
            except:
                pass

# ============================================================
# Handler principal
# ============================================================

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    message = update.message
    chat_id = message.chat_id

    # Si hay recorrida activa para este chat -> acumular en vez de procesar normal
    if chat_id in recorridas_activas:
        try:
            sesion = recorridas_activas[chat_id]

            if message.voice:
                await context.bot.send_message(chat_id=chat_id, text="Transcribiendo audio de recorrida...")
                file = await context.bot.get_file(message.voice.file_id)
                with tempfile.NamedTemporaryFile(suffix=".ogg", delete=False) as tmp:
                    await file.download_to_drive(tmp.name)
                    texto = transcribe_audio(tmp.name)
                    os.unlink(tmp.name)
                sesion["items"].append({"tipo": "audio", "texto": texto})
                await context.bot.send_message(
                    chat_id=chat_id,
                    text=f"Audio agregado ({len(sesion['items'])} items). Seguí mandando o /cerrar_recorrida"
                )
                return

            if message.photo:
                await context.bot.send_message(chat_id=chat_id, text="Analizando foto de recorrida...")
                photo = message.photo[-1]
                file = await context.bot.get_file(photo.file_id)
                # Guardamos la foto en disco para subirla a Drive despues
                tmp = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
                tmp.close()
                await file.download_to_drive(tmp.name)
                descripcion = describir_imagen_recorrida(tmp.name)
                caption = message.caption or ""
                texto_final = descripcion if not caption else f"{caption}. {descripcion}"
                sesion["items"].append({"tipo": "foto", "texto": texto_final, "foto_path": tmp.name})
                await context.bot.send_message(
                    chat_id=chat_id,
                    text=f"Foto agregada ({len(sesion['items'])} items). Seguí mandando o /cerrar_recorrida"
                )
                return

            if message.text:
                sesion["items"].append({"tipo": "texto", "texto": message.text})
                await context.bot.send_message(
                    chat_id=chat_id,
                    text=f"Nota agregada ({len(sesion['items'])} items). Seguí mandando o /cerrar_recorrida"
                )
                return

            # Otros tipos: no los acumulamos en recorrida
            await context.bot.send_message(
                chat_id=chat_id,
                text="En modo recorrida solo acepto texto, audio o fotos. Usa /cerrar_recorrida para terminar."
            )
            return

        except Exception as e:
            logger.error(f"Error en modo recorrida: {e}")
            await context.bot.send_message(chat_id=chat_id, text=f"Error: {e}")
            return

    # --- Modo normal (sin recorrida activa) ---
    try:
        text = None
        es_archivo = False

        if message.voice:
            await context.bot.send_message(chat_id=chat_id, text="Transcribiendo audio...")
            file = await context.bot.get_file(message.voice.file_id)
            with tempfile.NamedTemporaryFile(suffix=".ogg", delete=False) as tmp:
                await file.download_to_drive(tmp.name)
                text = transcribe_audio(tmp.name)
                os.unlink(tmp.name)

        elif message.photo:
            await context.bot.send_message(chat_id=chat_id, text="Leyendo imagen...")
            photo = message.photo[-1]
            file = await context.bot.get_file(photo.file_id)
            with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
                await file.download_to_drive(tmp.name)
                text = transcribe_image(tmp.name)
                os.unlink(tmp.name)
            es_archivo = True

        elif message.document:
            mime = message.document.mime_type or ""
            if mime.startswith("image/"):
                await context.bot.send_message(chat_id=chat_id, text="Leyendo imagen...")
                file = await context.bot.get_file(message.document.file_id)
                with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
                    await file.download_to_drive(tmp.name)
                    text = transcribe_image(tmp.name)
                    os.unlink(tmp.name)
                es_archivo = True
            elif mime == "application/pdf":
                await context.bot.send_message(chat_id=chat_id, text="Leyendo PDF...")
                file = await context.bot.get_file(message.document.file_id)
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                    await file.download_to_drive(tmp.name)
                    text = transcribe_pdf(tmp.name)
                    os.unlink(tmp.name)
                es_archivo = True
            else:
                await context.bot.send_message(chat_id=chat_id, text="Formato no soportado. Manda texto, audio, imagen o PDF.")
                return

        elif message.text:
            text = message.text

        else:
            await context.bot.send_message(chat_id=chat_id, text="Solo puedo procesar texto, audio, imagenes o PDFs.")
            return

        await context.bot.send_message(chat_id=chat_id, text="Procesando...")

        if es_archivo and message.caption:
            text = message.caption + "\n\n" + text

        categoria = clasificar_mensaje(text)
        sh = get_google_sheet()

        if categoria == "receta":
            data = extract_receta(text)
            data["campo"] = (data.get("campo") or "").lower()
            data["lote"] = (data.get("lote") or "").lower()
            data["cultivo"] = (data.get("cultivo") or "").lower()

            superficie_desde_hoja2 = False
            if not data.get("superficie") or data.get("superficie") == "null":
                sup = get_superficie_from_hoja2(data.get("lote", ""))
                if sup:
                    data["superficie"] = sup
                    superficie_desde_hoja2 = True

            worksheet = sh.worksheet("Hoja 1")
            receta_num = get_next_receta_number(worksheet)
            rows = save_receta(worksheet, data, receta_num)

            productos_texto = "\n".join([
                "  - " + r[6] + ": " + str(r[7]) + " " + r[8] + " | orden: " + str(r[9]) + " | consumo: " + str(r[12])
                for r in rows
            ])
            sup_label = str(data.get("superficie")) + " ha (desde registro de lotes)" if superficie_desde_hoja2 else str(data.get("superficie")) + " ha"
            respuesta = (
                "Receta #" + str(receta_num) + " guardada!\n\n"
                "Fecha: " + str(data.get("fecha")) + "\n"
                "Campo: " + str(data.get("campo")) + "\n"
                "Cultivo: " + str(data.get("cultivo")) + "\n"
                "Lote: " + str(data.get("lote")) + "\n"
                "Superficie: " + sup_label + "\n\n"
                "Productos:\n" + productos_texto
            )

        elif categoria == "cliente_nuevo":
            data = extract_cliente_nuevo(text)
            sh.worksheet("clientes").append_row([
                data.get("fecha", ""),
                data.get("cliente", ""),
                data.get("empresa", ""),
                data.get("zona", ""),
                data.get("provincia", ""),
                data.get("contacto", ""),
                data.get("telefono", ""),
                data.get("email", ""),
                data.get("origen", ""),
                data.get("necesidad", ""),
                data.get("tipo_trabajo", ""),
                data.get("estado", ""),
                data.get("proxima_accion", ""),
                data.get("fecha_seguimiento", ""),
                data.get("presupuesto", ""),
                data.get("probabilidad_cierre", ""),
                data.get("prioridad", ""),
                data.get("observaciones", "")
            ])
            respuesta = (
                "Cliente guardado!\n\n"
                "Cliente: " + str(data.get("cliente", "")) + "\n"
                "Empresa: " + str(data.get("empresa", "") or "-") + "\n"
                "Zona: " + str(data.get("zona", "") or "-") + "\n"
                "Tipo de trabajo: " + str(data.get("tipo_trabajo", "") or "-") + "\n"
                "Estado: " + str(data.get("estado", "")) + "\n"
                "Proxima accion: " + str(data.get("proxima_accion", "") or "-") + "\n"
                "Seguimiento: " + str(data.get("fecha_seguimiento", "")) + "\n"
                "Prioridad: " + str(data.get("prioridad", ""))
            )

        elif categoria == "cliente_consulta":
            clientes = get_clientes_activos()
            activos = [c for c in clientes if str(c.get("Estado", "")).lower() not in ["cerrado", "perdido"]]
            if not activos:
                respuesta = "No hay clientes activos en seguimiento."
            else:
                lineas = ["Clientes activos:\n"]
                for c in activos:
                    lineas.append(
                        "- " + str(c.get("Cliente", "")) + "\n"
                        "  Estado: " + str(c.get("Estado", "")) + "\n"
                        "  Proxima accion: " + str(c.get("Proxima accion", "")) + "\n"
                        "  Seguimiento: " + str(c.get("Fecha seguimiento", "")) + "\n"
                    )
                respuesta = "\n".join(lineas)

        elif categoria == "cliente_update":
            data = extract_cliente_update(text)
            nombre = data.get("cliente", "")
            ok = buscar_y_actualizar_cliente(
                nombre,
                nuevo_estado=data.get("nuevo_estado"),
                proxima_accion=data.get("proxima_accion"),
                fecha_seguimiento=data.get("fecha_seguimiento"),
                observaciones=data.get("observaciones")
            )
            if ok:
                respuesta = "Cliente " + nombre + " actualizado."
            else:
                respuesta = "No encontre el cliente '" + nombre + "'. Verifica el nombre."

        elif categoria == "tarea":
            data = extract_tarea(text)
            sh.worksheet("tareas").append_row([
                data.get("fecha", ""),
                data.get("tarea", ""),
                data.get("cliente", ""),
                data.get("categoria", ""),
                data.get("responsable", ""),
                data.get("estado", ""),
                data.get("prioridad", ""),
                data.get("fecha_limite", ""),
                data.get("observaciones", "")
            ])
            respuesta = (
                "Tarea guardada!\n\n"
                + str(data.get("tarea", "")) + "\n"
                "Cliente: " + str(data.get("cliente", "") or "-") + "\n"
                "Categoria: " + str(data.get("categoria", "")) + "\n"
                "Prioridad: " + str(data.get("prioridad", "")) + "\n"
                "Fecha limite: " + str(data.get("fecha_limite", ""))
            )

        elif categoria == "recorrida":
            data = extract_recorrida(text)
            sh.worksheet("recorridas").append_row([
                data.get("fecha", ""),
                data.get("cliente", ""),
                data.get("campo", ""),
                data.get("zona", ""),
                data.get("resumen", ""),
                data.get("problemas", ""),
                data.get("recomendaciones", ""),
                data.get("urgencia", ""),
                data.get("proxima_visita", ""),
                data.get("observaciones", "")
            ])
            respuesta = (
                "Recorrida guardada!\n\n"
                "Cliente: " + str(data.get("cliente", "")) + "\n"
                "Campo: " + str(data.get("campo", "") or "-") + "\n"
                "Resumen: " + str(data.get("resumen", "")) + "\n"
                "Urgencia: " + str(data.get("urgencia", "")) + "\n"
                "Proxima visita: " + str(data.get("proxima_visita", ""))
                + "\n\nTip: para recorridas con muchas notas/fotos usa /recorrida_inicio <campo>"
            )

        elif categoria == "presupuesto":
            data = extract_presupuesto(text)
            sh.worksheet("presupuestos").append_row([
                data.get("fecha", ""),
                data.get("cliente", ""),
                data.get("trabajo", ""),
                data.get("descripcion", ""),
                data.get("honorarios", ""),
                data.get("viaticos", ""),
                data.get("total", ""),
                data.get("estado", ""),
                data.get("fecha_envio", ""),
                data.get("fecha_respuesta", ""),
                data.get("observaciones", "")
            ])
            respuesta = (
                "Presupuesto guardado!\n\n"
                "Cliente: " + str(data.get("cliente", "")) + "\n"
                "Trabajo: " + str(data.get("trabajo", "")) + "\n"
                "Total: " + str(data.get("total", "")) + "\n"
                "Estado: " + str(data.get("estado", "")) + "\n"
                "Fecha envio: " + str(data.get("fecha_envio", ""))
            )

        elif categoria == "compra":
            data = extract_compra(text)
            sh.worksheet("compras").append_row([
                data.get("fecha", ""),
                data.get("cliente_obra", ""),
                data.get("material", ""),
                data.get("cantidad", ""),
                data.get("unidad", ""),
                data.get("proveedor", ""),
                data.get("precio_unitario", ""),
                data.get("total", ""),
                data.get("estado", ""),
                data.get("observaciones", "")
            ])
            respuesta = (
                "Compra guardada!\n\n"
                + str(data.get("material", "")) + "\n"
                "Cantidad: " + str(data.get("cantidad", "") or "-") + " " + str(data.get("unidad", "") or "") + "\n"
                "Obra: " + str(data.get("cliente_obra", "") or "-") + "\n"
                "Estado: " + str(data.get("estado", ""))
            )

        elif categoria == "idea":
            data = extract_idea(text)
            sh.worksheet("ideas").append_row([
                data.get("fecha", ""),
                data.get("tipo", ""),
                data.get("idea", ""),
                data.get("cliente_tema", ""),
                data.get("estado", ""),
                data.get("observaciones", "")
            ])
            respuesta = (
                "Idea guardada!\n\n"
                + str(data.get("idea", "")) + "\n"
                "Tipo: " + str(data.get("tipo", "")) + "\n"
                "Tema: " + str(data.get("cliente_tema", "") or "-")
            )

        else:
            respuesta = "No pude clasificar el mensaje. Intenta ser mas especifico."

        await context.bot.send_message(chat_id=chat_id, text=respuesta)

    except Exception as e:
        logger.error(f"Error: {e}")
        await context.bot.send_message(chat_id=chat_id, text="Error al procesar: " + str(e))

def process_capataz_event(event, draft, source_text=""):
    """Ejecuta la cuadrilla y materializa trabajos internos, incluido el borrador de correo."""
    result = agent_crew.process_event(event, draft, source_text=source_text)
    email_draft = email_draft_manager.prepare(
        event,
        draft,
        result,
        source_text=source_text,
    )
    return {**result, "email_draft": email_draft}


def persist_telegram_asset(
    file_path,
    *,
    event,
    asset_type,
    file_name,
    content_type,
    transcript_text="",
):
    asset_id = f"asset-{uuid.uuid4().hex}"
    now = iso_now()
    row = {
        "id": asset_id,
        "event_id": event.get("id"),
        "client_id": event.get("client_id"),
        "client_name": event.get("client_name"),
        "source": "telegram",
        "asset_type": asset_type,
        "file_name": file_name,
        "content_type": content_type or "application/octet-stream",
        "transcript_text": str(transcript_text or "")[:20000],
        "storage_status": "local_only",
        "storage_provider": "local",
        "storage_path": "",
        "storage_public_url": "",
        "storage_error": "",
        "created_at": now,
        "updated_at": now,
    }
    if SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY and SUPABASE_BUCKET:
        date_path = argentina_now().strftime("%Y-%m-%d")
        object_path = (
            f"inbox/telegram/{date_path}/"
            f"{safe_storage_segment(event.get('client_name'), 'sin-cliente')}/"
            f"{asset_id}_{safe_storage_segment(file_name, 'archivo')}"
        )
        try:
            uploaded = upload_field_file_to_supabase(
                file_path,
                object_path,
                content_type=content_type,
            )
            row.update(
                {
                    "storage_status": "supabase_uploaded",
                    "storage_provider": "supabase",
                    "storage_path": uploaded.get("path") or object_path,
                    "storage_public_url": uploaded.get("public_url") or "",
                }
            )
        except Exception as exc:
            row.update(
                {
                    "storage_status": "supabase_error",
                    "storage_provider": "supabase",
                    "storage_path": object_path,
                    "storage_error": str(exc)[:2000],
                }
            )
    source, warning = capataz_store.save_rows("intake_assets", [row])
    if capataz_store.supabase_configured and source != "supabase":
        raise PersistentStorageError(warning or "No se pudo guardar el archivo recibido por Telegram")
    return row


async def cmd_client_profile(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if MY_CHAT_ID and update.effective_chat.id != MY_CHAT_ID:
        return
    from client_profile import build_client_profile, format_client_profile

    name = " ".join(context.args or []).strip()
    if not name:
        clients, _source, _warning = await asyncio.to_thread(capataz_store.list_clients)
        names = ", ".join(sorted(str(row.get("name") or "") for row in clients if row.get("name"))) or "sin clientes registrados"
        await context.bot.send_message(
            chat_id=update.effective_chat.id,
            text=f"Usa /cliente NOMBRE. Clientes conocidos: {names}."[:4000],
        )
        return
    try:
        profile = await asyncio.to_thread(build_client_profile, capataz_store, name)
        await context.bot.send_message(
            chat_id=update.effective_chat.id,
            text=format_client_profile(profile)[:4000],
        )
    except Exception as exc:
        await context.bot.send_message(
            chat_id=update.effective_chat.id,
            text=f"No pude armar la ficha: {str(exc)[:500]}",
        )


async def cmd_pending(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if MY_CHAT_ID and update.effective_chat.id != MY_CHAT_ID:
        return
    from capataz import format_pending_summary

    dashboard = await asyncio.to_thread(capataz_store.dashboard)
    summary = format_pending_summary(dashboard)
    await context.bot.send_message(
        chat_id=update.effective_chat.id,
        text=(summary or "Sin pendientes registrados. Cuando delegues algo en un audio, aparece aca.")[:4000],
    )


async def cmd_agriculture(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if MY_CHAT_ID and update.effective_chat.id != MY_CHAT_ID:
        return
    from agriculture import build_agriculture_overview, format_agriculture_overview

    name = " ".join(context.args or []).strip()
    try:
        overview = await asyncio.to_thread(build_agriculture_overview, capataz_store, name)
        await context.bot.send_message(
            chat_id=update.effective_chat.id,
            text=format_agriculture_overview(overview)[:4000],
        )
    except Exception as exc:
        await context.bot.send_message(
            chat_id=update.effective_chat.id,
            text=f"No pude armar el resumen agricola: {str(exc)[:500]}",
        )


async def cmd_cleanup_storage(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if MY_CHAT_ID and update.effective_chat.id != MY_CHAT_ID:
        return
    await context.bot.send_message(
        chat_id=update.effective_chat.id,
        text="Limpiando almacenamiento (Supabase + disco de Render)...",
    )
    summary = await asyncio.to_thread(run_storage_cleanup)
    archive_status = summary.get("archive_status") or {}
    counts = archive_status.get("counts") or {}
    health = summary.get("health_purge") or {}
    prune = summary.get("local_prune") or {}
    lines = ["Limpieza terminada."]
    lines.append(
        f"Archivos marcados para pasar a la PC: {summary.get('archive_candidates', 0)} "
        f"(pendientes {counts.get('pending', 0)}, ya archivados {counts.get('archived', 0)})."
    )
    lines.append(f"Restos de chequeos borrados de Supabase: {health.get('deleted', 0)}.")
    lines.append(
        f"Disco de Render liberado: {prune.get('deleted', 0)} archivos, "
        f"{prune.get('freed_mb', 0)} MB."
    )
    if counts.get("pending"):
        lines.append(
            "Los pendientes bajan a tu PC cuando corre el archivador de Windows "
            "(tarea diaria o al iniciar sesion). Recien despues de verificar la copia "
            "se borran de Supabase."
        )
    errors = [
        text for text in (
            summary.get("archive_candidates_error"),
            health.get("error"),
            prune.get("error"),
        ) if text
    ]
    if errors:
        lines.append("Advertencias: " + " | ".join(errors)[:500])
    await context.bot.send_message(
        chat_id=update.effective_chat.id, text="\n".join(lines)[:4000]
    )


async def cmd_capataz_status(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if MY_CHAT_ID and update.effective_chat.id != MY_CHAT_ID:
        return
    dashboard = await asyncio.to_thread(capataz_store.dashboard)
    cdse_status = cdse_configuration_status(validate=False)
    await context.bot.send_message(
        chat_id=update.effective_chat.id,
        text=(
            "Capataz Campo activo.\n"
            f"Agentes recientes: {len(dashboard.get('agent_activity') or [])}.\n"
            f"Borradores de correo: {len(dashboard.get('email_drafts') or [])}.\n"
            f"NDVI Sentinel: {'configurado' if cdse_status['configured'] else 'BLOQUEADO - faltan credenciales CDSE'}.\n"
            "Compartime desde WhatsApp texto, audio, foto, PDF, Word, Excel o un paquete geoespacial "
            "(ZIP con SHP/SHX/DBF/PRJ + DEM GeoTIFF).\n"
            "Usa /informes para ver los entregables que la cuadrilla puede producir."
        ),
    )


async def cmd_report_catalog(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if MY_CHAT_ID and update.effective_chat.id != MY_CHAT_ID:
        return
    catalog = public_report_catalog()
    lines = ["Entregables profesionales disponibles:"]
    for item in catalog:
        agents = ", ".join(item["agents"])
        lines.append(f"- {item['title']}\n  Cuadrilla: {agents}")
    lines.extend([
        "",
        "Pedi el resultado en lenguaje normal. Ejemplos:",
        '"Compara estos tres presupuestos y devolveme PDF y Word."',
        '"Arma una propuesta tecnica para La Susana; deja honorarios pendientes."',
        '"Evalua la compra de esta isla contra La Tigra, sin inventar precios."',
        '"Prepara un dossier tecnico para vender este campo."',
    ])
    await context.bot.send_message(chat_id=update.effective_chat.id, text="\n".join(lines)[:4000])


async def cmd_send_confirmed_email(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Send only the exact draft id explicitly confirmed by Lucas."""
    if MY_CHAT_ID and update.effective_chat.id != MY_CHAT_ID:
        return
    draft_id = str((context.args or [""])[0]).strip()
    if not draft_id:
        await context.bot.send_message(
            chat_id=update.effective_chat.id,
            text=(
                "No envie nada. Para confirmar necesito el comando completo que aparece "
                "debajo del borrador, por ejemplo: /enviar_correo email-..."
            ),
        )
        return
    try:
        sent = await asyncio.to_thread(email_draft_manager.send_confirmed, draft_id)
    except Exception as exc:
        logger.exception("No se pudo enviar el borrador confirmado %s", draft_id)
        await context.bot.send_message(
            chat_id=update.effective_chat.id,
            text=f"No envie el correo: {str(exc)[:800]}",
        )
        return
    if sent.get("already_sent"):
        message = f"Ese correo ya habia sido enviado: {sent.get('subject') or draft_id}"
    else:
        message = (
            "Correo enviado por confirmacion explicita.\n"
            f"Para: {sent.get('to_email')}\n"
            f"Asunto: {sent.get('subject')}"
        )
    await context.bot.send_message(chat_id=update.effective_chat.id, text=message)


def telegram_asset_type(file_name, content_type=""):
    name = str(file_name or "").lower()
    content_type = str(content_type or "").lower()
    if name.endswith((".tif", ".tiff")):
        return "ndvi" if "ndvi" in name else "geotiff"
    if name.endswith(".kml"):
        return "kml"
    if name.endswith((".geojson", ".json")):
        return "geojson"
    if name.endswith(".zip"):
        return "paquete_geoespacial"
    if name.endswith(".shp"):
        return "shapefile"
    if name.endswith((".shx", ".dbf", ".prj", ".cpg", ".sbn", ".sbx", ".qix")):
        return "shapefile_componente"
    if name.endswith(".xml"):
        return "metadata"
    if content_type.startswith("image/"):
        return "foto"
    if content_type.startswith("audio/"):
        return "audio"
    if name.endswith(".pdf") or content_type == "application/pdf":
        return "informe_pdf" if name.startswith("informe_") else "pdf"
    return "archivo"


async def execute_capataz_telegram_work(
    chat_id,
    context,
    text,
    assets=None,
    preface="",
    allow_consulting_report=True,
):
    """Confirma una sola entrada y ejecuta la cuadrilla con todos sus archivos."""
    assets = list(assets or [])
    if not str(text or "").strip():
        raise ValueError("No pude extraer contenido del mensaje")
    draft = await asyncio.to_thread(
        analyze_intake,
        text,
        source="telegram",
        openai_client=openai_client,
    )
    confirmed = await asyncio.to_thread(capataz_store.confirm_intake, draft, source_text=text)
    plan = await asyncio.to_thread(
        agent_crew.queue_event,
        confirmed["event"],
        draft,
        source_text=text,
    )
    for asset in assets:
        await asyncio.to_thread(
            persist_telegram_asset,
            asset["path"],
            event=confirmed["event"],
            asset_type=asset.get("asset_type") or telegram_asset_type(
                asset.get("file_name"), asset.get("content_type")
            ),
            file_name=asset.get("file_name") or Path(asset["path"]).name,
            content_type=asset.get("content_type") or "application/octet-stream",
            transcript_text=asset.get("transcript_text") or text,
        )
    await context.bot.send_message(
        chat_id=chat_id,
        text="Trabajando con: " + ", ".join(plan.get("agents") or ["Capataz"]),
    )
    processed = await asyncio.to_thread(
        process_capataz_event,
        confirmed["event"],
        draft,
        text,
    )
    generated_report = None
    report_error = None
    if allow_consulting_report:
        await context.bot.send_message(
            chat_id=chat_id,
            text=(
                "Redactando el informe final. Puede tardar varios minutos; "
                "te aviso apenas este o si algo falla."
            ),
        )
        try:
            generated_report = await asyncio.to_thread(
                generate_consulting_report,
                event=confirmed["event"],
                draft=draft,
                crew_result=processed,
                source_text=text,
                assets=assets,
                output_dir=DATA_DIR / "consulting_reports",
                logo_path=str(logo_path()),
                openai_client=openai_client,
            )
        except Exception as exc:
            report_error = str(exc)[:1000]
            logger.exception("Fallo la generacion del entregable profesional")
            await context.bot.send_message(
                chat_id=chat_id,
                text=(
                    "No te entregue un informe de relleno: la redaccion final fallo y quedo bloqueada.\n"
                    f"Detalle: {report_error}\n"
                    "La entrada y el trabajo de los agentes quedaron guardados."
                ),
            )
    if generated_report:
        await context.bot.send_message(
            chat_id=chat_id,
            text=(
                f"{generated_report.title}: armando y entregando PDF + Word. "
                f"Estado tecnico: {generated_report.status}. "
                f"Redaccion final: {generated_report.model}."
            ),
        )
        generated_files = [
            (generated_report.pdf_path, "application/pdf", "informe_pdf"),
            (generated_report.docx_path, REPORT_DOCX_CONTENT_TYPE, "informe_docx"),
        ]
        for generated_path, generated_content_type, generated_type in generated_files:
            file_name = Path(generated_path).name
            await asyncio.to_thread(
                persist_telegram_asset,
                generated_path,
                event=confirmed["event"],
                asset_type=generated_type,
                file_name=file_name,
                content_type=generated_content_type,
                transcript_text=f"Entregable generado: {generated_report.title}",
            )
            with Path(generated_path).open("rb") as report_file:
                await context.bot.send_document(
                    chat_id=chat_id,
                    document=report_file,
                    filename=file_name,
                    caption=(
                        f"{generated_report.title}\n"
                        f"Estado: {generated_report.status}. "
                        "Los datos faltantes quedan marcados; no se completaron con supuestos."
                    ),
                )
        processed["generated_report"] = {
            "playbook": generated_report.playbook_key,
            "title": generated_report.title,
            "status": generated_report.status,
            "pdf_path": generated_report.pdf_path,
            "docx_path": generated_report.docx_path,
            "missing_data": list(generated_report.missing_data),
            "model": generated_report.model,
        }
    elif report_error:
        processed["generated_report_error"] = report_error
    runs = processed.get("runs") or []
    decision = processed.get("decision") or {}
    summaries = [
        f"{run.get('agent')}: {(run.get('output') or {}).get('summary')}"
        for run in runs
        if (run.get("output") or {}).get("summary")
    ]
    response_lines = [value for value in [preface, "Trabajo terminado."] if value]
    if generated_report:
        response_lines.append(
            f"Entregable real: PDF + Word ({generated_report.playbook_key}, "
            f"{generated_report.status}, {generated_report.model})."
        )
        if generated_report.missing_data:
            response_lines.append(
                "No invente lo que falta. Pendientes principales:\n- "
                + "\n- ".join(generated_report.missing_data[:5])
            )
    response_lines.extend(summaries[:6])
    if decision:
        response_lines.append("Decision lista para revisar en Capataz Campo.")
    email_draft = processed.get("email_draft")
    if email_draft:
        email_status = (
            "guardado en Gmail"
            if email_draft.get("status") == "gmail_created"
            else "preparado en Capataz"
        )
        response_lines.append(
            f"Correo {email_status}: {email_draft.get('subject')}\n"
            f"Para: {email_draft.get('to_email') or 'FALTA DESTINATARIO'}\n"
            "No fue enviado. Revisalo en Gmail y, solo si esta correcto, confirma con:\n"
            f"/enviar_correo {email_draft.get('id')}"
        )
    response_lines.append("Abrir: https://bot-agro-campo.onrender.com/campo")
    await context.bot.send_message(chat_id=chat_id, text="\n\n".join(response_lines)[:4000])
    return processed


def _cancel_geo_batch_jobs(context, chat_id):
    if not context.job_queue:
        return
    for name in (f"geo-batch-{chat_id}", f"geo-expire-{chat_id}"):
        for job in context.job_queue.get_jobs_by_name(name):
            job.schedule_removal()


def _cancel_document_batch_jobs(context, chat_id):
    if not context.job_queue:
        return
    for name in (f"doc-batch-{chat_id}", f"doc-expire-{chat_id}"):
        for job in context.job_queue.get_jobs_by_name(name):
            job.schedule_removal()


def _cleanup_telegram_assets(assets):
    for asset in assets or []:
        path = asset.get("path") if isinstance(asset, dict) else getattr(asset, "path", None)
        if path:
            Path(path).unlink(missing_ok=True)


async def process_telegram_geo_batch(chat_id, context):
    batch = telegram_geo_batches.get(chat_id)
    if not batch or batch.get("processing"):
        return
    instruction = str(batch.get("instruction") or "").strip()
    if not instruction:
        if not batch.get("asked_for_instruction"):
            batch["asked_for_instruction"] = True
            await context.bot.send_message(
                chat_id=chat_id,
                text=(
                    f"Recibi {len(batch.get('assets') or [])} archivo(s) geoespacial(es). "
                    "Mandame ahora, en un mensaje de texto, que analisis queres que haga."
                ),
            )
            if context.job_queue:
                context.job_queue.run_once(
                    expire_telegram_geo_batch_job,
                    when=900,
                    data={"chat_id": chat_id},
                    name=f"geo-expire-{chat_id}",
                    job_kwargs={"misfire_grace_time": None},  # Render Free es lento: ejecutar aunque llegue tarde
                    chat_id=chat_id,
                )
        return
    batch["processing"] = True
    source_assets = list(batch.get("assets") or [])
    generated_assets = []
    try:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Paquete completo. Calculando series NDVI, lotes, relieve y productos geoespaciales...",
        )
        package = await asyncio.to_thread(
            analyze_geospatial_package,
            [
                GeoAsset(asset["path"], asset["file_name"], asset.get("content_type") or "application/octet-stream")
                for asset in source_assets
            ],
            instruction,
        )
        for asset in package.get("generated_assets") or []:
            generated_assets.append({
                "path": asset.path,
                "file_name": asset.file_name,
                "content_type": asset.content_type,
                "asset_type": telegram_asset_type(asset.file_name, asset.content_type),
            })
        calculated_text = package["summary_text"]
        reports = [asset for asset in generated_assets if asset.get("asset_type") == "informe_pdf"]
        for report in reports:
            is_ndvi_report = "ndvi" in str(report.get("file_name") or "").lower()
            with Path(report["path"]).open("rb") as report_file:
                await context.bot.send_document(
                    chat_id=chat_id,
                    document=report_file,
                    filename=report["file_name"],
                    caption=(
                        "Informe NDVI multianual por lote listo. Incluye forestaciones separadas, "
                        "estabilidad, cambio reciente, ranking y limitaciones."
                        if is_ndvi_report
                        else "Informe topografico profesional listo. Incluye elevacion, pendientes, "
                        "cuencas, vias de escurrimiento, alternativas, economia y limitaciones."
                    ),
                )
        await context.bot.send_message(
            chat_id=chat_id,
            text=("Calculos terminados:\n\n" + calculated_text)[:3900],
        )
        full_text = f"PEDIDO DE LUCAS:\n{instruction}\n\n{calculated_text}"
        for asset in source_assets + generated_assets:
            asset["transcript_text"] = full_text
        await execute_capataz_telegram_work(
            chat_id,
            context,
            full_text,
            assets=source_assets + generated_assets,
            preface="El motor geoespacial termino los calculos y la cuadrilla audito el resultado.",
            allow_consulting_report=False,
        )
    except Exception as exc:
        logger.exception("Error procesando paquete geoespacial de Telegram")
        await context.bot.send_message(
            chat_id=chat_id,
            text=f"No pude terminar el paquete geoespacial: {str(exc)[:1000]}",
        )
    finally:
        telegram_geo_batches.pop(chat_id, None)
        _cleanup_telegram_assets(source_assets + generated_assets)


async def process_telegram_geo_batch_job(context):
    await process_telegram_geo_batch(context.job.data["chat_id"], context)


async def expire_telegram_geo_batch_job(context):
    chat_id = context.job.data["chat_id"]
    batch = telegram_geo_batches.pop(chat_id, None)
    if not batch or batch.get("processing"):
        return
    _cleanup_telegram_assets(batch.get("assets") or [])
    await context.bot.send_message(
        chat_id=chat_id,
        text="El paquete geoespacial vencio despues de 15 minutos sin instrucciones. Volve a compartirlo cuando quieras.",
    )


async def process_telegram_document_batch(chat_id, context):
    batch = telegram_document_batches.get(chat_id)
    if not batch or batch.get("processing"):
        return
    instruction = str(batch.get("instruction") or "").strip()
    if not instruction:
        if not batch.get("asked_for_instruction"):
            batch["asked_for_instruction"] = True
            await context.bot.send_message(
                chat_id=chat_id,
                text=(
                    f"Lei {len(batch.get('assets') or [])} archivo(s). Decime ahora que queres que haga "
                    "con el conjunto: comparar, auditar, preparar propuesta, informe, presupuesto, etc."
                ),
            )
            if context.job_queue:
                context.job_queue.run_once(
                    expire_telegram_document_batch_job,
                    when=900,
                    data={"chat_id": chat_id},
                    name=f"doc-expire-{chat_id}",
                    job_kwargs={"misfire_grace_time": None},  # Render Free es lento: ejecutar aunque llegue tarde
                    chat_id=chat_id,
                )
        return
    batch["processing"] = True
    assets = list(batch.get("assets") or [])
    try:
        blocks = [f"PEDIDO DE LUCAS:\n{instruction}"]
        for asset in assets:
            blocks.append(
                f"--- ARCHIVO: {asset.get('file_name') or 'sin nombre'} ---\n"
                f"{asset.get('extracted_text') or 'Sin texto extraido'}"
            )
        combined_text = "\n\n".join(blocks)[:280000]
        for asset in assets:
            asset["transcript_text"] = combined_text
        await execute_capataz_telegram_work(
            chat_id,
            context,
            combined_text,
            assets=assets,
            preface=f"La cuadrilla trabajo sobre {len(assets)} archivo(s) como un solo expediente.",
        )
    except Exception as exc:
        logger.exception("Error procesando expediente documental de Telegram")
        await context.bot.send_message(
            chat_id=chat_id,
            text=f"No pude terminar el expediente: {str(exc)[:1000]}",
        )
    finally:
        telegram_document_batches.pop(chat_id, None)
        _cleanup_telegram_assets(assets)


async def process_telegram_document_batch_job(context):
    await process_telegram_document_batch(context.job.data["chat_id"], context)


async def expire_telegram_document_batch_job(context):
    chat_id = context.job.data["chat_id"]
    batch = telegram_document_batches.pop(chat_id, None)
    if not batch or batch.get("processing"):
        return
    _cleanup_telegram_assets(batch.get("assets") or [])
    await context.bot.send_message(
        chat_id=chat_id,
        text="El expediente vencio despues de 15 minutos sin instrucciones. Volve a compartir los archivos cuando quieras.",
    )


async def handle_capataz_telegram_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    message = update.message
    if not message:
        return
    chat_id = message.chat_id
    if MY_CHAT_ID and chat_id != MY_CHAT_ID:
        await context.bot.send_message(chat_id=chat_id, text="Este Capataz es privado.")
        return

    temp_path = None
    try:
        text = str(message.text or message.caption or "").strip()
        pending_geo = telegram_geo_batches.get(chat_id)
        if message.text and pending_geo:
            pending_geo["instruction"] = text
            _cancel_geo_batch_jobs(context, chat_id)
            if context.job_queue:
                context.job_queue.run_once(
                    process_telegram_geo_batch_job,
                    when=8,
                    data={"chat_id": chat_id},
                    name=f"geo-batch-{chat_id}",
                    job_kwargs={"misfire_grace_time": None},  # Render Free es lento: ejecutar aunque llegue tarde
                    chat_id=chat_id,
                )
            else:
                await process_telegram_geo_batch(chat_id, context)
            return
        pending_documents = telegram_document_batches.get(chat_id)
        if message.text and pending_documents:
            pending_documents["instruction"] = text
            _cancel_document_batch_jobs(context, chat_id)
            if context.job_queue:
                context.job_queue.run_once(
                    process_telegram_document_batch_job,
                    when=8,
                    data={"chat_id": chat_id},
                    name=f"doc-batch-{chat_id}",
                    job_kwargs={"misfire_grace_time": None},  # Render Free es lento: ejecutar aunque llegue tarde
                    chat_id=chat_id,
                )
            else:
                await process_telegram_document_batch(chat_id, context)
            return

        assets = []
        if message.voice or message.audio:
            media = message.voice or message.audio
            telegram_file = await context.bot.get_file(media.file_id)
            file_name = getattr(media, "file_name", None) or f"audio_{uuid.uuid4().hex[:8]}.ogg"
            content_type = getattr(media, "mime_type", None) or "audio/ogg"
            suffix = Path(file_name).suffix or ".ogg"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp:
                temp_path = temp.name
            await telegram_file.download_to_drive(temp_path)
            await context.bot.send_message(chat_id=chat_id, text="Recibido. Transcribiendo y asignando agentes...")
            transcription = await asyncio.to_thread(transcribe_audio, temp_path)
            text = "\n\n".join(value for value in (text, transcription) if value).strip()
            assets.append({
                "path": temp_path,
                "file_name": file_name,
                "content_type": content_type,
                "asset_type": "audio",
            })
        elif message.photo:
            telegram_file = await context.bot.get_file(message.photo[-1].file_id)
            file_name = f"foto_{uuid.uuid4().hex[:8]}.jpg"
            content_type = "image/jpeg"
            with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as temp:
                temp_path = temp.name
            await telegram_file.download_to_drive(temp_path)
            await context.bot.send_message(chat_id=chat_id, text="Recibido. Analizando la imagen y asignando agentes...")
            description = await asyncio.to_thread(transcribe_image, temp_path)
            text = "\n\n".join(value for value in (text, description) if value).strip()
            assets.append({
                "path": temp_path,
                "file_name": file_name,
                "content_type": content_type,
                "asset_type": "foto",
            })
        elif message.document:
            document = message.document
            document_instruction = text
            telegram_file = await context.bot.get_file(document.file_id)
            file_name = document.file_name or f"documento_{uuid.uuid4().hex[:8]}"
            content_type = document.mime_type or mimetypes.guess_type(file_name)[0] or "application/octet-stream"
            suffix = Path(file_name).suffix or mimetypes.guess_extension(content_type) or ".bin"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp:
                temp_path = temp.name
            await telegram_file.download_to_drive(temp_path)
            if is_geospatial_filename(file_name):
                batch = telegram_geo_batches.setdefault(chat_id, {
                    "assets": [],
                    "instruction": "",
                    "processing": False,
                    "asked_for_instruction": False,
                })
                batch["assets"].append({
                    "path": temp_path,
                    "file_name": file_name,
                    "content_type": content_type,
                    "asset_type": telegram_asset_type(file_name, content_type),
                })
                temp_path = None  # el batch se hace responsable de borrar el temporal
                if text:
                    batch["instruction"] = text
                if len(batch["assets"]) == 1:
                    await context.bot.send_message(
                        chat_id=chat_id,
                        text="Archivo geoespacial recibido. Esperando el resto del paquete...",
                    )
                _cancel_geo_batch_jobs(context, chat_id)
                if context.job_queue:
                    context.job_queue.run_once(
                        process_telegram_geo_batch_job,
                        when=8,
                        data={"chat_id": chat_id},
                        name=f"geo-batch-{chat_id}",
                        job_kwargs={"misfire_grace_time": None},  # Render Free es lento: ejecutar aunque llegue tarde
                        chat_id=chat_id,
                    )
                elif batch.get("instruction"):
                    await process_telegram_geo_batch(chat_id, context)
                return

            await context.bot.send_message(chat_id=chat_id, text="Recibido. Leyendo el archivo y asignando agentes...")
            if content_type.startswith("image/"):
                extracted = await asyncio.to_thread(transcribe_image, temp_path)
            elif content_type.startswith("audio/"):
                extracted = await asyncio.to_thread(transcribe_audio, temp_path)
            elif content_type == "application/pdf" or suffix.lower() == ".pdf":
                extracted = await asyncio.to_thread(transcribe_pdf, temp_path)
            elif suffix.lower() in {".docx", ".xlsx", ".xlsm", ".csv", ".tsv", ".txt", ".md"}:
                extracted = await asyncio.to_thread(
                    extract_office_document,
                    temp_path,
                    file_name,
                )
            else:
                raise ValueError(
                    "Formato no soportado. Compartime texto, audio, imagen, PDF, DOCX, XLSX, CSV "
                    "o un ZIP geoespacial con SHP/SHX/DBF/PRJ y DEM GeoTIFF."
                )
            text = "\n\n".join(value for value in (text, extracted) if value).strip()
            assets.append({
                "path": temp_path,
                "file_name": file_name,
                "content_type": content_type,
                "asset_type": telegram_asset_type(file_name, content_type),
                "extracted_text": extracted,
            })
            batch = telegram_document_batches.setdefault(chat_id, {
                "assets": [],
                "instruction": "",
                "processing": False,
                "asked_for_instruction": False,
            })
            batch["assets"].extend(assets)
            assets = []
            temp_path = None  # el expediente se hace responsable del temporal
            if document_instruction:
                batch["instruction"] = document_instruction
            _cancel_document_batch_jobs(context, chat_id)
            if batch.get("instruction"):
                await context.bot.send_message(
                    chat_id=chat_id,
                    text=(
                        f"Expediente: {len(batch['assets'])} archivo(s) leido(s). "
                        "Espero unos segundos por si compartis otro y despues trabaja la cuadrilla."
                    ),
                )
                if context.job_queue:
                    context.job_queue.run_once(
                        process_telegram_document_batch_job,
                        when=8,
                        data={"chat_id": chat_id},
                        name=f"doc-batch-{chat_id}",
                        job_kwargs={"misfire_grace_time": None},  # Render Free es lento: ejecutar aunque llegue tarde
                        chat_id=chat_id,
                    )
                else:
                    await process_telegram_document_batch(chat_id, context)
            else:
                await process_telegram_document_batch(chat_id, context)
            return
        await execute_capataz_telegram_work(chat_id, context, text, assets=assets)
    except Exception as exc:
        logger.exception("Error procesando entrada de Telegram")
        await context.bot.send_message(chat_id=chat_id, text=f"No pude terminar el trabajo: {str(exc)[:1000]}")
    finally:
        if temp_path:
            Path(temp_path).unlink(missing_ok=True)


def build_telegram_application():
    app = Application.builder().token(TELEGRAM_TOKEN).build()

    app.add_handler(CommandHandler("start", cmd_capataz_status))
    app.add_handler(CommandHandler("status", cmd_capataz_status))
    app.add_handler(CommandHandler("informes", cmd_report_catalog))
    app.add_handler(CommandHandler("limpiar", cmd_cleanup_storage))
    app.add_handler(CommandHandler("cliente", cmd_client_profile))
    app.add_handler(CommandHandler("agro", cmd_agriculture))
    app.add_handler(CommandHandler("pendientes", cmd_pending))
    app.add_handler(CommandHandler("enviar_correo", cmd_send_confirmed_email))
    app.add_handler(MessageHandler(
        filters.TEXT | filters.VOICE | filters.AUDIO | filters.PHOTO | filters.Document.ALL,
        handle_capataz_telegram_message,
    ))

    return app

def main():
    app = build_telegram_application()
    logger.info("Bot iniciado!")
    app.run_polling()

def safe_field_extension(upload: UploadFile, item_type: str) -> str:
    extension = Path(upload.filename or "").suffix.lower()
    allowed = {
        "foto": {".jpg", ".jpeg", ".png", ".webp", ".heic", ".heif"},
        "audio": {".webm", ".m4a", ".mp3", ".wav", ".ogg", ".mp4", ".aac"},
    }.get(item_type, set())
    if extension in allowed:
        return extension

    guessed = mimetypes.guess_extension(upload.content_type or "")
    if guessed in allowed:
        return guessed

    return ".webm" if item_type == "audio" else ".jpg"

def validate_field_upload(upload: UploadFile, item_type: str):
    content_type = str(upload.content_type or "").lower()
    extension = Path(upload.filename or "").suffix.lower()
    if item_type == "foto":
        allowed_extensions = {".jpg", ".jpeg", ".png", ".webp", ".heic", ".heif"}
        if not content_type.startswith("image/") and extension not in allowed_extensions:
            raise HTTPException(status_code=415, detail="El archivo no parece ser una foto valida")
        return MAX_PHOTO_UPLOAD_BYTES
    allowed_extensions = {".webm", ".m4a", ".mp3", ".wav", ".ogg", ".mp4", ".aac"}
    if not content_type.startswith("audio/") and extension not in allowed_extensions:
        raise HTTPException(status_code=415, detail="El archivo no parece ser un audio valido")
    return MAX_AUDIO_UPLOAD_BYTES

def save_upload_with_limit(upload: UploadFile, destination: Path, max_bytes: int):
    total = 0
    try:
        with destination.open("wb") as output:
            while True:
                chunk = upload.file.read(1024 * 1024)
                if not chunk:
                    break
                total += len(chunk)
                if total > max_bytes:
                    raise HTTPException(
                        status_code=413,
                        detail=f"Archivo demasiado grande; limite {max_bytes // (1024 * 1024)} MB",
                    )
                output.write(chunk)
    except Exception:
        destination.unlink(missing_ok=True)
        raise
    return total

@asynccontextmanager
async def fastapi_lifespan(app: FastAPI):
    telegram_app = None
    telegram_mode = "disabled"
    enable_bot = os.environ.get("ENABLE_TELEGRAM_BOT", "false").lower() in {"1", "true", "yes", "si"}

    if enable_bot:
        if not TELEGRAM_TOKEN:
            logger.warning("ENABLE_TELEGRAM_BOT=true pero falta TELEGRAM_TOKEN. FastAPI inicia sin bot.")
        else:
            telegram_app = build_telegram_application()
            await telegram_app.initialize()
            await telegram_app.start()
            app.state.telegram_app = telegram_app
            webhook_base_url = telegram_webhook_base_url()
            if webhook_base_url:
                webhook_url = f"{webhook_base_url}{TELEGRAM_WEBHOOK_PATH}"
                webhook_registered = await telegram_app.bot.set_webhook(
                    url=webhook_url,
                    secret_token=telegram_webhook_secret(),
                    allowed_updates=Update.ALL_TYPES,
                    drop_pending_updates=False,
                )
                if not webhook_registered:
                    raise RuntimeError("Telegram rechazo el registro del webhook")
                telegram_mode = "webhook"
                logger.info("Bot de Telegram iniciado en modo webhook: %s", webhook_url)
            else:
                await telegram_app.updater.start_polling(drop_pending_updates=False)
                telegram_mode = "polling"
                logger.info("Bot de Telegram iniciado en modo polling local.")

    app.state.telegram_mode = telegram_mode

    try:
        yield
    finally:
        app.state.telegram_app = None
        app.state.telegram_mode = "stopped"
        if telegram_app:
            if telegram_mode == "polling":
                await telegram_app.updater.stop()
            await telegram_app.stop()
            await telegram_app.shutdown()

fastapi_app = FastAPI(title="Capataz Campo", lifespan=fastapi_lifespan)
fastapi_app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

@fastapi_app.middleware("http")
async def protect_private_api(request: Request, call_next):
    is_private_api = (
        request.url.path.startswith("/api/")
        and request.url.path != "/api/health/campo"
    )
    if is_private_api and PRIVATE_API_REQUIRES_TOKEN and not FIELD_APP_TOKEN:
        return JSONResponse(
            status_code=503,
            content={"ok": False, "detail": "FIELD_APP_TOKEN no configurado en el servidor"},
        )
    if is_private_api and FIELD_APP_TOKEN:
        supplied = request.headers.get("X-Field-App-Token", "")
        authorization = request.headers.get("Authorization", "")
        if authorization.lower().startswith("bearer "):
            supplied = authorization[7:].strip()
        if not supplied or not secrets.compare_digest(supplied, FIELD_APP_TOKEN):
            return JSONResponse(
                status_code=401,
                content={"ok": False, "detail": "Token de Capataz Campo requerido"},
            )
    return await call_next(request)

@fastapi_app.get("/")
async def root():
    return RedirectResponse(url="/campo")

@fastapi_app.get("/campo")
async def campo():
    return FileResponse(STATIC_DIR / "index.html")

@fastapi_app.get("/sw.js", include_in_schema=False)
async def campo_service_worker():
    response = FileResponse(STATIC_DIR / "sw.js", media_type="application/javascript")
    response.headers["Service-Worker-Allowed"] = "/"
    response.headers["Cache-Control"] = "no-cache"
    return response


@fastapi_app.post(TELEGRAM_WEBHOOK_PATH, include_in_schema=False)
async def receive_telegram_webhook(request: Request):
    """Acknowledge Telegram quickly and let Application process the update in its queue."""
    expected_secret = telegram_webhook_secret()
    supplied_secret = request.headers.get("X-Telegram-Bot-Api-Secret-Token", "")
    if not expected_secret or not secrets.compare_digest(supplied_secret, expected_secret):
        raise HTTPException(status_code=403, detail="Webhook de Telegram no autorizado")

    telegram_app = getattr(request.app.state, "telegram_app", None)
    if telegram_app is None:
        raise HTTPException(status_code=503, detail="Bot de Telegram iniciando")
    try:
        payload = await request.json()
        update = Update.de_json(payload, telegram_app.bot)
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Actualizacion de Telegram invalida") from exc
    if update is not None:
        await telegram_app.update_queue.put(update)
    return {"ok": True}

@fastapi_app.get("/api/health/campo")
async def health_campo():
    tables = {
        "field_items": check_supabase_table_health("field_items", FIELD_ITEMS_COLUMNS),
        "field_sessions": check_supabase_table_health("field_sessions", FIELD_SESSIONS_COLUMNS),
        "field_reports": check_supabase_table_health("field_reports", FIELD_REPORTS_COLUMNS),
    }
    return {
        "ok": True,
        "env": {
            "SUPABASE_URL": bool(SUPABASE_URL),
            "SUPABASE_SERVICE_ROLE_KEY": bool(SUPABASE_SERVICE_ROLE_KEY),
            "SUPABASE_BUCKET": bool(SUPABASE_BUCKET),
            "OPENAI_API_KEY": bool(OPENAI_API_KEY),
            "DATA_DIR": str(DATA_DIR),
            "ENABLE_TELEGRAM_BOT": os.environ.get("ENABLE_TELEGRAM_BOT", "false"),
        },
        "storage": check_supabase_storage_health(),
        "tables": tables,
        "capataz_tables": capataz_store.schema_health(),
        "push": push_notifier.status(),
        "gmail": gmail_service.status(),
        "geospatial": {"engine": "rasterio", **cdse_configuration_status(validate=False)},
        "telegram": {
            "enabled": os.environ.get("ENABLE_TELEGRAM_BOT", "false").lower()
            in {"1", "true", "yes", "si"},
            "mode": getattr(fastapi_app.state, "telegram_mode", "not_started"),
            "webhook_origin_configured": bool(telegram_webhook_base_url()),
        },
        "archive": archive_manager.status(),
    }


@fastapi_app.get("/api/health/cdse")
async def health_cdse():
    """Validate Copernicus OAuth without exposing any credential or token."""
    return cdse_configuration_status(validate=True)

@fastapi_app.post("/share-target")
async def share_target(
    title: str = Form(""),
    text: str = Form(""),
    url: str = Form(""),
):
    """Recibe texto compartido desde Android y lo entrega a la PWA para confirmar."""
    shared_text = "\n".join(part.strip() for part in (title, text, url) if part and part.strip())
    encoded = base64.b64encode(shared_text.encode("utf-8")).decode("ascii")
    return HTMLResponse(
        "<!doctype html><html lang='es'><meta charset='utf-8'>"
        "<meta name='viewport' content='width=device-width, initial-scale=1'>"
        "<title>Capataz Campo</title><body><p>Abriendo Capataz Campo...</p>"
        "<script>"
        f"const bytes=Uint8Array.from(atob('{encoded}'),c=>c.charCodeAt(0));"
        "localStorage.setItem('capataz.sharedText',new TextDecoder().decode(bytes));"
        "location.replace('/campo?shared=1');"
        "</script></body></html>"
    )

@fastapi_app.post("/api/capataz/analyze")
async def analyze_capataz_intake(payload: dict = Body(...)):
    text = str(payload.get("text") or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="text es obligatorio")
    draft = analyze_intake(
        text,
        field_name=str(payload.get("field_name") or payload.get("campo") or "").strip(),
        source=str(payload.get("source") or "app").strip(),
        openai_client=openai_client,
    )
    return {"ok": True, "draft": draft}

@fastapi_app.post("/api/capataz/confirm")
async def confirm_capataz_intake(background_tasks: BackgroundTasks, payload: dict = Body(...)):
    draft = payload.get("draft") or {}
    if not isinstance(draft, dict):
        raise HTTPException(status_code=400, detail="draft invalido")
    source_text = str(payload.get("source_text") or "")
    try:
        result = capataz_store.confirm_intake(draft, source_text=source_text)
    except PersistentStorageError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    try:
        plan = agent_crew.queue_event(result["event"], draft, source_text=source_text)
    except PersistentStorageError as exc:
        raise HTTPException(
            status_code=503,
            detail=(
                "La entrada quedo guardada, pero la cuadrilla no pudo iniciar. "
                "El borrador se conserva para reintentar."
            ),
        ) from exc
    queued = plan.get("storage") in {"supabase", "local"} and not (
        capataz_store.supabase_configured and plan.get("storage") != "supabase"
    )
    if capataz_store.supabase_configured and not queued:
        raise HTTPException(
            status_code=503,
            detail=(
                "La entrada quedo guardada, pero la cuadrilla no pudo iniciar. "
                "El borrador se conserva para reintentar."
            ),
        )
    if queued:
        background_tasks.add_task(
            process_capataz_event,
            result["event"],
            draft,
            source_text,
        )
    return {"ok": True, **result, "crew_plan": plan, "crew_queued": queued}

@fastapi_app.get("/api/capataz/dashboard")
async def get_capataz_dashboard():
    dashboard = capataz_store.dashboard()
    if capataz_store.supabase_configured and dashboard.get("warnings"):
        raise HTTPException(status_code=503, detail="No se pudo leer el seguimiento desde Supabase")
    try:
        reports = list_recent_field_reports_from_supabase(limit=30)
        today = argentina_now().date()

        def happened_today(row):
            value = row.get("finished_at") or row.get("updated_at") or row.get("created_at")
            if not value:
                return False
            try:
                parsed = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
                if parsed.tzinfo is None:
                    parsed = parsed.replace(tzinfo=argentina_now().tzinfo)
                return parsed.astimezone(argentina_now().tzinfo).date() == today
            except (TypeError, ValueError):
                return False

        dashboard["recent_reports"] = [report for report in reports if happened_today(report)][:10]
        dashboard["agent_activity"] = [
            run for run in dashboard.get("agent_activity") or [] if happened_today(run)
        ]
    except Exception as exc:
        logger.warning(f"No se pudieron agregar informes al tablero: {exc}")
        dashboard["recent_reports"] = []
    return {"ok": True, **dashboard}

@fastapi_app.get("/api/capataz/crew")
async def get_capataz_crew():
    return {"ok": True, "agents": agent_crew.registry()}

@fastapi_app.get("/api/capataz/daily-review")
async def get_capataz_daily_review():
    return {"ok": True, **agent_crew.daily_review()}

@fastapi_app.get("/api/capataz/push/public-key")
async def get_capataz_push_public_key():
    if not push_notifier.configured:
        raise HTTPException(status_code=503, detail="Notificaciones push no configuradas")
    return {"ok": True, "public_key": push_notifier.public_key}

@fastapi_app.post("/api/capataz/push/subscribe")
async def subscribe_capataz_push(payload: dict = Body(...)):
    try:
        subscription = push_notifier.subscribe(payload.get("subscription") or payload)
        return {"ok": True, "subscription_id": subscription["id"]}
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PersistentStorageError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

@fastapi_app.post("/api/capataz/reminders/dispatch")
async def dispatch_capataz_reminders(request: Request):
    try:
        daily_review = await asyncio.to_thread(agent_crew.persist_daily_review)
        followup_drafts = await asyncio.to_thread(email_draft_manager.prepare_due_followups)
        reminders = (
            await asyncio.to_thread(push_notifier.dispatch_due)
            if push_notifier.configured
            else {"sent": 0, "failed": 0, "skipped": "Web Push no configurado"}
        )
        email_sync = await asyncio.to_thread(email_draft_manager.sync_prepared)
        storage_cleanup = await asyncio.to_thread(run_storage_cleanup)
        telegram_summary_sent = False
        try:
            from capataz import format_pending_summary

            telegram_app = getattr(request.app.state, "telegram_app", None)
            dashboard = await asyncio.to_thread(capataz_store.dashboard)
            summary = format_pending_summary(dashboard)
            if summary and telegram_app is not None and MY_CHAT_ID:
                await telegram_app.bot.send_message(chat_id=MY_CHAT_ID, text=summary[:4000])
                telegram_summary_sent = True
        except Exception:
            logger.exception("No se pudo enviar el resumen diario por Telegram")
        return {
            "telegram_summary_sent": telegram_summary_sent,
            "ok": True,
            **reminders,
            "daily_review": daily_review,
            "followup_drafts": followup_drafts,
            "email_sync": email_sync,
            "storage_cleanup": storage_cleanup,
        }
    except (PersistentStorageError, RuntimeError) as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc


@fastapi_app.get("/api/archive/status")
async def get_archive_status():
    return {"ok": True, "archive": await asyncio.to_thread(archive_manager.status)}


@fastapi_app.get("/api/archive/manifest")
async def get_archive_manifest(limit: int = Query(100, ge=1, le=500)):
    try:
        objects = await asyncio.to_thread(archive_manager.manifest, limit=limit)
        return {"ok": True, "objects": objects}
    except (RuntimeError, ValueError) as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc


@fastapi_app.post("/api/archive/confirm")
async def confirm_local_archive(payload: dict = Body(...)):
    try:
        archived = await asyncio.to_thread(
            archive_manager.confirm,
            str(payload.get("archive_id") or ""),
            str(payload.get("sha256") or ""),
            payload.get("size_bytes"),
            str(payload.get("relative_path") or ""),
            machine=str(payload.get("machine") or ""),
        )
        return {"ok": True, "archive": archived}
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc


@fastapi_app.get("/api/capataz/gmail/status")
async def get_gmail_status():
    return {"ok": True, **gmail_service.status()}


@fastapi_app.post("/api/capataz/email-drafts/sync")
async def sync_gmail_drafts():
    try:
        result = await asyncio.to_thread(email_draft_manager.sync_prepared)
        return {"ok": True, **result}
    except (PersistentStorageError, RuntimeError) as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc


@fastapi_app.post("/api/capataz/events/{event_id}/email-draft")
async def prepare_event_email_draft(event_id: str):
    event_id = validate_record_id(event_id, "event_id")
    events, source, warning = capataz_store.list_rows("client_events", order="created_at.desc")
    event = next((row for row in events if row.get("id") == event_id), None)
    if not event:
        raise HTTPException(status_code=404, detail="evento no encontrado")
    if capataz_store.supabase_configured and source != "supabase":
        raise HTTPException(status_code=503, detail=warning or "evento no disponible")
    runs, runs_source, runs_warning = capataz_store.list_rows("agent_runs", order="created_at.desc")
    decisions, decisions_source, decisions_warning = capataz_store.list_rows("decisions", order="created_at.desc")
    if capataz_store.supabase_configured and (
        runs_source != "supabase" or decisions_source != "supabase"
    ):
        raise HTTPException(
            status_code=503,
            detail=runs_warning or decisions_warning or "trabajo de agentes no disponible",
        )
    draft = {
        "client_name": event.get("client_name") or "",
        "summary": event.get("summary") or "",
        "event_type": event.get("event_type") or "nota",
        "agents": event.get("agents") or ["Cartera"],
    }
    result = {
        "runs": [row for row in runs if row.get("event_id") == event_id],
        "decision": next((row for row in decisions if row.get("event_id") == event_id), None),
    }
    try:
        email_draft = email_draft_manager.prepare(
            event,
            draft,
            result,
            source_text=event.get("source_text") or "",
            force=True,
        )
        return {"ok": True, "email_draft": email_draft}
    except PersistentStorageError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

@fastapi_app.get("/api/capataz/events/{event_id}/runs")
async def get_capataz_event_runs(event_id: str):
    event_id = validate_record_id(event_id, "event_id")
    runs, source, warning = capataz_store.list_rows("agent_runs", order="created_at.desc")
    return {
        "ok": True,
        "runs": [run for run in runs if run.get("event_id") == event_id],
        "source": source,
        "warning": warning,
    }

@fastapi_app.post("/api/capataz/events/{event_id}/dispatch")
async def dispatch_capataz_event(event_id: str):
    event_id = validate_record_id(event_id, "event_id")
    events, source, warning = capataz_store.list_rows("client_events", order="created_at.desc")
    event = next((row for row in events if row.get("id") == event_id), None)
    if not event:
        raise HTTPException(status_code=404, detail="evento no encontrado")
    if capataz_store.supabase_configured and source != "supabase":
        raise HTTPException(status_code=503, detail=warning or "evento no disponible en Supabase")
    draft = {
        "draft_id": event_id.removeprefix("event-"),
        "client_name": event.get("client_name") or "",
        "summary": event.get("summary") or "",
        "event_type": event.get("event_type") or "nota",
        "agents": event.get("agents") or ["Cartera"],
        "economic_review": bool(event.get("economic_review")),
        "water_project": bool(event.get("water_project")),
        "field_name": event.get("field_name") or "",
        "source": event.get("source") or "app",
        "tasks": [],
    }
    try:
        result = process_capataz_event(event, draft, source_text=event.get("source_text") or "")
        return {"ok": True, **result}
    except PersistentStorageError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

@fastapi_app.post("/api/capataz/decisions/{decision_id}/approve")
async def approve_capataz_decision(decision_id: str):
    decision_id = validate_record_id(decision_id, "decision_id")
    try:
        result = agent_crew.approve_decision(decision_id)
        return {"ok": True, **result}
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except PersistentStorageError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

@fastapi_app.post("/api/capataz/decisions/{decision_id}/reject")
async def reject_capataz_decision(decision_id: str):
    decision_id = validate_record_id(decision_id, "decision_id")
    try:
        decision = capataz_store.update_row(
            "decisions",
            decision_id,
            {"status": "rejected", "updated_at": datetime.utcnow().isoformat() + "Z"},
        )
        return {"ok": True, "decision": decision}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc

@fastapi_app.patch("/api/capataz/tasks/{task_id}")
async def update_capataz_task(task_id: str, payload: dict = Body(...)):
    task_id = validate_record_id(task_id, "task_id")
    try:
        task = capataz_store.update_task(task_id, payload)
        return {"ok": True, "task": task}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc

@fastapi_app.patch("/api/capataz/clients/{client_id}")
async def update_capataz_client(client_id: str, payload: dict = Body(...)):
    client_id = validate_record_id(client_id, "client_id")
    try:
        client = capataz_store.update_client(client_id, payload)
        return {"ok": True, "client": client}
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PersistentStorageError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

@fastapi_app.post("/api/field-sessions")
async def create_field_session(payload: dict = Body(...)):
    session = normalize_field_session(payload)
    logger.info(f"Creando recorrida: {session['id']}")
    save_local_session(session)

    supabase_error = ""
    try:
        upsert_field_session_to_supabase(session)
    except Exception as e:
        supabase_error = str(e)
        logger.error(f"Recorrida ERROR: {e}")

    if supabase_database_configured() and supabase_error:
        raise HTTPException(status_code=503, detail=f"Supabase no guardo la recorrida: {supabase_error}")

    return {"ok": True, "id": session["id"], "supabase_error": supabase_error}

@fastapi_app.post("/api/field-sessions/{session_id}/close")
async def close_field_session(
    session_id: str,
    background_tasks: BackgroundTasks,
    payload: dict = Body(default={}),
):
    session_id = validate_record_id(session_id, "session_id")
    closed_at = clean_db_value(payload.get("closed_at")) or datetime.utcnow().isoformat() + "Z"
    logger.info(f"Cerrando recorrida: {session_id}")
    session = load_local_session(session_id) or {
        "id": session_id,
        "nombre": "",
        "campo": "",
        "sector": "",
        "started_at": None,
        "latitud_inicio": None,
        "longitud_inicio": None,
        "precision_gps_inicio": None,
        "notas": "",
        "created_at": datetime.utcnow().isoformat() + "Z",
    }
    session["estado"] = "cerrada"
    session["closed_at"] = closed_at
    session["updated_at"] = datetime.utcnow().isoformat() + "Z"
    save_local_session(session)

    supabase_error = ""
    try:
        close_field_session_in_supabase(session_id, closed_at)
    except Exception as e:
        supabase_error = str(e)
        logger.error(f"Recorrida ERROR: {e}")

    if supabase_database_configured() and supabase_error:
        raise HTTPException(status_code=503, detail=f"Supabase no cerro la recorrida: {supabase_error}")

    report_queued = False
    if supabase_database_configured():
        try:
            _remote_session, session_items = get_items_for_session_from_supabase(session_id)
            if session_items:
                background_tasks.add_task(generate_field_report, session_id, False)
                report_queued = True
        except Exception as exc:
            logger.warning(f"No se pudo encolar el informe automatico de {session_id}: {exc}")

    return {
        "ok": True,
        "id": session_id,
        "supabase_error": supabase_error,
        "report_queued": report_queued,
    }

@fastapi_app.get("/api/field-sessions")
async def list_field_sessions():
    supabase_error = ""
    try:
        supabase_sessions = list_field_sessions_from_supabase()
        if supabase_sessions is not None:
            return {"ok": True, "sessions": supabase_sessions, "source": "supabase"}
    except Exception as e:
        supabase_error = str(e)
        logger.error(f"Recorrida ERROR: {e}")

    if supabase_database_configured() and supabase_error:
        raise HTTPException(status_code=503, detail="No se pudieron leer las recorridas desde Supabase")

    sessions = load_local_sessions()
    item_counts = count_items_by_session(load_local_field_items())
    for session in sessions:
        session["items_count"] = item_counts.get(session.get("id"), 0)
        session["has_items"] = session["items_count"] > 0
    return {"ok": True, "sessions": sessions, "source": "local", "supabase_error": supabase_error}

@fastapi_app.post("/api/field-sessions/{session_id}/generate-report")
async def generate_field_session_report(
    session_id: str,
    background_tasks: BackgroundTasks,
    force: bool = Query(False),
):
    session_id = validate_record_id(session_id, "session_id")
    if not supabase_database_configured():
        raise HTTPException(status_code=503, detail="Supabase Database no configurado")
    try:
        session, _items = get_items_for_session_from_supabase(session_id)
    except Exception as exc:
        raise HTTPException(status_code=503, detail="No se pudo leer la recorrida desde Supabase") from exc
    if not session:
        raise HTTPException(status_code=404, detail="recorrida no encontrada")
    background_tasks.add_task(generate_field_report, session_id, force)
    return {
        "ok": True,
        "queued": True,
        "report": {
            "session_id": session_id,
            "estado": "generando",
            "progress_message": "Informe en cola",
        },
    }

@fastapi_app.get("/api/field-sessions/{session_id}/report")
async def get_field_session_report(session_id: str):
    session_id = validate_record_id(session_id, "session_id")
    session = None
    supabase_error = ""
    try:
        session = get_field_session_from_supabase(session_id)
        report = get_field_report_from_supabase(session_id)
        return {
            "ok": True,
            "session": session,
            "report": report,
            "docx_public_url": report.get("docx_public_url") if report else "",
            "pdf_public_url": report.get("pdf_public_url") if report else "",
            "informe_markdown": report.get("informe_markdown") if report else "",
            "estado": report.get("estado") if report else "sin informe",
            "progress_message": report.get("progress_message") if report else "",
            "error": report.get("error") if report else "",
            "source": "supabase" if supabase_database_configured() else "local",
        }
    except Exception as e:
        supabase_error = str(e)
        logger.error(f"Informe ERROR: {e}")

    session = session or load_local_session(session_id)
    return {
        "ok": True,
        "session": session,
        "report": None,
        "docx_public_url": "",
        "pdf_public_url": "",
        "informe_markdown": "",
        "estado": "sin informe",
        "progress_message": "",
        "error": supabase_error,
        "source": "local",
        "supabase_error": supabase_error,
    }

@fastapi_app.get("/api/field-sessions/{session_id}")
async def get_field_session(session_id: str):
    session_id = validate_record_id(session_id, "session_id")
    supabase_error = ""
    try:
        session, items = get_items_for_session_from_supabase(session_id)
        if session is not None:
            session["items_count"] = len(items)
            session["has_items"] = bool(items)
            return {"ok": True, "session": session, "items": items, "source": "supabase"}
    except Exception as e:
        supabase_error = str(e)
        logger.error(f"Recorrida ERROR: {e}")

    session = load_local_session(session_id)
    if not session:
        return {
            "ok": False,
            "error": supabase_error or "recorrida no encontrada",
            "session": None,
            "items": [],
            "source": "local",
            "supabase_error": supabase_error,
        }
    items = [item for item in load_local_field_items() if item.get("session_id") == session_id]
    session["items_count"] = len(items)
    session["has_items"] = bool(items)
    return {"ok": True, "session": session, "items": items, "source": "local", "supabase_error": supabase_error}

@fastapi_app.get("/api/field-items")
async def list_field_items():
    supabase_error = ""
    try:
        supabase_items = list_field_items_from_supabase()
        if supabase_items is not None:
            return {"ok": True, "items": supabase_items, "source": "supabase"}
    except Exception as e:
        supabase_error = str(e)
        logger.error(f"Metadata ERROR: {e}")

    if supabase_database_configured() and supabase_error:
        raise HTTPException(status_code=503, detail="No se pudieron leer los items desde Supabase")

    return {"ok": True, "items": load_local_field_items(), "source": "local", "supabase_error": supabase_error}

@fastapi_app.post("/api/field-items/{item_id}/assign-session")
async def assign_field_item_session(item_id: str, payload: dict = Body(...)):
    item_id = validate_record_id(item_id, "item_id")
    session_id = str(payload.get("session_id") or "").strip()
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id es obligatorio")
    session_id = validate_record_id(session_id, "session_id")

    linked_session = None
    if supabase_database_configured():
        try:
            linked_session = get_field_session_from_supabase(session_id)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"No se pudo validar la recorrida: {e}") from e
    else:
        linked_session = load_local_session(session_id)
    if not linked_session:
        raise HTTPException(status_code=404, detail="recorrida no encontrada")

    local_updated = False
    if FIELD_ITEMS_DIR.exists():
        for metadata_path in FIELD_ITEMS_DIR.rglob("*.json"):
            try:
                metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
            except Exception:
                continue
            if str(metadata.get("id") or "") != item_id:
                continue
            metadata["session_id"] = session_id
            metadata_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")
            local_updated = True
            break

    if supabase_database_configured():
        response = requests.patch(
            f"{SUPABASE_URL}/rest/v1/field_items",
            headers=supabase_headers("return=minimal"),
            params={"id": f"eq.{item_id}"},
            json={"session_id": session_id},
            timeout=30,
        )
        if not response.ok:
            raise HTTPException(status_code=500, detail=response.text)
        try:
            verify_field_item_session_assignment(item_id, session_id)
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e)) from e
    elif not local_updated:
        raise HTTPException(status_code=404, detail="item no encontrado")

    return {
        "ok": True,
        "item_id": item_id,
        "assigned_session_id": session_id,
        "session_name": linked_session.get("nombre") or linked_session.get("campo") or session_id,
    }

@fastapi_app.post("/api/field-items")
async def create_field_item(
    item_type: str = Form(...),
    campo: str = Form(...),
    sector: str = Form(""),
    captured_at: str = Form(...),
    latitude: str = Form(""),
    longitude: str = Form(""),
    gps_accuracy: str = Form(""),
    client_id: str = Form(""),
    session_id: str = Form(""),
    photo_label: str = Form(""),
    audio_label: str = Form(""),
    file: UploadFile = File(...),
):
    item_type = item_type.strip().lower()
    if item_type not in {"audio", "foto"}:
        raise HTTPException(status_code=400, detail="item_type debe ser audio o foto")
    if not campo.strip():
        raise HTTPException(status_code=400, detail="campo es obligatorio")
    session_id = validate_record_id(session_id.strip(), "session_id")
    if not session_id:
        raise HTTPException(
            status_code=400,
            detail="La foto o el audio deben pertenecer a una recorrida activa",
        )

    linked_session = None
    if supabase_database_configured():
        try:
            linked_session = get_field_session_from_supabase(session_id)
        except Exception as e:
            logger.warning(f"No se pudo validar la recorrida {session_id} en Supabase: {e}")
    else:
        linked_session = load_local_session(session_id)
    if linked_session is None:
        raise HTTPException(
            status_code=409,
            detail="La recorrida no existe en el servidor. Sincronizala antes de subir archivos.",
        )
    linked_campo = str(linked_session.get("campo") or "").strip()
    if linked_campo and normalize_key(campo) != normalize_key(linked_campo):
        logger.warning(
            f"Campo del item corregido por la recorrida: recibido={campo!r} recorrido={linked_campo!r}"
        )
        campo = linked_campo

    now = datetime.utcnow()
    try:
        captured_datetime = datetime.fromisoformat(captured_at.replace("Z", "+00:00"))
    except (TypeError, ValueError):
        captured_datetime = now
    requested_item_id = re.sub(r"[^A-Za-z0-9_-]", "", client_id.strip())[:80]
    item_id = requested_item_id or uuid.uuid4().hex
    item_dir = FIELD_ITEMS_DIR / captured_datetime.strftime("%Y-%m-%d")
    item_dir.mkdir(parents=True, exist_ok=True)

    extension = safe_field_extension(file, item_type)
    stored_filename = f"{item_id}_{item_type}{extension}"
    file_path = item_dir / stored_filename

    logger.info(f"Archivo recibido: {stored_filename}")
    max_upload_bytes = validate_field_upload(file, item_type)
    uploaded_bytes = save_upload_with_limit(file, file_path, max_upload_bytes)

    metadata = {
        "ok": True,
        "id": item_id,
        "client_id": client_id,
        "item_type": item_type,
        "campo": campo.strip(),
        "sector": sector.strip(),
        "captured_at": captured_at,
        "latitude": latitude,
        "longitude": longitude,
        "gps_accuracy": gps_accuracy,
        "original_filename": file.filename,
        "content_type": file.content_type,
        "size_bytes": uploaded_bytes,
        "stored_file": str(file_path),
        "received_at": now.isoformat() + "Z",
        "ai_processed": False,
        "storage_status": "local_only",
        "storage_provider": "local",
        "session_id": session_id,
        "photo_label": photo_label.strip(),
        "audio_label": audio_label.strip(),
    }
    logger.info(f"Item asociado a session_id: {session_id}")

    if SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY and SUPABASE_BUCKET:
        campo_segment = safe_storage_segment(campo, "sin-campo")
        sector_segment = safe_storage_segment(sector, "sin-sector")
        storage_path = f"{campo_segment}/{sector_segment}/{captured_datetime.strftime('%Y-%m-%d')}/{stored_filename}"
        try:
            supabase_file = upload_field_file_to_supabase(
                file_path,
                storage_path,
                content_type=file.content_type,
                upsert=True,
            )
            metadata["storage_status"] = "supabase_uploaded"
            metadata["storage_provider"] = "supabase"
            metadata["storage_path"] = supabase_file.get("path", "")
            metadata["storage_public_url"] = supabase_file.get("public_url", "")
            logger.info(f"Supabase OK: {metadata['storage_path']}")
        except Exception as e:
            metadata["storage_status"] = "supabase_error"
            metadata["storage_provider"] = "supabase"
            metadata["storage_path"] = storage_path
            metadata["storage_error"] = str(e)
            logger.error(f"Supabase ERROR: {e}")

    capataz_draft = None
    capataz_error = ""
    source_text = audio_label.strip() if item_type == "audio" else photo_label.strip()
    if item_type == "audio" and openai_client:
        try:
            source_text = transcribe_audio(file_path)
            metadata.update({
                "ai_processed": True,
                "transcript_status": "done",
                "transcript_text": source_text,
                "transcript_model": "whisper-1",
                "transcript_at": datetime.utcnow().isoformat() + "Z",
                "transcript_error": "",
            })
        except Exception as e:
            capataz_error = f"No se pudo transcribir el audio: {e}"
            metadata.update({
                "transcript_status": "error",
                "transcript_error": str(e),
                "transcript_at": datetime.utcnow().isoformat() + "Z",
            })
            logger.error(f"Capataz audio ERROR: {e}")

    if source_text:
        try:
            capataz_draft = analyze_intake(
                source_text,
                field_name=campo.strip(),
                source=f"{item_type}_campo",
                openai_client=openai_client,
            )
        except Exception as e:
            capataz_error = str(e)
            logger.error(f"Capataz analisis ERROR: {e}")

    metadata_path = file_path.with_suffix(file_path.suffix + ".json")
    metadata_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")

    metadata_error = ""
    try:
        upsert_field_item_metadata(metadata)
        verify_field_item_session_assignment(item_id, session_id)
    except Exception as e:
        metadata_error = str(e)
        logger.error(f"Metadata ERROR: {e}")

    return {
        "ok": not metadata_error and not metadata.get("storage_error"),
        "id": item_id,
        "storage_status": metadata.get("storage_status", ""),
        "storage_error": metadata.get("storage_error", ""),
        "metadata_error": metadata_error,
        "assigned_session_id": session_id,
        "transcript_text": metadata.get("transcript_text", ""),
        "capataz_draft": capataz_draft,
        "capataz_error": capataz_error,
    }

if __name__ == "__main__":
    main()
