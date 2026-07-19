"""Lectura deterministica de documentos compartidos con Capataz Campo."""

from __future__ import annotations

import csv
import mimetypes
import re
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None


LUCAS_HIGHLIGHT_START = "[APORTE DE LUCAS - RESALTADO AMARILLO]"
LUCAS_HIGHLIGHT_END = "[/APORTE DE LUCAS]"


def _docx_paragraph_text(paragraph):
    """Preserve yellow authorship instead of flattening every DOCX run."""
    chunks = []
    yellow_open = False
    has_drawing = False
    for run in paragraph.runs:
        highlighted = run.font.highlight_color == WD_COLOR_INDEX.YELLOW
        if highlighted and not yellow_open:
            chunks.append(LUCAS_HIGHLIGHT_START)
            yellow_open = True
        elif yellow_open and not highlighted:
            chunks.append(LUCAS_HIGHLIGHT_END)
            yellow_open = False
        if run.text:
            chunks.append(run.text)
        if run._element.xpath(".//a:blip"):
            has_drawing = True
    if yellow_open:
        chunks.append(LUCAS_HIGHLIGHT_END)
    text = "".join(chunks).strip()
    if has_drawing:
        text = "\n".join(value for value in (text, "[IMAGEN INCORPORADA EN EL DOCUMENTO]") if value)
    return text


def extract_docx_embedded_images(file_path, max_images=6, max_bytes=5_000_000):
    """Return bounded DOCX images for visual grounding and final-report figures."""
    result = []
    try:
        with ZipFile(file_path) as archive:
            names = sorted(
                name for name in archive.namelist()
                if name.startswith("word/media/") and not name.endswith("/")
            )
            for name in names:
                if len(result) >= max_images:
                    break
                mime_type = mimetypes.guess_type(name)[0] or "application/octet-stream"
                if not mime_type.startswith("image/"):
                    continue
                info = archive.getinfo(name)
                if info.file_size <= 0 or info.file_size > max_bytes:
                    continue
                result.append({
                    "name": Path(name).name,
                    "mime_type": mime_type,
                    "data": archive.read(name),
                })
    except (BadZipFile, KeyError, OSError):
        return []
    return result


def extract_office_document(file_path, file_name=""):
    """Extract auditable text from DOCX/XLSX/CSV/TXT without model guesses."""
    suffix = Path(file_name or file_path).suffix.lower()
    if suffix == ".docx":
        document = Document(file_path)
        blocks = []
        for paragraph in document.paragraphs:
            text = _docx_paragraph_text(paragraph)
            if text:
                blocks.append(text)
        for table_index, table in enumerate(document.tables, 1):
            blocks.append(f"--- Tabla {table_index} ---")
            for row in table.rows:
                values = []
                for cell in row.cells:
                    paragraph_values = []
                    for paragraph in cell.paragraphs:
                        paragraph_text = _docx_paragraph_text(paragraph)
                        if paragraph_text:
                            paragraph_values.append(paragraph_text)
                    cell_text = " ".join(paragraph_values)
                    values.append(re.sub(r"\s+", " ", cell_text).strip())
                if any(values):
                    blocks.append(" | ".join(values))
        image_count = len(extract_docx_embedded_images(file_path))
        if image_count:
            blocks.append(f"[IMAGENES INCORPORADAS DETECTADAS: {image_count}]")
        return "\n".join(blocks)[:100000]
    if suffix in {".xlsx", ".xlsm"}:
        if load_workbook is None:
            raise RuntimeError("Falta openpyxl para leer la planilla")
        workbook = load_workbook(file_path, read_only=True, data_only=False)
        blocks = []
        cell_count = 0
        try:
            for sheet in workbook.worksheets[:12]:
                blocks.append(f"--- Hoja: {sheet.title} ---")
                for row in sheet.iter_rows():
                    values = [str(cell.value).strip() if cell.value is not None else "" for cell in row[:60]]
                    while values and not values[-1]:
                        values.pop()
                    if values and any(values):
                        blocks.append(" | ".join(values))
                        cell_count += len(values)
                    if cell_count >= 12000:
                        blocks.append("[Lectura truncada al limite de 12.000 celdas]")
                        return "\n".join(blocks)[:100000]
        finally:
            workbook.close()
        return "\n".join(blocks)[:100000]
    if suffix in {".csv", ".tsv"}:
        raw = Path(file_path).read_bytes()
        decoded = ""
        for encoding in ("utf-8-sig", "latin-1"):
            try:
                decoded = raw.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        dialect = csv.excel_tab if suffix == ".tsv" else csv.excel
        rows = []
        for index, row in enumerate(csv.reader(decoded.splitlines(), dialect=dialect)):
            rows.append(" | ".join(str(value).strip() for value in row[:60]))
            if index >= 1000:
                rows.append("[Lectura truncada a 1.000 filas]")
                break
        return "\n".join(rows)[:100000]
    if suffix in {".txt", ".md"}:
        raw = Path(file_path).read_bytes()
        try:
            return raw.decode("utf-8-sig")[:100000]
        except UnicodeDecodeError:
            return raw.decode("latin-1")[:100000]
    raise ValueError("Formato de oficina no soportado")
