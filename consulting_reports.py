"""Generador de PDF/DOCX para trabajos profesionales de la cuadrilla.

La redaccion final requiere un modelo configurado y falla de manera visible si
ese modelo no responde. Los tests pueden habilitar expresamente el borrador
determinista, pero produccion nunca lo presenta como un informe terminado.
"""

from __future__ import annotations

import base64
import json
import os
import re
import unicodedata
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import reportlab
from capataz import extract_json_object
from document_intake import extract_docx_embedded_images
from report_playbooks import ReportPlaybook, detect_report_playbook, playbook_prompt

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


GREEN = "0B5D4D"
GREEN_LIGHT = "EAF5F0"
GOLD = "9A6A2F"
INK = "17211F"
MUTED = "5E6B68"
LIGHT = "F2F5F3"
RED = "8F2D2D"
PAGE_WIDTH_DXA = 9360


@dataclass(frozen=True)
class GeneratedConsultingReport:
    playbook_key: str
    title: str
    status: str
    pdf_path: str
    docx_path: str
    missing_data: tuple[str, ...]
    model: str


def _clean_text(value: Any, limit: int = 8000) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    return text[:limit]


def _clean_list(values: Any, *, limit: int = 20, item_limit: int = 800) -> list[str]:
    if isinstance(values, str):
        values = [values]
    if not isinstance(values, (list, tuple)):
        return []
    result = []
    for value in values:
        cleaned = _clean_text(value, item_limit)
        if cleaned and cleaned not in result:
            result.append(cleaned)
    return result[:limit]


def _safe_filename(value: str, fallback: str = "Informe") -> str:
    text = "".join(
        char for char in unicodedata.normalize("NFKD", str(value or ""))
        if not unicodedata.combining(char)
    )
    text = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
    return (text or fallback)[:90]


def _materialize_docx_figures(assets: list[dict], root: Path) -> list[dict]:
    figures = []
    figure_dir = root / "figuras_fuente"
    for asset in assets:
        asset_path = Path(str(asset.get("path") or ""))
        if asset_path.suffix.lower() != ".docx" or not asset_path.exists():
            continue
        for embedded in extract_docx_embedded_images(asset_path):
            figure_dir.mkdir(parents=True, exist_ok=True)
            suffix = Path(embedded["name"]).suffix or ".png"
            output_path = figure_dir / f"figura_{len(figures) + 1:02d}{suffix.lower()}"
            output_path.write_bytes(embedded["data"])
            figures.append({
                "path": str(output_path),
                "caption": (
                    f"Figura {len(figures) + 1}. Imagen incorporada en "
                    f"{asset.get('file_name') or asset_path.name}."
                ),
            })
    return figures


def _run_outputs(crew_result: dict) -> list[dict]:
    return [
        {"agent": run.get("agent"), "output": run.get("output") or {}}
        for run in (crew_result or {}).get("runs") or []
        if run.get("agent") != "Contralor"
    ]


def _fallback_payload(
    playbook: ReportPlaybook,
    event: dict,
    draft: dict,
    crew_result: dict,
    source_text: str,
    assets: list[dict],
) -> dict:
    runs = _run_outputs(crew_result)
    decision = (crew_result or {}).get("decision") or {}
    findings = []
    recommendations = []
    risks = list(decision.get("risks") or [])
    missing = list(decision.get("missing_data") or [])
    for run in runs:
        output = run["output"]
        summary = _clean_text(output.get("summary"), 1000)
        if summary:
            findings.append(f"{run['agent']}: {summary}")
        findings.extend(_clean_list(output.get("findings"), limit=8))
        recommendations.extend(_clean_list(output.get("recommendations"), limit=8))
        risks.extend(_clean_list(output.get("risks"), limit=8))
        missing.extend(_clean_list(output.get("missing_data"), limit=8))

    summary = _clean_text(
        decision.get("summary") or draft.get("summary") or source_text,
        2400,
    )
    recommendation = _clean_text(
        decision.get("recommendation")
        or (recommendations[0] if recommendations else "Revisar los datos pendientes antes de ejecutar."),
        1600,
    )
    field = _clean_text(draft.get("field_name") or event.get("client_name"), 120)
    sections = []
    evidence = findings or ["Entrada recibida sin hallazgos tecnicos estructurados."]
    for index, title in enumerate(playbook.sections):
        if index == 0:
            paragraphs = [summary or "No se registro un resumen suficiente."]
            bullets = []
        elif index == 1:
            paragraphs = [
                "Esta seccion se construye exclusivamente con la nota y los archivos recibidos. "
                "Los datos no demostrados quedan consignados como pendientes."
            ]
            bullets = evidence[:10]
        elif "riesg" in title.lower() or "pendiente" in title.lower() or "limit" in title.lower():
            paragraphs = []
            bullets = _clean_list(risks + missing, limit=15) or ["Sin datos suficientes para cerrar esta seccion."]
        elif "recomend" in title.lower() or "decision" in title.lower():
            paragraphs = [recommendation]
            bullets = recommendations[:8]
        else:
            paragraphs = []
            bullets = []
        sections.append({"title": title, "paragraphs": paragraphs, "bullets": bullets})

    return {
        "title": f"{playbook.title}{(' - ' + field) if field else ''}",
        "subtitle": playbook.purpose,
        "executive_summary": summary or "Trabajo recibido; faltan datos para emitir una conclusion tecnica.",
        "recommendation": recommendation,
        "status": "preliminar" if missing else "auditado",
        "sections": sections,
        "calculations": [],
        "risks": _clean_list(risks, limit=20),
        "missing_data": _clean_list(missing, limit=20),
        "sources": [
            f"Nota original recibida por {event.get('source') or draft.get('source') or 'Capataz'}",
            *[f"Archivo recibido: {asset.get('file_name')}" for asset in assets if asset.get("file_name")],
            *[f"Revision del agente {run['agent']}" for run in runs],
        ],
    }


def _normalize_calculation(value: Any) -> dict | None:
    if not isinstance(value, dict):
        return None
    status = _clean_text(value.get("status"), 40).lower()
    if status not in {"calculado", "pendiente", "no_aplica"}:
        status = "pendiente"
    label = _clean_text(value.get("label"), 180)
    if not label:
        return None
    formula = _clean_text(value.get("formula"), 500)
    result = _clean_text(value.get("result"), 300)
    source = _clean_text(value.get("source"), 500)
    if status == "calculado" and (not formula or not result or not source):
        status = "pendiente"
        result = ""
    return {
        "label": label,
        "formula": formula or "Pendiente de datos",
        "result": result or "Pendiente",
        "source": source or "Fuente no identificada",
        "status": status,
    }


def _normalize_payload(data: Any, fallback: dict, playbook: ReportPlaybook) -> dict:
    if not isinstance(data, dict):
        return fallback
    sections = []
    raw_sections = data.get("sections") or []
    if isinstance(raw_sections, list):
        for section in raw_sections[:20]:
            if not isinstance(section, dict):
                continue
            title = _clean_text(section.get("title"), 180)
            if not title:
                continue
            sections.append({
                "title": title,
                "paragraphs": _clean_list(section.get("paragraphs"), limit=8, item_limit=2200),
                "bullets": _clean_list(section.get("bullets"), limit=18, item_limit=1200),
            })
    section_by_key = {section["title"].lower(): section for section in sections}
    ordered_sections = []
    for required in playbook.sections:
        matched = next(
            (section for key, section in section_by_key.items() if required.lower() == key),
            None,
        )
        fallback_section = next(
            (section for section in fallback["sections"] if section["title"].lower() == required.lower()),
            {"title": required, "paragraphs": [], "bullets": []},
        )
        ordered_sections.append(matched or fallback_section)
    calculations = []
    for raw in data.get("calculations") or []:
        normalized = _normalize_calculation(raw)
        if normalized:
            calculations.append(normalized)
    missing = _clean_list(data.get("missing_data"), limit=25) or fallback["missing_data"]
    status = _clean_text(data.get("status"), 40).lower()
    if status not in {"preliminar", "auditado", "bloqueado"}:
        status = "preliminar" if missing else "auditado"
    if any(item["status"] == "pendiente" for item in calculations) and status == "auditado":
        status = "preliminar"
    return {
        "title": _clean_text(data.get("title"), 220) or fallback["title"],
        "subtitle": _clean_text(data.get("subtitle"), 500) or fallback["subtitle"],
        "executive_summary": _clean_text(data.get("executive_summary"), 3000) or fallback["executive_summary"],
        "recommendation": _clean_text(data.get("recommendation"), 2200) or fallback["recommendation"],
        "status": status,
        "sections": ordered_sections,
        "calculations": calculations[:25],
        "risks": _clean_list(data.get("risks"), limit=25) or fallback["risks"],
        "missing_data": missing,
        "sources": _clean_list(data.get("sources"), limit=30) or fallback["sources"],
    }


def _generate_payload(
    playbook: ReportPlaybook,
    event: dict,
    draft: dict,
    crew_result: dict,
    source_text: str,
    assets: list[dict],
    openai_client=None,
    model: str = "gpt-5.6-sol",
    reasoning_effort: str = "high",
    allow_fallback: bool = False,
) -> dict:
    fallback = _fallback_payload(playbook, event, draft, crew_result, source_text, assets)
    if openai_client is None:
        if allow_fallback:
            return fallback
        raise RuntimeError(
            "No se genero el informe: falta un cliente OpenAI activo. "
            "Verifica OPENAI_API_KEY y CAPATAZ_REPORT_MODEL."
        )
    evidence = {
        "cliente": event.get("client_name"),
        "campo": draft.get("field_name"),
        # Un comparativo puede contener varias ofertas extensas. 180k caracteres
        # permite incluir el conjunto sin acercarse al contexto maximo del modelo.
        "nota_original": str(source_text or "")[:180000],
        "archivos": [
            {"nombre": asset.get("file_name"), "tipo": asset.get("asset_type")}
            for asset in assets
        ],
        "especialistas": _run_outputs(crew_result),
    }
    prompt = f"""
Sos el agente Informes del estudio del Ing. Agr. Lucas Estecho (M.P. 2009 LE),
Entre Rios, Argentina. Debes preparar el contenido de un entregable profesional.

{playbook_prompt(playbook)}

REGLAS DE REDACCION Y VERIFICACION:
- Usa solo la evidencia incluida abajo. No completes huecos con conocimiento general.
- Conserva todo el contenido material de las notas; mejora la redaccion y la estructura sin sustituirlo por relleno generico.
- Los segmentos entre [APORTE DE LUCAS - RESALTADO AMARILLO] y [/APORTE DE LUCAS] fueron agregados por Lucas: identificalos explicitamente como aportes de Lucas. El resto proviene del registro original.
- Las imagenes adjuntas son evidencia del documento fuente. Interpreta solo lo que sea legible y no inventes valores ocultos.
- Todo numero debe conservar unidad y fuente. No inventes precios, superficies, fechas ni mediciones.
- Un calculo con datos insuficientes queda status "pendiente", sin resultado supuesto.
- Distingui claramente medido/provisto, calculado, inferido y pendiente.
- La recomendacion debe responder a la decision, ser condicionada y tener proximo paso.
- Tono profesional, claro, argentino, sin mencionar IA ni estos prompts.
- No autorices compra, obra, envio o compromiso externo.
- El documento es un ENTREGABLE PARA EL CLIENTE. PROHIBIDO mencionar Contralor,
  auditorias, estados internos (pending_review, bloqueado, borrador), IDs de
  decisiones, agentes internos o el funcionamiento del sistema.
- Lo que falta va SOLO en "missing_data" (uso interno de Lucas). Las secciones
  del documento desarrollan lo que SI esta registrado, bien redactado; no
  repitas listas de faltantes ni llenes secciones con "pendiente".
- El subtitulo es profesional (cliente, campo o fecha); nunca "borrador" ni
  "sujeto a validacion".

Responde SOLO JSON puro con esta forma:
{{
  "title": "titulo",
  "subtitle": "subtitulo",
  "executive_summary": "resumen con conclusion y limites",
  "recommendation": "recomendacion condicionada",
  "status": "preliminar|auditado|bloqueado",
  "sections": [
    {{"title":"titulo exacto del contrato","paragraphs":["texto"],"bullets":["punto"]}}
  ],
  "calculations": [
    {{"label":"nombre","formula":"formula con operandos","result":"resultado y unidad",
      "source":"origen exacto de los operandos","status":"calculado|pendiente|no_aplica"}}
  ],
  "risks": ["riesgo"],
  "missing_data": ["dato concreto faltante"],
  "sources": ["fuente efectivamente recibida"]
}}

EVIDENCIA:
{json.dumps(evidence, ensure_ascii=False)}
""".strip()
    content: str | list[dict] = prompt
    image_parts = []
    for asset in assets:
        asset_path = Path(str(asset.get("path") or ""))
        if asset_path.suffix.lower() != ".docx" or not asset_path.exists():
            continue
        for embedded in extract_docx_embedded_images(asset_path):
            encoded = base64.b64encode(embedded["data"]).decode("ascii")
            image_parts.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:{embedded['mime_type']};base64,{encoded}",
                    "detail": "high",
                },
            })
    if image_parts:
        content = [{"type": "text", "text": prompt}, *image_parts]
    try:
        request = {"model": model, "messages": [{"role": "user", "content": content}]}
        if str(model).startswith("gpt-5.6"):
            request["reasoning_effort"] = reasoning_effort
        else:
            request["temperature"] = 0
        # Tope duro: sin esto, una llamada colgada deja al bot "clavado" sin aviso.
        request["timeout"] = float(os.environ.get("CAPATAZ_REPORT_TIMEOUT_SECONDS", "420"))
        response = openai_client.chat.completions.create(**request)
        return _normalize_payload(
            extract_json_object(response.choices[0].message.content),
            fallback,
            playbook,
        )
    except Exception as exc:
        if allow_fallback:
            return fallback
        raise RuntimeError(
            f"No se genero el informe: el modelo {model} fallo: {str(exc)[:500]}"
        ) from exc


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[min(index, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_docx_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, before, after, color in (
        ("Title", 26, 0, 8, GREEN),
        ("Subtitle", 13, 0, 12, MUTED),
        ("Heading 1", 16, 16, 8, GREEN),
        ("Heading 2", 13, 12, 6, GREEN),
        ("Heading 3", 12, 8, 4, GOLD),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        # Algunas plantillas de Word traen una linea azul heredada en Title.
        # La identidad de Capataz usa jerarquia tipografica, no ese residuo.
        if name == "Title":
            p_pr = style.element.get_or_add_pPr()
            border = p_pr.find(qn("w:pBdr"))
            if border is not None:
                p_pr.remove(border)
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def _docx_header_footer(document: Document, logo_path: str | None) -> None:
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_path and Path(logo_path).exists():
        paragraph.add_run().add_picture(str(logo_path), width=Inches(0.36))
    run = paragraph.add_run("  LUCAS ESTECHO - INGENIERO AGRONOMO")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(GREEN)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Pasto · Agua · Ganaderia rentable  |  M.P. 2009 LE")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def _add_docx_callout(document: Document, label: str, value: str, fill: str = GREEN_LIGHT) -> None:
    table = document.add_table(rows=1, cols=1)
    _set_table_geometry(table, [PAGE_WIDTH_DXA])
    cell = table.cell(0, 0)
    _shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    label_run = paragraph.add_run(label.upper() + "\n")
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor.from_string(GREEN)
    value_run = paragraph.add_run(value)
    value_run.font.size = Pt(11)


def create_consulting_docx(payload: dict, output_path: str, *, event: dict, playbook: ReportPlaybook,
                           logo_path: str | None = None) -> str:
    document = Document()
    _set_docx_styles(document)
    _docx_header_footer(document, logo_path)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(34)
    run = kicker.add_run("CAPATAZ CAMPO · ENTREGABLE PROFESIONAL")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(GOLD)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(payload["title"])
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(payload["subtitle"])
    meta = document.add_table(rows=0, cols=2)
    meta.style = "Table Grid"
    rows = [
        ("Cliente / campo", event.get("client_name") or "No identificado"),
        ("Fecha", datetime.now().strftime("%d/%m/%Y")),
        ("Tipo", playbook.title),
    ]
    for label, value in rows:
        cells = meta.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)
        _shade_cell(cells[0], LIGHT)
        cells[0].paragraphs[0].runs[0].bold = True
    _set_table_geometry(meta, [2200, 7160])
    document.add_paragraph()
    _add_docx_callout(document, "Conclusion ejecutiva", payload["executive_summary"])
    document.add_paragraph()
    _add_docx_callout(document, "Recomendacion", payload["recommendation"], fill="F7F0E7")
    document.add_page_break()

    for section in payload["sections"]:
        if not section["paragraphs"] and not section["bullets"]:
            continue
        document.add_heading(section["title"], level=1)
        for paragraph_text in section["paragraphs"]:
            document.add_paragraph(paragraph_text)
        for bullet in section["bullets"]:
            document.add_paragraph(bullet, style="List Bullet")

    if payload.get("figures"):
        document.add_heading("Figuras del documento fuente", level=1)
        for figure in payload["figures"]:
            figure_path = Path(figure["path"])
            if not figure_path.exists():
                continue
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(figure_path), width=Inches(6.0))
            caption = document.add_paragraph(figure["caption"])
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption.runs:
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor.from_string(MUTED)

    payload = {**payload, "calculations": [
        calculation for calculation in payload.get("calculations") or []
        if str(calculation.get("status") or "").lower() == "calculado"
    ]}
    if payload["calculations"]:
        document.add_heading("Calculos realizados", level=1)
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["Calculo", "Formula", "Resultado", "Fuente / estado"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
            _shade_cell(table.rows[0].cells[index], GREEN)
            for run in table.rows[0].cells[index].paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        for calculation in payload["calculations"]:
            cells = table.add_row().cells
            cells[0].text = calculation["label"]
            cells[1].text = calculation["formula"]
            cells[2].text = calculation["result"]
            cells[3].text = str(calculation['source'])
        _set_table_geometry(table, [1900, 2700, 1700, 3060])

    if payload["sources"]:
        document.add_heading("Fuentes", level=1)
        for value in payload["sources"]:
            document.add_paragraph(value, style="List Bullet")
    document.add_paragraph()
    closing = document.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = closing.add_run("Ing. Agr. Lucas Estecho\nM.P. 2009 LE")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def _register_pdf_fonts() -> tuple[str, str]:
    regular = "CapatazSans"
    bold = "CapatazSans-Bold"
    if regular not in pdfmetrics.getRegisteredFontNames():
        fonts = Path(reportlab.__file__).parent / "fonts"
        pdfmetrics.registerFont(TTFont(regular, str(fonts / "Vera.ttf")))
        pdfmetrics.registerFont(TTFont(bold, str(fonts / "VeraBd.ttf")))
        pdfmetrics.registerFont(TTFont("CapatazSans-Italic", str(fonts / "VeraIt.ttf")))
        pdfmetrics.registerFont(TTFont("CapatazSans-BoldItalic", str(fonts / "VeraBI.ttf")))
        pdfmetrics.registerFontFamily(
            regular,
            normal=regular,
            bold=bold,
            italic="CapatazSans-Italic",
            boldItalic="CapatazSans-BoldItalic",
        )
    return regular, bold


def _pdf_styles():
    regular, bold = _register_pdf_fonts()
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "CapatazTitle", parent=styles["Title"], fontName=bold,
            fontSize=25, leading=29, textColor=colors.HexColor("#" + GREEN),
            alignment=TA_CENTER, spaceAfter=10,
        ),
        "subtitle": ParagraphStyle(
            "CapatazSubtitle", parent=styles["Normal"], fontName=regular,
            fontSize=12, leading=16, textColor=colors.HexColor("#" + MUTED),
            alignment=TA_CENTER, spaceAfter=18,
        ),
        "h1": ParagraphStyle(
            "CapatazH1", parent=styles["Heading1"], fontName=bold,
            fontSize=15, leading=18, textColor=colors.HexColor("#" + GREEN),
            spaceBefore=12, spaceAfter=8, keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "CapatazBody", parent=styles["BodyText"], fontName=regular,
            fontSize=9.5, leading=13.2, textColor=colors.HexColor("#" + INK),
            spaceAfter=6,
        ),
        "bullet": ParagraphStyle(
            "CapatazBullet", parent=styles["BodyText"], fontName=regular,
            fontSize=9.3, leading=12.5, leftIndent=13, firstLineIndent=-7,
            bulletIndent=5, spaceAfter=4, textColor=colors.HexColor("#" + INK),
        ),
        "small": ParagraphStyle(
            "CapatazSmall", parent=styles["BodyText"], fontName=regular,
            fontSize=7.5, leading=9.5, textColor=colors.HexColor("#" + MUTED),
        ),
        "callout_label": ParagraphStyle(
            "CapatazCalloutLabel", parent=styles["BodyText"], fontName=bold,
            fontSize=8, leading=10, textColor=colors.HexColor("#" + GREEN),
            spaceAfter=3,
        ),
    }


def _pdf_header_footer(canvas, doc, *, logo_path: str | None, report_title: str) -> None:
    regular, bold = _register_pdf_fonts()
    canvas.saveState()
    width, height = A4
    if logo_path and Path(logo_path).exists():
        canvas.drawImage(str(logo_path), 1.45 * cm, height - 1.35 * cm, width=0.65 * cm,
                         height=0.65 * cm, preserveAspectRatio=True, mask="auto")
    canvas.setFont(bold, 8)
    canvas.setFillColor(colors.HexColor("#" + GREEN))
    canvas.drawString(2.25 * cm, height - 1.05 * cm, "LUCAS ESTECHO · INGENIERO AGRONOMO")
    canvas.setStrokeColor(colors.HexColor("#C7D8D1"))
    canvas.line(1.45 * cm, height - 1.48 * cm, width - 1.45 * cm, height - 1.48 * cm)
    canvas.setFont(regular, 7.5)
    canvas.setFillColor(colors.HexColor("#" + MUTED))
    canvas.drawString(1.45 * cm, 0.72 * cm, "Pasto · Agua · Ganaderia rentable · M.P. 2009 LE")
    canvas.drawRightString(width - 1.45 * cm, 0.72 * cm, f"Pagina {doc.page}")
    canvas.restoreState()


def _pdf_callout(label: str, value: str, styles: dict, *, fill: str = GREEN_LIGHT):
    table = Table([[Paragraph(label.upper(), styles["callout_label"]), Paragraph(value, styles["body"])]],
                  colWidths=[3.5 * cm, 12.5 * cm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#" + fill)),
        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#B8CEC5")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 9),
        ("RIGHTPADDING", (0, 0), (-1, -1), 9),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    return table


def create_consulting_pdf(payload: dict, output_path: str, *, event: dict, playbook: ReportPlaybook,
                          logo_path: str | None = None) -> str:
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    styles = _pdf_styles()
    regular, bold = _register_pdf_fonts()
    doc = SimpleDocTemplate(
        output_path, pagesize=A4, rightMargin=1.55 * cm, leftMargin=1.55 * cm,
        topMargin=1.8 * cm, bottomMargin=1.35 * cm,
        title=payload["title"], author="Ing. Agr. Lucas Estecho",
    )
    story = [Spacer(1, 2.1 * cm)]
    if logo_path and Path(logo_path).exists():
        story.extend([Image(str(logo_path), width=3.2 * cm, height=3.2 * cm), Spacer(1, 0.45 * cm)])
        story[-2].hAlign = "CENTER"
    story.extend([
        Paragraph("CAPATAZ CAMPO · ENTREGABLE PROFESIONAL", styles["callout_label"]),
        Paragraph(payload["title"], styles["title"]),
        Paragraph(payload["subtitle"], styles["subtitle"]),
    ])
    meta_data = [
        ["Cliente / campo", _clean_text(event.get("client_name") or "No identificado", 200)],
        ["Fecha", datetime.now().strftime("%d/%m/%Y")],
        ["Tipo", playbook.title],
    ]
    meta = Table(meta_data, colWidths=[4 * cm, 12 * cm])
    meta.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#D2DDD8")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#" + LIGHT)),
        ("FONTNAME", (0, 0), (0, -1), bold),
        ("FONTNAME", (1, 0), (1, -1), regular),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#" + INK)),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.extend([
        meta, Spacer(1, 0.55 * cm),
        _pdf_callout("Conclusion", payload["executive_summary"], styles),
        Spacer(1, 0.35 * cm),
        _pdf_callout("Recomendacion", payload["recommendation"], styles, fill="F7F0E7"),
        PageBreak(),
    ])
    for section in payload["sections"]:
        if not section["paragraphs"] and not section["bullets"]:
            continue
        story.append(Paragraph(section["title"], styles["h1"]))
        for value in section["paragraphs"]:
            story.append(Paragraph(value, styles["body"]))
        for value in section["bullets"]:
            story.append(Paragraph("• " + value, styles["bullet"]))

    if payload.get("figures"):
        story.append(Paragraph("Figuras del documento fuente", styles["h1"]))
        for figure in payload["figures"]:
            figure_path = Path(figure["path"])
            if not figure_path.exists():
                continue
            image = Image(str(figure_path))
            image._restrictSize(16 * cm, 10 * cm)
            image.hAlign = "CENTER"
            story.extend([
                image,
                Spacer(1, 0.12 * cm),
                Paragraph(figure["caption"], styles["small"]),
                Spacer(1, 0.3 * cm),
            ])

    payload = {**payload, "calculations": [
        calculation for calculation in payload.get("calculations") or []
        if str(calculation.get("status") or "").lower() == "calculado"
    ]}
    if payload["calculations"]:
        story.append(Paragraph("Calculos realizados", styles["h1"]))
        rows = [[
            Paragraph("Calculo", styles["small"]), Paragraph("Formula", styles["small"]),
            Paragraph("Resultado", styles["small"]), Paragraph("Fuente / estado", styles["small"]),
        ]]
        for calculation in payload["calculations"]:
            rows.append([
                Paragraph(calculation["label"], styles["small"]),
                Paragraph(calculation["formula"], styles["small"]),
                Paragraph(calculation["result"], styles["small"]),
                Paragraph(f"{calculation['source']}<br/>Estado: {calculation['status']}", styles["small"]),
            ])
        table = Table(rows, colWidths=[3.4 * cm, 4.5 * cm, 3.0 * cm, 5.1 * cm], repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + GREEN)),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), bold),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD8D2")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(table)

    for title, values, color in (
        ("Fuentes", payload["sources"], MUTED),
    ):
        if not values:
            continue
        story.append(Paragraph(title, styles["h1"]))
        for value in values:
            story.append(Paragraph("• " + value, styles["bullet"]))
    story.extend([
        Spacer(1, 0.5 * cm),
        Paragraph("<b>Ing. Agr. Lucas Estecho</b><br/>M.P. 2009 LE", styles["body"]),
    ])
    callback = lambda canvas, document: _pdf_header_footer(
        canvas, document, logo_path=logo_path, report_title=payload["title"]
    )
    doc.build(story, onFirstPage=callback, onLaterPages=callback)
    return output_path


def generate_consulting_report(
    *,
    event: dict,
    draft: dict,
    crew_result: dict,
    source_text: str,
    assets: list[dict] | None,
    output_dir: str | Path,
    logo_path: str | None = None,
    openai_client=None,
    model: str | None = None,
    reasoning_effort: str | None = None,
    allow_fallback: bool = False,
) -> GeneratedConsultingReport | None:
    playbook = detect_report_playbook(source_text or draft.get("summary") or "")
    if playbook is None:
        return None
    selected_model = model or os.environ.get("CAPATAZ_REPORT_MODEL", "gpt-5.6-sol")
    payload = _generate_payload(
        playbook,
        event,
        draft,
        crew_result,
        source_text,
        list(assets or []),
        openai_client=openai_client,
        model=selected_model,
        reasoning_effort=reasoning_effort or os.environ.get("CAPATAZ_REPORT_REASONING", "high"),
        allow_fallback=allow_fallback,
    )
    root = Path(output_dir) / str(event.get("id") or "sin-evento")
    root.mkdir(parents=True, exist_ok=True)
    payload["figures"] = _materialize_docx_figures(list(assets or []), root)
    stamp = datetime.now().strftime("%Y-%m-%d")
    base = _safe_filename(f"{playbook.title}_{event.get('client_name') or draft.get('field_name') or ''}_{stamp}")
    pdf_path = root / f"{base}.pdf"
    docx_path = root / f"{base}.docx"
    create_consulting_pdf(payload, str(pdf_path), event=event, playbook=playbook, logo_path=logo_path)
    create_consulting_docx(payload, str(docx_path), event=event, playbook=playbook, logo_path=logo_path)
    return GeneratedConsultingReport(
        playbook_key=playbook.key,
        title=payload["title"],
        status=payload["status"],
        pdf_path=str(pdf_path),
        docx_path=str(docx_path),
        missing_data=tuple(payload["missing_data"]),
        model=selected_model,
    )
